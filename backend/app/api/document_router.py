# backend/app/api/document_router.py
"""
文档生成 API 路由
用于将 AI 生成的内容（Markdown 格式）转换为 Word 或 PDF 文档

技术方案：AI 结构化提取 + 专业模板渲染
"""
import os
import uuid
import re
import shutil
import logging
from typing import Optional
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.database import get_db
from app.models.contract import ContractDoc, ContractStatus
from app.utils.office_utils import OfficeTokenManager
from app.services.document_preprocessor import get_preprocessor
from app.services.document_structurer import get_structurer
from app.services.document_renderer import get_renderer

logger = logging.getLogger(__name__)

router = APIRouter(
    prefix="/api/document",
    tags=["Document Generation"]
)

UPLOAD_DIR = "storage/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


class GenerateDocumentRequest(BaseModel):
    """文档生成请求"""
    content: str  # Markdown 或纯文本内容
    output_format: str = "docx"  # docx 或 pdf
    filename: Optional[str] = None  # 可选的文件名
    doc_type: Optional[str] = None  # 文档类型提示: contract/letter/judicial
    use_ai_structuring: bool = True  # 是否使用 AI 结构化（默认启用）


def parse_markdown_to_docx(content: str) -> 'docx.Document':
    """
    将 Markdown 格式内容解析为 Word 文档

    Args:
        content: Markdown 格式的文本

    Returns:
        python-docx Document 对象
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    # 预处理：清理常见的 AI 输出特殊字符
    # 1. 替换 HTML <br> 标签为换行
    content = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'<BR\s*/?>', '\n', content)

    # 2. 替换其他 HTML 标签
    content = re.sub(r'<[^>]+>', '', content)

    # 3. 替换 ≈ 为约或删除
    content = re.sub(r'≈', '约', content)

    # 4. 处理连续的星号（粗体标记）
    content = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', content)  # ***粗斜体***
    content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)      # **粗体**
    content = re.sub(r'\*(.+?)\*', r'\1', content)          # *斜体*

    lines = content.split('\n')
    in_list = False
    list_level = 0

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 空行
        if not line:
            if in_list:
                in_list = False
            else:
                doc.add_paragraph()  # 添加空段落
            i += 1
            continue

        # 一级标题 # 标题
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_list = False

        # 二级标题 ## 标题
        elif line.startswith('## '):
            title = line[3:].strip()
            doc.add_heading(title, level=2)
            in_list = False

        # 三级标题 ### 标题
        elif line.startswith('### '):
            title = line[4:].strip()
            doc.add_heading(title, level=3)
            in_list = False

        # 无序列表 - 项目 或 * 项目
        elif line.startswith(('- ', '* ', '• ')):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            in_list = True

        # 有序列表 1. 项目
        elif re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line).strip()
            p = doc.add_paragraph(text, style='List Number')
            in_list = True

        # 普通段落
        else:
            # 清理剩余的特殊字符
            text = line.strip()

            if in_list:
                # 如果之前在列表中，现在不是列表项，结束列表
                in_list = False

            # 检查是否是特殊的合同格式（如：甲方：xxx）
            if '：' in line or ':' in line:
                parts = re.split(r'[：:]', line, 1)
                if len(parts) == 2:
                    label = parts[0].strip()
                    value = parts[1].strip()
                    p = doc.add_paragraph()
                    run = p.add_run(f"{label}：")
                    run.bold = True
                    p.add_run(value)
                else:
                    doc.add_paragraph(text)
            else:
                doc.add_paragraph(text)

        i += 1

    return doc


def convert_docx_to_pdf(docx_path: str) -> str:
    """
    将 DOCX 转换为 PDF

    Args:
        docx_path: DOCX 文件路径

    Returns:
        PDF 文件路径
    """
    try:
        from app.services.converter import convert_to_pdf_via_onlyoffice

        filename = os.path.basename(docx_path)
        success, pdf_content = convert_to_pdf_via_onlyoffice(filename)

        if success:
            pdf_path = docx_path.replace('.docx', '.pdf')
            with open(pdf_path, 'wb') as f:
                f.write(pdf_content)
            return pdf_path
        else:
            raise Exception("PDF 转换失败")
    except Exception as e:
        raise Exception(f"PDF 转换失败: {str(e)}")


@router.post("/generate-from-content")
async def generate_document_from_content(
    request: GenerateDocumentRequest,
    db: Session = Depends(get_db)
):
    """
    从 AI 生成的内容创建 Word 或 PDF 文档

    新方案：AI 结构化提取 + 专业模板渲染
    1. 使用 AI 分析文本，提取文档结构
    2. 使用专业模板渲染为规范格式
    3. 如需要，转换为 PDF
    4. 保存文档记录
    5. 返回预览和下载链接

    兜底方案：如果 AI 不可用，使用原有的 Markdown 解析
    """
    try:
        # 尝试使用 AI 结构化渲染
        used_ai_structuring = False
        structure = None
        final_title = None

        if request.use_ai_structuring:
            structurer = get_structurer()
            renderer = get_renderer()

            if structurer.is_available():
                # 使用 AI 提取结构
                structure = structurer.extract_structure(
                    request.content,
                    request.doc_type
                )

                if structure and structure.title:
                    # 使用 AI 提取的标题作为文件名
                    final_title = structure.title

                    # 根据文档类型渲染
                    if structure.doc_type == "letter" or request.doc_type == "letter":
                        # 先生成临时文件路径
                        temp_docx_path = os.path.join(UPLOAD_DIR, f"temp_{uuid.uuid4().hex[:8]}.docx")
                        success = renderer.render_letter(structure, temp_docx_path)
                    else:
                        # 默认使用合同格式
                        temp_docx_path = os.path.join(UPLOAD_DIR, f"temp_{uuid.uuid4().hex[:8]}.docx")
                        success = renderer.render_contract(structure, temp_docx_path)

                    if success:
                        used_ai_structuring = True
                        logging.info(f"使用 AI 结构化渲染成功: {temp_docx_path}")
                    else:
                        logging.warning("AI 渲染失败，将使用 Markdown 兜底方案")
                    # 渲染后会重命名文件，这里暂时记录临时路径
                    docx_path = temp_docx_path
                else:
                    logging.warning("AI 结构提取失败，将使用 Markdown 兜底方案")
            else:
                logging.info("AI 服务不可用，使用 Markdown 兜底方案")

        # 生成最终文件名
        if not request.filename:
            ext = 'docx' if request.output_format == 'docx' else 'pdf'

            # 如果成功提取了标题，使用标题作为文件名
            if final_title:
                # 清理文件名中的非法字符
                import re
                safe_title = re.sub(r'[<>:"/\\|?*]', '', final_title)
                safe_title = safe_title.strip()
                # 限制文件名长度
                if len(safe_title) > 50:
                    safe_title = safe_title[:50]
                request.filename = f"{safe_title}.{ext}"
                logging.info(f"[DEBUG] 使用 AI 标题作为文件名: {request.filename}")
            else:
                request.filename = f"generated_{uuid.uuid4().hex[:8]}.{ext}"
                logging.info(f"[DEBUG] 使用随机文件名: {request.filename}")
        else:
            logging.info(f"[DEBUG] 使用请求中的文件名: {request.filename}")

        # 确保文件名有正确的扩展名
        if not request.filename.endswith('.docx') and not request.filename.endswith('.pdf'):
            request.filename = f"{request.filename}.{request.output_format}"

        base_filename = request.filename.rsplit('.', 1)[0]
        docx_filename = f"{base_filename}.docx"
        final_docx_path = os.path.join(UPLOAD_DIR, docx_filename)

        # 如果使用了 AI 渲染，需要重命名临时文件
        if used_ai_structuring and 'docx_path' in locals() and docx_path != final_docx_path:
            import shutil
            shutil.move(docx_path, final_docx_path)
            docx_path = final_docx_path
        elif not used_ai_structuring:
            docx_path = final_docx_path

        # 兜底方案：使用原有的 Markdown 解析
        if not used_ai_structuring:
            doc = parse_markdown_to_docx(request.content)
            doc.save(docx_path)

        # 如果需要 PDF，进行转换
        pdf_path = None
        if request.output_format == 'pdf':
            try:
                pdf_path = convert_docx_to_pdf(docx_path)
            except Exception as e:
                # PDF 转换失败，返回 docx 并警告
                pdf_path = None

        # 保存文档记录到数据库
        db_contract = ContractDoc(
            title=request.filename,
            status=ContractStatus.DRAFT.value,
            original_file_path=docx_path,
            pdf_converted_path=pdf_path or docx_path,
            owner_id=1  # 临时，后面接用户系统
        )
        db.add(db_contract)
        db.commit()
        db.refresh(db_contract)

        # 生成 OnlyOffice 配置用于预览
        file_url = f"http://backend:8000/storage/uploads/{docx_filename}"
        from datetime import datetime

        config = {
            "document": {
                "fileType": "docx",
                "key": str(db_contract.id) + "_generated_" + str(int(datetime.now().timestamp())),
                "title": request.filename,
                "url": file_url,
            },
            "editorConfig": {
                "mode": "view",
                "user": {"id": "1", "name": "用户"}
            }
        }
        token = OfficeTokenManager.create_token(config)

        return {
            "success": True,
            "contract_id": db_contract.id,
            "filename": request.filename,
            "docx_path": docx_path,
            "pdf_path": pdf_path,
            "status": "completed",
            "used_ai_structuring": used_ai_structuring,
            # 预览相关
            "preview_url": f"/api/contract/{db_contract.id}/onlyoffice-config",
            # 下载相关
            "download_docx_url": f"/api/contract/{db_contract.id}/download",
            "download_pdf_url": f"/api/contract/{db_contract.id}/download-revised" if pdf_path else None,
            # OnlyOffice 配置
            "config": config,
            "token": token
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"文档生成失败: {str(e)}")


@router.get("/preview/{contract_id}")
def get_document_preview(contract_id: int, db: Session = Depends(get_db)):
    """
    获取文档预览配置（使用 OnlyOffice）

    返回 OnlyOffice 编辑器配置和 JWT token
    """
    contract = db.query(ContractDoc).filter(ContractDoc.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="文档不存在")

    file_url = f"http://backend:8000/storage/uploads/{os.path.basename(contract.original_file_path)}"

    from datetime import datetime

    config = {
        "document": {
            "fileType": "docx",
            "key": str(contract.id) + "_preview_" + str(int(datetime.now().timestamp())),
            "title": contract.title,
            "url": file_url,
        },
        "editorConfig": {
            "mode": "view",
            "user": {"id": "1", "name": "用户"}
        }
    }

    token = OfficeTokenManager.create_token(config)
    return {"config": config, "token": token}


@router.get("/preview/by-filename/{filename}")
def get_document_preview_by_filename(filename: str, db: Session = Depends(get_db)):
    """
    根据文件名获取文档预览配置（用于文档预处理等没有 contract_id 的场景）

    返回 OnlyOffice 编辑器配置和 JWT token
    """
    # 检查文件是否存在
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")

    file_url = f"http://backend:8000/storage/uploads/{filename}"

    from datetime import datetime

    config = {
        "document": {
            "fileType": "docx" if filename.endswith('.docx') else 'pdf',
            "key": filename + "_preview_" + str(int(datetime.now().timestamp())),
            "title": filename,
            "url": file_url,
        },
        "editorConfig": {
            "mode": "view",
            "user": {"id": "1", "name": "用户"}
        }
    }

    token = OfficeTokenManager.create_token(config)
    return {"config": config, "token": token}


@router.post("/process-file-to-standard")
async def process_file_to_standard(
    file: UploadFile = File(...),
    output_format: str = Form("docx"),
    document_type: str = Form("contract"),
    db: Session = Depends(get_db)
):
    """
    将不规范的文件处理为标准的合同或函件格式

    流程：
    1. 保存上传的文件
    2. 使用文档预处理中心转换为 Word 格式
    3. 根据文档类型（合同/函件）进行格式化处理
    4. 如需要，转换为 PDF
    5. 保存文档记录
    6. 返回预览和下载链接
    """
    try:
        # 保存上传文件
        file_ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
        unique_name = f"upload_{uuid.uuid4().hex}.{file_ext}"
        file_path = os.path.join(UPLOAD_DIR, unique_name)

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 使用预处理中心转换为 docx
        preprocessor = get_preprocessor()
        result, output_path, metadata = preprocessor.preprocess(
            file_path=file_path,
            delete_original=False
        )

        if result != "success" and result != "already_converted":
            raise HTTPException(status_code=400, detail=f"文件转换失败: {output_path}")

        # 使用转换后的文件路径
        docx_path = output_path

        # 如果需要 PDF，进行转换
        pdf_path = None
        if output_format == 'pdf':
            try:
                pdf_path = convert_docx_to_pdf(docx_path)
            except Exception as e:
                # PDF 转换失败，返回 docx 并警告
                pdf_path = None

        # 生成文档标题
        doc_type_name = "标准合同" if document_type == "contract" else "正式函件"
        base_name = os.path.splitext(file.filename)[0]
        filename = f"{base_name}_{doc_type_name}.{output_format}"

        # 保存文档记录到数据库
        db_contract = ContractDoc(
            title=filename,
            status=ContractStatus.DRAFT.value,
            original_file_path=docx_path,
            pdf_converted_path=pdf_path or docx_path,
            owner_id=1  # 临时，后面接用户系统
        )
        db.add(db_contract)
        db.commit()
        db.refresh(db_contract)

        # 生成 OnlyOffice 配置用于预览
        docx_filename = os.path.basename(docx_path)
        file_url = f"http://backend:8000/storage/uploads/{docx_filename}"
        from datetime import datetime

        config = {
            "document": {
                "fileType": "docx",
                "key": str(db_contract.id) + "_standardized_" + str(int(datetime.now().timestamp())),
                "title": filename,
                "url": file_url,
            },
            "editorConfig": {
                "mode": "view",
                "user": {"id": "1", "name": "用户"}
            }
        }
        token = OfficeTokenManager.create_token(config)

        return {
            "success": True,
            "contract_id": db_contract.id,
            "filename": filename,
            "docx_path": docx_path,
            "pdf_path": pdf_path,
            "status": "completed",
            "document_type": document_type,
            # 预览相关
            "preview_url": f"/api/document/preview/{db_contract.id}",
            # 下载相关
            "download_docx_url": f"/api/contract/{db_contract.id}/download",
            "download_pdf_url": f"/api/contract/{db_contract.id}/download-revised" if pdf_path else None,
            # OnlyOffice 配置
            "config": config,
            "token": token,
            # 元数据
            "metadata": metadata
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")
