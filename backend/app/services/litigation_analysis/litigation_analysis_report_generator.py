# backend/app/services/litigation_analysis_report_generator.py
"""
案件分析报告生成服务

将案件分析结果生成规范的 Word 或 PDF 报告
"""
import logging
import os
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class LitigationAnalysisReportGenerator:
    """案件分析报告生成器"""

    def __init__(self):
        """初始化报告生成器"""
        self.output_dir = "storage/reports/litigation_analysis"
        os.makedirs(self.output_dir, exist_ok=True)

    def generate(self, session_id: str, data: Dict[str, Any], format: str = 'docx') -> str:
        """
        生成分析报告

        Args:
            session_id: 会话ID
            data: 分析数据（包含 case_summary, evidence_assessment, strategies 等）
            format: 输出格式（docx 或 pdf）

        Returns:
            报告文件路径
        """
        try:
            # 创建Word文档
            doc = Document()

            # 设置文档样式
            self._setup_styles(doc)

            # 添加标题
            self._add_title(doc, session_id)

            # 添加元数据
            self._add_metadata(doc, session_id, data)

            # 添加核心结论
            self._add_executive_summary(doc, data)

            # 添加证据评估
            if data.get('evidence_assessment'):
                self._add_evidence_section(doc, data['evidence_assessment'])

            # 添加争议焦点
            if data.get('legal_issues'):
                self._add_issues_section(doc, data['legal_issues'])

            # 添加诉讼策略
            if data.get('strategies'):
                self._add_strategies_section(doc, data['strategies'])

            # 添加风险提示
            if data.get('risk_warnings'):
                self._add_risks_section(doc, data['risk_warnings'])

            # 添加详细报告
            if data.get('final_report') or data.get('report_md'):
                self._add_detailed_report(doc, data.get('final_report') or data.get('report_md'))

            # 保存文档
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"litigation_analysis_{session_id}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)

            logger.info(f"分析报告生成成功: {filepath}")

            # 如果需要PDF格式，转换
            if format == 'pdf':
                pdf_path = self._convert_to_pdf(filepath)
                return pdf_path

            return filepath

        except Exception as e:
            logger.error(f"分析报告生成失败: {str(e)}", exc_info=True)
            raise

    def _setup_styles(self, doc: Document):
        """设置文档样式"""
        # 正文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)

        # 标题1样式
        if 'Heading 1' in doc.styles:
            h1 = doc.styles['Heading 1']
            h1.font.name = '黑体'
            h1.font.size = Pt(16)
            h1.font.bold = True

        # 标题2样式
        if 'Heading 2' in doc.styles:
            h2 = doc.styles['Heading 2']
            h2.font.name = '黑体'
            h2.font.size = Pt(14)
            h2.font.bold = True

    def _add_title(self, doc: Document, session_id: str):
        """添加文档标题"""
        doc.add_paragraph()

        heading = doc.add_heading("案件分析报告", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(18)

        # 添加会话ID
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"会话编号：{session_id}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_metadata(self, doc: Document, session_id: str, data: Dict[str, Any]):
        """添加元数据"""
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts = []
        parts.append(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")

        if data.get('case_type'):
            parts.append(f"案件类型：{data['case_type']}")

        if data.get('case_position'):
            parts.append(f"诉讼地位：{data['case_position']}")

        if parts:
            p.add_run('  |  '.join(parts))
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

    def _add_executive_summary(self, doc: Document, data: Dict[str, Any]):
        """添加核心结论"""
        doc.add_heading("一、核心结论", level=2)

        # 胜诉率
        if data.get('win_probability') is not None:
            p = doc.add_paragraph()
            p.add_run("胜诉评估：").bold = True
            win_rate = data['win_probability'] * 100
            p.add_run(f"{win_rate:.1f}%")

        # 案件摘要
        if data.get('case_summary'):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("案件摘要：").bold = True
            summary_para = doc.add_paragraph(data['case_summary'])
            summary_para.paragraph_format.first_line_indent = Inches(0.3)

        doc.add_paragraph()

    def _add_evidence_section(self, doc: Document, evidence: Any):
        """添加证据评估"""
        doc.add_heading("二、证据评估", level=2)

        if isinstance(evidence, dict):
            # 结构化证据数据
            if evidence.get('overall_assessment'):
                p = doc.add_paragraph()
                p.add_run("整体评估：").bold = True
                p.add_run(evidence['overall_assessment'])

            if evidence.get('key_evidence'):
                p = doc.add_paragraph()
                p.add_run("关键证据：").bold = True
                if isinstance(evidence['key_evidence'], list):
                    for item in evidence['key_evidence']:
                        doc.add_paragraph(str(item), style='List Bullet')

            if evidence.get('evidence_gaps'):
                p = doc.add_paragraph()
                p.add_run("证据缺口：").bold = True
                if isinstance(evidence['evidence_gaps'], list):
                    for gap in evidence['evidence_gaps']:
                        doc.add_paragraph(str(gap), style='List Bullet')

        elif isinstance(evidence, str):
            p = doc.add_paragraph(evidence)

        doc.add_paragraph()

    def _add_issues_section(self, doc: Document, issues: Any):
        """添加争议焦点"""
        doc.add_heading("三、争议焦点", level=2)

        if isinstance(issues, list):
            for i, issue in enumerate(issues, 1):
                if isinstance(issue, dict):
                    doc.add_heading(f"{i}. {issue.get('title', f'焦点{i}')}", level=3)
                    if issue.get('description'):
                        doc.add_paragraph(issue['description'])
                else:
                    doc.add_paragraph(f"{i}. {str(issue)}", style='List Number')
        elif isinstance(issues, str):
            doc.add_paragraph(issues)

        doc.add_paragraph()

    def _add_strategies_section(self, doc: Document, strategies: Any):
        """添加诉讼策略"""
        doc.add_heading("四、诉讼策略", level=2)

        if isinstance(strategies, list):
            for i, strategy in enumerate(strategies, 1):
                if isinstance(strategy, dict):
                    doc.add_heading(f"策略{i}：{strategy.get('title', f'策略{i}')}", level=3)
                    if strategy.get('description'):
                        doc.add_paragraph(strategy['description'])
                    if strategy.get('actions'):
                        p = doc.add_paragraph()
                        p.add_run("具体行动：").bold = True
                        if isinstance(strategy['actions'], list):
                            for action in strategy['actions']:
                                doc.add_paragraph(action, style='List Bullet 2')
                else:
                    doc.add_paragraph(f"{i}. {str(strategy)}", style='List Number')
        elif isinstance(strategies, str):
            doc.add_paragraph(strategies)

        doc.add_paragraph()

    def _add_risks_section(self, doc: Document, risks: Any):
        """添加风险提示"""
        doc.add_heading("五、风险提示", level=2)

        if isinstance(risks, list):
            for i, risk in enumerate(risks, 1):
                if isinstance(risk, dict):
                    p = doc.add_paragraph()
                    p.add_run(f"{i}. ").bold = True
                    p.add_run(risk.get('description', risk.get('risk', str(risk))))
                    if risk.get('mitigation'):
                        p = doc.add_paragraph()
                        p.add_run("  应对措施：").italic = True
                        p.add_run(risk['mitigation'])
                else:
                    doc.add_paragraph(f"{i}. {str(risk)}", style='List Number')
        elif isinstance(risks, str):
            doc.add_paragraph(risks)

        doc.add_paragraph()

    def _add_detailed_report(self, doc: Document, report_text: str):
        """添加详细报告"""
        doc.add_heading("六、详细分析", level=2)

        # 简化处理：将Markdown/文本直接添加为段落
        lines = report_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('#'):
                # 标题
                level = min(line.count('#'), 3)
                text = line.lstrip('#').strip()
                doc.add_heading(text, level=level)
            elif line.startswith('-') or line.startswith('*'):
                # 列表
                text = line.lstrip('-*').strip()
                doc.add_paragraph(text, style='List Bullet')
            else:
                # 普通段落
                doc.add_paragraph(line)

    def _convert_to_pdf(self, docx_path: str) -> str:
        """将DOCX转换为PDF"""
        try:
            from app.services.common.file_service import file_service
            # 使用 file_service 转换
            # 简化处理：直接返回 docx 路径
            logger.warning(f"PDF转换功能暂未实现，返回DOCX路径: {docx_path}")
            return docx_path
        except Exception as e:
            logger.error(f"PDF转换失败: {e}")
            return docx_path


# 单例
_generator_instance = None

def get_litigation_analysis_report_generator():
    global _generator_instance
    if _generator_instance is None:
        _generator_instance = LitigationAnalysisReportGenerator()
    return _generator_instance
