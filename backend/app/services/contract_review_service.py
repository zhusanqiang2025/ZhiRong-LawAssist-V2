# backend/app/services/contract_review_service.py
import os
import json
import httpx
import logging
import traceback
import subprocess
import shutil
import time
from docx import Document as DocxDocument
try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except Exception:
    PYPANDOC_AVAILABLE = False
from sqlalchemy.orm import Session
from pydantic import BaseModel, Field, create_model

from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.prompts import ChatPromptTemplate

# 数据库模型
from app.models.contract import (
    ContractDoc,
    ContractReviewItem,
    ContractStatus,
)
from app.models.rule import ReviewRule  # ✅ 从 contract.py 移动到独立的 rule.py

# Pydantic Schemas（现在已经齐全）
from app.schemas import ContractMetadataSchema, ReviewOutput

# PDF 解析服务
from app.services.pdf_service import PdfService

# ⭐ 新增：规则组装器
from app.services.contract_review.rule_assembler import rule_assembler


# === 配置读取 ===
# ⭐ 使用 Qwen3-Thinking 模型进行元数据提取和合同审查
QWEN3_API_KEY = os.getenv("QWEN3_THINKING_API_KEY")
QWEN3_API_URL = os.getenv("QWEN3_THINKING_API_URL")
QWEN3_MODEL = os.getenv("QWEN3_THINKING_MODEL", "Qwen3-235B-A22B-Thinking-2507")
QWEN3_ENABLED = os.getenv("QWEN3_THINKING_ENABLED", "true").lower() == "true"

if not QWEN3_API_KEY:
    print("警告: QWEN3_THINKING_API_KEY 未配置，合同审查功能将不可用")
    QWEN3_API_KEY = "fake-key-for-startup"  # 设置一个假密钥以允许启动

if not QWEN3_API_URL:
    QWEN3_API_URL = "https://api.deepseek.com"

print(f"[配置] Qwen3-Thinking 模型: {QWEN3_MODEL}, 启用: {QWEN3_ENABLED}")


class ContractReviewService:
    def __init__(self, db: Session):
        self.db = db
        self.llm = self._init_llm()
        self.logger = logging.getLogger(__name__)  # 初始化 logger

    def _init_llm(self):
        """初始化兼容 OpenAI 格式的 Qwen3-Thinking 模型"""
        http_client = httpx.Client(
            verify=False,      # 国内环境常需关闭 SSL 验证
            trust_env=False,   # 防止读取系统代理
        )
        return ChatOpenAI(
            model=QWEN3_MODEL,
            api_key=QWEN3_API_KEY,
            base_url=QWEN3_API_URL,
            temperature=0.1,
            http_client=http_client,
            max_tokens=4096,
        )

    def extract_metadata(self, contract_id: int) -> ContractMetadataSchema | None:
        """
        阶段一：从合同提取基本元数据（名称、当事人、金额等）

        增强功能：
        - 使用知识图谱数据库进行合同类型精确匹配
        - 支持别名匹配和关键词搜索
        - 结合 LLM 验证和补充其他字段
        """
        import time
        start_time = time.time()
        self.logger.info(f"[元数据提取] 开始提取合同 {contract_id} 的元数据")

        contract = self.db.query(ContractDoc).filter(ContractDoc.id == contract_id).first()
        if not contract:
            raise ValueError("合同不存在")

        # 优先直接从原始 docx/doc 提取文本（避免依赖 OnlyOffice 转换失败）
        text = ""
        structured_meta: dict = {}
        try:
            self.logger.info(f"[元数据提取] 步骤1: 从原始文件提取文本...")
            file_parse_start = time.time()
            orig = (contract.original_file_path or "").strip()
            if orig:
                lower = orig.lower()
                if lower.endswith('.doc'):
                    # 尝试用 pypandoc 转换到 docx
                    orig_to_read = orig
                    if PYPANDOC_AVAILABLE:
                        try:
                            converted = orig + '.converted.docx'
                            pypandoc.convert_file(orig, 'docx', outputfile=converted)
                            orig_to_read = converted
                        except Exception as e:
                            print('pypandoc doc -> docx 转换失败，错误：', e)
                            traceback.print_exc()
                            orig_to_read = orig

                    # 如果仍未转换，尝试使用 LibreOffice (soffice) 命令行转换
                    if orig_to_read == orig:
                        try:
                            soffice_path = shutil.which('soffice')
                            if soffice_path:
                                outdir = os.path.dirname(orig)
                                cmd = [soffice_path, '--headless', '--convert-to', 'docx', '--outdir', outdir, orig]
                                print('尝试使用 LibreOffice 转换 doc -> docx，命令：', ' '.join(cmd))
                                res = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                                if res.returncode == 0:
                                    converted = os.path.join(outdir, os.path.splitext(os.path.basename(orig))[0] + '.docx')
                                    if os.path.exists(converted):
                                        orig_to_read = converted
                                        print('LibreOffice 转换成功：', converted)
                                else:
                                    print('LibreOffice 转换返回非零状态，stderr:', res.stderr)
                        except Exception as e:
                            print('使用 LibreOffice 转换时发生异常：', e)
                            traceback.print_exc()
                else:
                    orig_to_read = orig

                # 若为 docx，尝试直接解析结构化信息与纯文本
                if orig_to_read.lower().endswith('.docx'):
                    try:
                        doc = DocxDocument(orig_to_read)
                        paragraphs = [p.text for p in doc.paragraphs if p.text]
                        text = "\n".join(paragraphs)
                        # 提取结构化信息（表格、签署方、标题等）
                        structured_meta = self._extract_structured_from_docx(doc)
                    except Exception as e:
                        self.logger.warning(f"[元数据提取] 读取 DOCX 失败: {e}")
                        text = ""

            self.logger.info(f"[元数据提取] 文件解析完成，耗时: {time.time() - file_parse_start:.2f}秒，文本长度: {len(text)}字符")

        except Exception as e:
            self.logger.error(f"[元数据提取] 尝试从原始文件读取时发生异常: {e}")
            traceback.print_exc()
            text = ""

        # 如果没有从 docx 成功提取，再尝试 PDF 解析
        if not text:
            self.logger.info(f"[元数据提取] DOCX提取失败，尝试从PDF解析...")
            if not contract.pdf_converted_path:
                raise ValueError("合同文件未就绪或 PDF 未转换")
            text, _ = PdfService.parse_pdf(contract.pdf_converted_path)
            if not text:
                raise ValueError("无法从 PDF 或 DOCX 提取文本内容")

        # === 新增：基于法律特征的知识图谱匹配 ===
        self.logger.info(f"[元数据提取] 步骤2: 知识图谱匹配...")
        kg_match_start = time.time()
        from app.services.common.contract_knowledge_db_service import contract_knowledge_db_service
        import re

        contract_type_match = None
        matched_features = None
        match_source = None

        try:
            # 1. 尝试从标题中精确匹配
            if structured_meta.get('contract_name'):
                title = structured_meta['contract_name']
                match = contract_knowledge_db_service.get_by_name(title)
                if match:
                    contract_type_match = match['name']
                    matched_features = match.get('legal_features', {})
                    match_source = 'title_exact_match'
                    print(f"[知识图谱] 从标题精确匹配到合同类型: {contract_type_match}")

            # 2. 如果标题未匹配，使用 LLM 提取法律特征，然后基于特征匹配
            if not contract_type_match:
                print(f"[知识图谱] 标题未匹配，开始基于法律特征匹配...")

                # === 动态构建法律特征提取 Schema ===
                # 从知识图谱数据库中获取所有法律特征字段，动态生成 Schema
                all_types = contract_knowledge_db_service.get_all()

                # 收集所有法律特征字段（动态适应知识图谱变化）
                legal_feature_fields = {}
                feature_descriptions = {}

                for kg_type in all_types:
                    features = kg_type.get('legal_features', {})
                    for key, value in features.items():
                        if key not in legal_feature_fields:
                            # 从第一个遇到的值推断字段类型
                            if isinstance(value, bool):
                                legal_feature_fields[key] = bool
                                feature_descriptions[key] = f"布尔值：{key}"
                            elif isinstance(value, list):
                                legal_feature_fields[key] = list
                                feature_descriptions[key] = f"列表：{key}"
                            else:
                                legal_feature_fields[key] = str
                                feature_descriptions[key] = f"{key}"

                # 构建字段定义字典（用于 create_model）
                field_definitions = {}
                for field_name, field_type in legal_feature_fields.items():
                    field_definitions[field_name] = (field_type, Field(default='', description=feature_descriptions[field_name]))
                # 添加建议合同类型字段
                field_definitions['suggested_contract_type'] = (str, Field(description="LLM 推测的合同类型名称"))

                # 使用 create_model 动态创建 Pydantic 模型
                LegalFeaturesExtraction = create_model(
                    'LegalFeaturesExtraction',
                    __base__=BaseModel,
                    **field_definitions
                )

                # 构建 Schema 说明文本（用于 Prompt）
                schema_descriptions = []
                for field_name, description in feature_descriptions.items():
                    # 从知识图谱中找一些示例值
                    example_values = set()
                    for kg_type in all_types[:5]:  # 只看前5个类型
                        val = kg_type.get('legal_features', {}).get(field_name)
                        if val:
                            if isinstance(val, list):
                                example_values.add(f"{str(val[:2])}..." if len(val) > 2 else str(val))
                            else:
                                example_values.add(str(val)[:30])
                    example_text = f" 示例: {', '.join(list(example_values)[:3])}" if example_values else ""
                    schema_descriptions.append(f"- **{field_name}**: {description}{example_text}")

                schema_text = "\n".join(schema_descriptions)

                # 使用 LLM 提取法律特征（使用动态 Schema）
                feature_parser = PydanticOutputParser(pydantic_object=LegalFeaturesExtraction)
                feature_prompt = ChatPromptTemplate.from_messages([
                    ("system", f"""你是一名专业的合同法律特征分析专家。请分析合同文本，提取以下法律特征：

**需要提取的法律特征字段：**
{schema_text}

**输出要求：**
- 对于无法确定或文本中没有明确提到的字段，请填写"未提及"或空字符串
- suggested_contract_type 字段请根据你的专业知识推测最可能的合同类型

{{format_instructions}}"""),
                    ("user", "合同文本（前2000字）：\n{text}")
                ])

                feature_chain = feature_prompt | self.llm | feature_parser

                try:
                    features_result = feature_chain.invoke({
                        "text": text[:2000],
                        "format_instructions": feature_parser.get_format_instructions()
                    })

                    # 打印提取的关键特征
                    print(f"[法律特征提取] 提取到 {len([k for k in dir(features_result) if not k.startswith('_')])} 个特征字段")
                    for key in ['transaction_nature', 'contract_object', 'transaction_characteristics', 'usage_scenario']:
                        if hasattr(features_result, key):
                            val = getattr(features_result, key)
                            if val:
                                print(f"[法律特征提取] {key}: {val}")
                    print(f"[法律特征提取] LLM推测类型: {features_result.suggested_contract_type}")

                    # === 使用提取的特征进行匹配（动态匹配所有字段）===
                    best_match = None
                    best_score = 0
                    match_details = {}

                    # 字段权重配置：越具体的字段权重越高
                    field_weights = {
                        'transaction_nature': 5,      # 通用，权重低
                        'contract_object': 5,          # 通用，权重低
                        'stance': 3,                   # 很通用，权重更低
                        'consideration_type': 5,       # 通用
                        'consideration_detail': 8,     # 较具体
                        'transaction_characteristics': 15,  # 非常重要！能区分合同具体内容
                        'usage_scenario': 20,          # 最重要！使用场景最能区分合同类型
                        'legal_basis': 2,              # 参考性弱
                    }

                    for kg_type in all_types:
                        kg_features = kg_type.get('legal_features', {})
                        score = 0
                        details = {}

                        # 对每个提取的特征进行匹配
                        for field_name in legal_feature_fields.keys():
                            extracted_val = getattr(features_result, field_name, None)
                            kg_val = kg_features.get(field_name)

                            if not extracted_val or not kg_val:
                                continue

                            # 跳过"未提及"或空值
                            if isinstance(extracted_val, str) and extracted_val in ['未提及', '', '无']:
                                continue

                            # 获取字段权重
                            weight = field_weights.get(field_name, 10)

                            # 根据字段类型计算匹配分数
                            if isinstance(extracted_val, str) and isinstance(kg_val, str):
                                # 字符串匹配
                                if extracted_val == kg_val:
                                    score += weight * 2  # 完全匹配，双倍权重
                                    details[f'{field_name}_match'] = 'exact'
                                elif extracted_val in kg_val or kg_val in extracted_val:
                                    score += weight * 0.5  # 部分匹配，半倍权重
                                    details[f'{field_name}_match'] = 'partial'

                            elif isinstance(extracted_val, list) and isinstance(kg_val, list):
                                # 列表匹配：计算重叠度
                                overlap = len(set(extracted_val) & set(kg_val))
                                if overlap > 0:
                                    score += overlap * weight * 0.3
                                    details[f'{field_name}_overlap'] = overlap

                        if score > best_score:
                            best_score = score
                            best_match = kg_type
                            match_details = details

                    # 提高阈值：要求至少有 transaction_characteristics 或 usage_scenario 的匹配
                    dynamic_threshold = 25  # 提高基础阈值

                    # 额外检查：如果没有匹配到重要的具体字段，降低分数
                    if match_details.get('usage_scenario_match') != 'exact' and \
                       match_details.get('transaction_characteristics_match') != 'exact':
                        # 没有匹配到核心具体字段，增加惩罚
                        best_score = best_score * 0.6

                    if best_match and best_score >= dynamic_threshold:
                        contract_type_match = best_match['name']
                        matched_features = best_match.get('legal_features', {})
                        match_source = 'legal_feature_matching'
                        print(f"[知识图谱] 基于法律特征匹配到合同类型: {contract_type_match} (匹配分数: {best_score}/{dynamic_threshold})")
                        print(f"[知识图谱] 匹配详情: {match_details}")
                    else:
                        print(f"[知识图谱] 法律特征匹配分数不足 (最高分: {best_score}/{dynamic_threshold})，使用 LLM 推测结果")
                        if best_match:
                            print(f"[知识图谱] 最佳候选: {best_match['name']}, 匹配详情: {match_details}")
                        # 使用 LLM 推测的合同类型
                        if features_result.suggested_contract_type:
                            contract_type_match = features_result.suggested_contract_type
                            # 将提取的所有特征作为 matched_features
                            matched_features = {}
                            for field_name in legal_feature_fields.keys():
                                val = getattr(features_result, field_name, None)
                                if val and val not in ['未提及', '', '无']:
                                    matched_features[field_name] = val
                            match_source = 'llm_suggestion'

                except Exception as e:
                    print(f"[法律特征提取] 失败: {e}")
                    # 回退到原来的关键词匹配方法
                    text_sample = text[:800]
                    patterns = [
                        r'([^，。；\n]{2,8})(?:服务)?合同',
                        r'([^，。；\n]{2,8})(?:服务)?协议',
                        r'([^，。；\n]{2,8})(?:合作)?契约',
                    ]
                    for pattern in patterns:
                        matches = re.findall(pattern, text_sample)
                        for match_candidate in matches:
                            match_candidate = match_candidate.strip()
                            if 2 <= len(match_candidate) <= 10:
                                match = contract_knowledge_db_service.get_by_name(match_candidate)
                                if match:
                                    contract_type_match = match['name']
                                    matched_features = match.get('legal_features', {})
                                    match_source = 'text_keyword_exact_match'
                                    print(f"[知识图谱] 回退到关键词匹配: {contract_type_match}")
                                    break
                                else:
                                    search_results = contract_knowledge_db_service.search_by_keywords(match_candidate)
                                    if search_results:
                                        contract_type_match = search_results[0]['name']
                                        matched_features = search_results[0].get('legal_features', {})
                                        match_source = 'text_keyword_fuzzy_match'
                                        print(f"[知识图谱] 回退到模糊匹配: {contract_type_match}")
                                        break
                            if contract_type_match:
                                break

            self.logger.info(f"[元数据提取] 知识图谱匹配完成，耗时: {time.time() - kg_match_start:.2f}秒，匹配结果: {contract_type_match or '无'}")

        except Exception as e:
            self.logger.error(f"[元数据提取] 知识图谱匹配失败: {e}")

        # === 步骤3: LLM 提取元数据 ===
        self.logger.info(f"[元数据提取] 步骤3: LLM提取元数据...")
        llm_extract_start = time.time()
        parser = PydanticOutputParser(pydantic_object=ContractMetadataSchema)

        # 根据是否匹配到知识图谱，使用不同的 prompt
        if contract_type_match:
            # 将 matched_features 转换为字符串，避免模板变量冲突
            # 使用三重大括号转义 JSON 内容
            matched_features_str = json.dumps(matched_features, ensure_ascii=False, indent=2) if matched_features else '无'
            # 转义花括号，避免 LangChain 解析为模板变量
            matched_features_str = matched_features_str.replace('{', '{{').replace('}', '}}')

            system_prompt = f"""你是一名专业的合同信息提取助手。合同已通过知识图谱识别为：【{contract_type_match}】

请严格按照以下 JSON 格式从合同文本中提取基本信息。

知识图谱参考特征：
{matched_features_str}

要求：
1. 合同类型应优先使用知识图谱识别的：{contract_type_match}（除非文本中有更明确的不同表述）
2. **当事人提取规则**：
   - parties 字段应按照"角色：名称"的格式，用分号分隔多个当事人
   - 正确识别甲方/乙方/丙方，以及委托人/受托人/出卖人/买受人等角色
   - 如果是三方或多方合同，必须列出所有签约方
   - 示例格式："甲方：某某公司；乙方：某某个人；丙方：担保方"
3. 提取金额、签署日期等关键信息
4. 识别合同的核心条款和标的

{{{{format_instructions}}}}"""
        else:
            system_prompt = """你是一名专业的合同信息提取助手。请严格按照以下 JSON 格式从合同文本中提取基本信息。

提示：请仔细识别合同类型，参考以下常见类型：
- 劳动合同、劳务合同、承揽合同、技术开发合同、买卖合同、租赁合同等
- 如果无法确定具体类型，请标注为"其他合同"并说明原因

**当事人提取规则**：
- parties 字段应按照"角色：名称"的格式，用分号分隔多个当事人
- 正确识别甲方/乙方/丙方，以及委托人/受托人/出卖人/买受人等角色
- 如果是三方或多方合同，必须列出所有签约方
- 示例格式："甲方：某某公司；乙方：某某个人；丙方：担保方"

{format_instructions}"""

        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("user", "合同全文：\n{text}"),
        ])

        chain = prompt | self.llm | parser

        try:
            # ⭐ 优化：智能选择文本范围，而不是固定的前8000字
            # 策略：优先提取合同开头部分（通常包含关键信息），如果文本过长则智能截取
            prompt_text = self._prepare_metadata_extraction_text(text, structured_meta)

            self.logger.info(f"[元数据提取] 准备LLM提取，输入文本长度: {len(prompt_text)}字符")

            result: ContractMetadataSchema = chain.invoke({
                "text": prompt_text,
                "format_instructions": parser.get_format_instructions(),
            })

            # 如果知识图谱匹配到了，且 LLM 识别的类型不一致，使用知识图谱结果
            if contract_type_match and result.contract_type != contract_type_match:
                print(f"[知识图谱] 覆盖 LLM 识别结果: {result.contract_type} -> {contract_type_match}")
                result.contract_type = contract_type_match

            # 更新数据库（包含知识图谱匹配信息）
            metadata_dict = result.model_dump()

            # 添加知识图谱匹配信息到元数据字典
            if contract_type_match:
                metadata_dict['knowledge_graph_match'] = True
                metadata_dict['match_source'] = match_source or 'kg_keyword_matching'
                metadata_dict['legal_features'] = matched_features

            contract.metadata_info = metadata_dict
            self.db.commit()

            total_time = time.time() - start_time
            self.logger.info(f"[元数据提取] ✅ 合同 {contract_id} 元数据已保存!")
            self.logger.info(f"[元数据提取] 总耗时: {total_time:.2f}秒 (文件解析: {file_parse_start and time.time() - file_parse_start or 0:.2f}s | 知识图谱: {kg_match_start and time.time() - kg_match_start or 0:.2f}s | LLM提取: {llm_extract_start and time.time() - llm_extract_start or 0:.2f}s)")
            self.logger.info(f"[元数据提取] 提取结果: contract_type={metadata_dict.get('contract_type')}, parties={metadata_dict.get('parties')}")
            return result
        except Exception as e:
            self.logger.error(f"[元数据提取] ❌ 元数据提取失败: {e}")
            traceback.print_exc()
            return None

    def _extract_structured_from_docx(self, doc: DocxDocument) -> dict:
        """
        从 docx 文档中抽取结构化信息：表格、签署方(甲方/乙方)、标题、金额等。
        返回一个字典，键可能包含: contract_name, parties (list), amount, tables

        增强当事人识别：
        - 支持甲方/乙方/丙方/丁方等多方识别
        - 支持委托人/受托人/出卖人/买受人等角色识别
        - 从表格中提取当事人信息
        - 识别公司名称、统一社会信用代码
        """
        meta: dict = {}

        # 1) 标题：尝试从 document.core_properties.title 或首段推断
        try:
            title = doc.core_properties.title
            if title and title.strip():
                meta['contract_name'] = title.strip()
            else:
                # 首个较短段落或加粗段落作为标题候选
                for p in doc.paragraphs[:6]:
                    text = (p.text or '').strip()
                    if text and len(text) < 120 and text.split():
                        meta['contract_name'] = text
                        break
        except Exception:
            pass

        # 2) 表格内容聚合 - 优先从表格提取当事人
        tables = []
        parties_from_tables = []

        try:
            for table in doc.tables[:10]:  # 只检查前10个表格
                rows = []
                for r in table.rows:
                    cells = [c.text.strip() for c in r.cells]
                    rows.append(cells)

                    # 检查表格中是否包含当事人信息
                    for i, cell in enumerate(cells):
                        # 模式匹配：甲方/乙方等
                        import re
                        party_patterns = [
                            r'^(甲方|乙方|丙方|丁方|戊方)\s*[:：]\s*(.+)$',
                            r'^(委托人|受托人|出卖人|买受人|出租人|承租人|发包人|承包人|转让方|受让方)\s*[:：]\s*(.+)$',
                            r'^(Party\s*[ABC])\s*[:：]\s*(.+)$',
                        ]
                        for pattern in party_patterns:
                            m = re.match(pattern, cell)
                            if m:
                                role = m.group(1)
                                name = m.group(2).strip()
                                if name and len(name) > 1:
                                    parties_from_tables.append({'role': role, 'name': name})
                                    print(f"[当事人提取] 从表格识别: {role} = {name}")

                if rows:
                    tables.append(rows)
        except Exception as e:
            print(f"[当事人提取] 表格解析失败: {e}")

        if tables:
            meta['tables'] = tables

        # 3) 从段落中提取当事人（改进版）
        parties = []
        try:
            # 合并前50个段落用于当事人提取
            relevant_text = []
            for i, p in enumerate(doc.paragraphs[:50]):
                t = (p.text or '').strip()
                if t:
                    relevant_text.append(t)

            import re

            # 定义更全面的当事人角色模式
            role_patterns = [
                # 标准甲方乙方模式
                r'^(甲方|乙方|丙方|丁方|戊方)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
                # 委托代理模式
                r'^(委托人|受托人|被委托人)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
                # 买卖模式
                r'^(出卖人|买受人|卖方|买方)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
                # 租赁模式
                r'^(出租人|承租人|租赁方)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
                # 工程模式
                r'^(发包人|承包人|分包人|监理人)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
                # 转让模式
                r'^(转让方|受让方)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
                # 服务模式
                r'^(委托方|服务方|提供方|接受方)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
                # 投资模式
                r'^(投资方|被投资方|融资方)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
                # 担保模式
                r'^(保证人|债务人|债权人)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
                # 英文模式
                r'^(Party\s*[ABC]|Client|Service\s*Provider|Employer|Employee)\s*[:：]\s*(.+?)(?:\s|$|，|。|\n)',
            ]

            seen_parties = set()  # 避免重复

            for line in relevant_text:
                for pattern in role_patterns:
                    m = re.match(pattern, line)
                    if m:
                        role = m.group(1)
                        name = m.group(2).strip()

                        # 清理名称中的多余符号
                        name = re.sub(r'^[:：\s]+|[:：\s,。，]+$|（.*?）|\(.*?\)', '', name).strip()

                        # 过滤掉太短或明显无效的名称
                        if name and len(name) >= 2 and len(name) <= 100:
                            # 检查是否已经记录过
                            party_key = f"{role}:{name}"
                            if party_key not in seen_parties:
                                parties.append({'role': role, 'name': name})
                                seen_parties.add(party_key)
                                print(f"[当事人提取] 从段落识别: {role} = {name}")

                            # 如果是表格中也有的，跳过后续匹配
                            break

        except Exception as e:
            print(f"[当事人提取] 段落解析失败: {e}")

        # 合并表格和段落提取的当事人（表格优先）
        all_parties = parties_from_tables + parties

        if all_parties:
            # 格式化为可读字符串
            party_strings = []
            for p in all_parties:
                party_strings.append(f"{p['role']}：{p['name']}")
            meta['parties'] = party_strings
            # 同时保留结构化格式
            meta['parties_structured'] = all_parties

        # 4) 提取金额（保留原有逻辑）
        amount = None
        try:
            for p in doc.paragraphs:
                t = (p.text or '').strip()
                if not t:
                    continue

                # 匹配金额，如 10000元、￥10000、10,000.50元
                m2 = re.search(r'([¥￥$]?\s*[0-9,]{1,10}(?:\.\d{1,2})?\s*(元|USD|美元|人民币|CNY)?)', t)
                if m2 and not amount:
                    amount = m2.group(1).strip()
                    # 只提取合理范围的金额（100-1亿）
                    try:
                        num_str = re.sub(r'[^\d.]', '', amount)
                        num_val = float(num_str)
                        if 100 <= num_val <= 100000000:
                            break
                        else:
                            amount = None
                    except:
                        amount = None
        except Exception:
            pass

        if amount:
            meta['amount'] = amount

        return meta

    def _extract_text_from_file(self, file_path: str) -> str:
        """
        从文件路径中提取文本内容，自动识别文件类型

        Args:
            file_path: 文件路径（可能是 .pdf, .docx, .doc 等）

        Returns:
            提取的文本内容
        """
        if not file_path or not os.path.exists(file_path):
            return ""

        file_path_lower = file_path.lower()

        # 如果是 PDF 文件，使用 PdfService
        if file_path_lower.endswith('.pdf'):
            text, _ = PdfService.parse_pdf(file_path)
            return text

        # 如果是 DOCX 文件，使用 python-docx
        if file_path_lower.endswith('.docx'):
            try:
                doc = DocxDocument(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                return text
            except Exception as e:
                print(f"Error reading DOCX file: {e}")
                traceback.print_exc()
                return ""

        # 如果是 DOC 文件，需要先转换
        if file_path_lower.endswith('.doc'):
            if PYPANDOC_AVAILABLE:
                try:
                    converted = file_path + '.converted.docx'
                    pypandoc.convert_file(file_path, 'docx', outputfile=converted)
                    doc = DocxDocument(converted)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return text
                except Exception as e:
                    print(f"Error converting DOC file: {e}")
                    traceback.print_exc()
                    return ""

        # 尝试作为 PDF 解析（如果 pdf_converted_path 指向的是 PDF）
        try:
            text, _ = PdfService.parse_pdf(file_path)
            if text:
                return text
        except Exception as e:
            print(f"Failed to parse as PDF: {e}")

        return ""

    def _prepare_metadata_extraction_text(self, full_text: str, structured_meta: dict) -> str:
        """
        ⭐ 智能准备元数据提取的文本

        策略：
        1. 优先提取合同开头部分（通常包含标题、当事人、金额等关键信息）
        2. 如果有结构化信息，优先使用
        3. 智能截取，避免截断关键信息
        4. 控制在合理长度内（15000字符），确保 LLM 能处理

        Args:
            full_text: 完整的合同文本
            structured_meta: 从文档结构中提取的元数据（表格、标题等）

        Returns:
            准备好的文本，用于 LLM 提取元数据
        """
        import re

        # 如果有结构化信息，优先使用
        if structured_meta:
            meta_text = json.dumps(structured_meta, ensure_ascii=False, indent=2)
            self.logger.info(f"[元数据提取] 使用结构化信息，长度: {len(meta_text)}字符")
        else:
            meta_text = ""

        # 智能截取策略：
        # 1. 合同开头通常包含关键信息（标题、当事人、金额等）
        # 2. 但也要保留足够长度的文本（15000字符）
        # 3. 如果合同较短，使用全文
        MAX_TEXT_LENGTH = 15000  # 增加到15000字符

        if len(full_text) <= MAX_TEXT_LENGTH:
            # 短合同，使用全文
            text_portion = full_text
            self.logger.info(f"[元数据提取] 使用全文，长度: {len(text_portion)}字符")
        else:
            # 长合同，智能截取
            # 策略：取前 12000 字符 + 后 3000 字符（通常签名、日期在最后）
            text_portion = full_text[:12000] + "\n\n...（省略中间内容）...\n\n" + full_text[-3000:]
            self.logger.info(f"[元数据提取] 长合同智能截取，总长度: {len(full_text)}字符 -> {len(text_portion)}字符")

        # 组合：结构化信息 + 合同文本
        if meta_text:
            result = f"# 从文档结构提取的信息:\n{meta_text}\n\n# 合同文本内容:\n{text_portion}"
        else:
            result = text_portion

        self.logger.info(f"[元数据提取] 最终输入长度: {len(result)}字符")
        return result

    def _prepare_review_text(self, full_text: str) -> str:
        """
        ⭐ 智能准备深度审查的文本

        与元数据提取不同，深度审查需要覆盖更多内容以发现所有风险点

        策略：
        1. 短合同（<20000字符）：使用全文
        2. 长合同（≥20000字符）：智能分块
           - 前 15000 字符（包含开头的重要条款）
           - 后 5000 字符（包含结尾的争议解决、签署等）

        Args:
            full_text: 完整的合同文本

        Returns:
            准备好的文本，用于深度审查
        """
        MAX_REVIEW_LENGTH = 20000  # 增加到20000字符

        if len(full_text) <= MAX_REVIEW_LENGTH:
            # 短合同，使用全文
            self.logger.info(f"[深度审查] 使用全文，长度: {len(full_text)}字符")
            return full_text
        else:
            # 长合同，智能截取
            text_portion = full_text[:15000] + "\n\n...（省略中间内容）...\n\n" + full_text[-5000:]
            self.logger.info(f"[深度审查] 长合同智能截取，总长度: {len(full_text)}字符 -> {len(text_portion)}字符")
            return text_portion

    def run_deep_review(
        self,
        contract_id: int,
        stance: str = "甲方",
        updated_meta: dict | None = None,
        enable_custom_rules: bool = False,
        user_id: int = 1,
        transaction_structures: list | None = None  # ⭐ 新增: 交易结构列表
    ) -> bool:
        """
        阶段二：深度合同审查（核心逻辑）
        返回 True 表示审查成功启动

        Args:
            contract_id: 合同ID
            stance: 审查立场
            updated_meta: 更新的元数据
            enable_custom_rules: 是否启用用户自定义规则
            user_id: 用户ID
            transaction_structures: 用户选择的交易结构列表 (新增)
        """
        contract = self.db.query(ContractDoc).filter(ContractDoc.id == contract_id).first()
        if not contract:
            raise ValueError("合同不存在")

        # 更新用户确认后的元数据和立场
        if updated_meta:
            contract.metadata_info = updated_meta
        contract.stance = stance
        contract.status = ContractStatus.REVIEWING.value

        # ⭐ 新增: 保存交易结构
        if transaction_structures:
            contract.transaction_structures = transaction_structures
            self.logger.info(f"[ContractReviewService] 使用交易结构: {transaction_structures}")

        self.db.commit()

        # 提取全文 - 优先使用 original_file_path，其次使用 pdf_converted_path
        text = ""
        if contract.original_file_path and os.path.exists(contract.original_file_path):
            text = self._extract_text_from_file(contract.original_file_path)

        # 如果从 original_file_path 提取失败，尝试 pdf_converted_path
        if not text and contract.pdf_converted_path and os.path.exists(contract.pdf_converted_path):
            text = self._extract_text_from_file(contract.pdf_converted_path)

        if not text:
            contract.status = ContractStatus.DRAFT.value
            self.db.commit()
            raise ValueError("无法读取合同文本")

        # ⭐ 使用 RuleAssembler 加载规则（支持 universal/stance/feature/industry 分类）
        self.logger.info(f"[ContractReviewService] 使用 RuleAssembler 加载规则")

        # 获取元数据（用户更新后的或数据库中的）
        metadata = contract.metadata_info or {}

        # 准备法律特征字典
        legal_features = {
            "contract_type": metadata.get("contract_type"),
            "transaction_nature": metadata.get("legal_features", {}).get("transaction_nature") if isinstance(metadata.get("legal_features"), dict) else None,
            "contract_object": metadata.get("legal_features", {}).get("contract_object") if isinstance(metadata.get("legal_features"), dict) else None,
        }

        # 使用 RuleAssembler 组装规则
        assembled_rules = rule_assembler.assemble_prompt_context(
            legal_features=legal_features,
            stance=stance,
            user_id=user_id if enable_custom_rules else None,
            transaction_structures=transaction_structures
        )

        rule_content = assembled_rules if assembled_rules else "（暂无审查规则）"
        self.logger.info(f"[ContractReviewService] 规则组装完成，规则长度: {len(rule_content)} 字符")

        parser = PydanticOutputParser(pydantic_object=ReviewOutput)

        system_prompt = """
# 企业级合同智能审查指引

你是一名资深的企业法务专家。请遵循以下 **"四层审查法"** 对用户提供的合同进行全方位深度审查。

---

## 四层审查法

{rules}

---

## 输出要求

1. **只输出发现的风险点**，不输出无风险的确认信息
2. **严格遵循 JSON Schema** 格式
3. **违反系统规则或自定义规则中明确禁止的内容**的，必须将 `action_type` 设为 `"Alert"`，`severity` 设为 `"Critical"` 或 `"High"`
4. **常规法律风险**的，将 `action_type` 设为 `"Revision"`，根据风险程度设置 `severity`
5. `issue_type` 必须使用**中文**描述，格式为：`[审查层级] 具体问题`，例如：
   - `宏观审查-主体资格瑕疵`
   - `中观审查-条款缺失`
   - `微观审查-金额表述错误`
   - `自定义规则-特殊要求不符`
6. `quote` 必须是合同原文中的**具体短语或句子**，便于定位
7. `explanation` 应简要说明**为什么这是问题**
8. `suggestion` 应提供**具体的修改建议或警告措辞**
9. **`legal_basis` 必须列明审查依据**，包括：
   - 相关法律法规名称及条款（如《民法典》第xxx条、《公司法》第xxx条）
   - 司法解释或裁判观点
   - 行业规范或标准条款

请严格按照以下 JSON 格式输出：
{format_instructions}
        """

        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("user", """【合同基本信息】
{metadata}

【审查立场】{stance}

【合同全文】（前10000字）：
{text}

请基于以上四层审查法，识别并列出所有风险点。"""),
        ])

        chain = prompt | self.llm | parser

        try:
            # ⭐ 优化：使用智能文本准备，而不是固定的前10000字
            review_text = self._prepare_review_text(text)

            self.logger.info(f"[ContractReviewService] 开始LLM深度审查，输入文本长度: {len(review_text)}字符")

            result: ReviewOutput = chain.invoke({
                "stance": stance,
                "metadata": json.dumps(contract.metadata_info, ensure_ascii=False, indent=2),
                "rules": rule_content,
                "text": review_text,  # ⭐ 使用智能准备的文本（最多20000字符）
                "format_instructions": parser.get_format_instructions(),
            })

            # 清空旧审查记录
            self.db.query(ContractReviewItem).filter(ContractReviewItem.contract_id == contract.id).delete()

            # 写入新审查项
            for issue in result.issues:
                item = ContractReviewItem(
                    contract_id=contract.id,
                    issue_type=issue.issue_type,
                    quote=issue.quote,
                    explanation=issue.explanation,
                    suggestion=issue.suggestion,
                    legal_basis=getattr(issue, 'legal_basis', ''),
                    severity=issue.severity,
                    action_type=issue.action_type,
                    item_status="Pending",  # 待人工确认
                )
                self.db.add(item)

            contract.status = ContractStatus.WAITING_HUMAN.value
            self.db.commit()

            return True

        except Exception as e:
            print(f"深度审查失败: {e}")
            contract.status = ContractStatus.DRAFT.value
            self.db.commit()
            return False