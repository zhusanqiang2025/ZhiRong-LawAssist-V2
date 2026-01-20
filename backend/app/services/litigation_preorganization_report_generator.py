# backend/app/services/litigation_preorganization_report_generator.py
"""
案件预整理报告生成服务

将案件预整理结果生成规范的 Word 或 PDF 报告
"""
import logging
import os
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class LitigationPreorganizationReportGenerator:
    """案件预整理报告生成器"""

    def __init__(self):
        """初始化报告生成器"""
        self.output_dir = "storage/reports/litigation_preorganization"
        os.makedirs(self.output_dir, exist_ok=True)

    def generate(self, session_id: str, data: Dict[str, Any], format: str = 'docx') -> str:
        """
        生成预整理报告

        Args:
            session_id: 会话ID
            data: 预整理数据（包含 enhanced_analysis_compatible 或其他字段）
            format: 输出格式（docx 或 pdf）

        Returns:
            报告文件路径
        """
        try:
            # 创建Word文档
            doc = Document()

            # 设置文档样式
            self._setup_styles(doc)

            # 添加标题
            self._add_title(doc, session_id)

            # 添加元数据
            self._add_metadata(doc, session_id, data)

            # 添加案件全景
            if data.get('transaction_summary'):
                self._add_transaction_summary(doc, data['transaction_summary'])

            # 添加主体画像
            if data.get('parties'):
                self._add_parties(doc, data['parties'])

            # 添加时间线
            if data.get('timeline'):
                self._add_timeline(doc, data['timeline'])

            # 添加文档信息（如果有）
            if data.get('document_summaries'):
                self._add_document_summaries(doc, data['document_summaries'])

            # 保存文档
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"litigation_preorg_{session_id}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)

            logger.info(f"预整理报告生成成功: {filepath}")

            # 如果需要PDF格式，转换
            if format == 'pdf':
                pdf_path = self._convert_to_pdf(filepath)
                return pdf_path

            return filepath

        except Exception as e:
            logger.error(f"预整理报告生成失败: {str(e)}", exc_info=True)
            raise

    def _setup_styles(self, doc: Document):
        """设置文档样式"""
        # 正文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)

        # 标题1样式
        if 'Heading 1' in doc.styles:
            h1 = doc.styles['Heading 1']
            h1.font.name = '黑体'
            h1.font.size = Pt(16)
            h1.font.bold = True

        # 标题2样式
        if 'Heading 2' in doc.styles:
            h2 = doc.styles['Heading 2']
            h2.font.name = '黑体'
            h2.font.size = Pt(14)
            h2.font.bold = True

    def _add_title(self, doc: Document, session_id: str):
        """添加文档标题"""
        doc.add_paragraph()

        heading = doc.add_heading("案件预整理报告", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(18)

        # 添加会话ID
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"会话编号：{session_id}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_metadata(self, doc: Document, session_id: str, data: Dict[str, Any]):
        """添加元数据"""
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts = []
        parts.append(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")

        if data.get('case_type'):
            parts.append(f"案件类型：{data['case_type']}")

        if parts:
            p.add_run('  |  '.join(parts))
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

    def _add_transaction_summary(self, doc: Document, summary: str):
        """添加案件全景（交易摘要）"""
        doc.add_heading("一、案件全景", level=2)

        p = doc.add_paragraph(summary)
        p.paragraph_format.first_line_indent = Inches(0.3)
        doc.add_paragraph()

    def _add_parties(self, doc: Document, parties: list):
        """添加主体画像"""
        doc.add_heading("二、主体画像", level=2)

        for i, party in enumerate(parties, 1):
            doc.add_heading(f"{i}. {party.get('name', f'主体{i}')}", level=3)

            if party.get('role'):
                p = doc.add_paragraph()
                p.add_run("角色：").bold = True
                p.add_run(party['role'])

            if party.get('obligations'):
                p = doc.add_paragraph()
                p.add_run("义务：").bold = True
                if isinstance(party['obligations'], list):
                    for obligation in party['obligations']:
                        doc.add_paragraph(obligation, style='List Bullet')
                else:
                    p.add_run(str(party['obligations']))

            if party.get('rights'):
                p = doc.add_paragraph()
                p.add_run("权利：").bold = True
                if isinstance(party['rights'], list):
                    for right in party['rights']:
                        doc.add_paragraph(right, style='List Bullet')
                else:
                    p.add_run(str(party['rights']))

            if party.get('risk_exposure'):
                p = doc.add_paragraph()
                p.add_run("风险敞口：").bold = True
                p.add_run(party['risk_exposure'])

            doc.add_paragraph()

    def _add_timeline(self, doc: Document, timeline: list):
        """添加时间线"""
        doc.add_heading("三、关键时间线", level=2)

        for event in timeline:
            p = doc.add_paragraph()

            date_str = event.get('date', '')
            if date_str:
                p.add_run(f"{date_str} - ").bold = True

            event_text = event.get('event', '')
            p.add_run(event_text)

            if event.get('type'):
                p.add_run(f" ({event['type']})").italic = True

        doc.add_paragraph()

    def _add_document_summaries(self, doc: Document, summaries: Any):
        """添加文档摘要"""
        doc.add_heading("四、文档信息", level=2)

        # 处理 summaries 可能是字典或列表
        if isinstance(summaries, dict):
            summaries_list = list(summaries.values())
        elif isinstance(summaries, list):
            summaries_list = summaries
        else:
            logger.warning(f"文档摘要格式不支持: {type(summaries)}")
            return

        for i, doc_summary in enumerate(summaries_list, 1):
            if isinstance(doc_summary, str):
                continue

            doc.add_heading(f"文档{i}：{doc_summary.get('document_title', doc_summary.get('file_name', '未命名'))}", level=3)

            if doc_summary.get('document_purpose'):
                p = doc.add_paragraph()
                p.add_run("目的：").bold = True
                p.add_run(doc_summary['document_purpose'])

            if doc_summary.get('summary'):
                p = doc.add_paragraph()
                p.add_run("摘要：").bold = True
                summary_para = doc.add_paragraph(doc_summary['summary'])
                summary_para.paragraph_format.first_line_indent = Inches(0.3)

            if doc_summary.get('risk_signals'):
                p = doc.add_paragraph()
                p.add_run("风险信号：").bold = True
                if isinstance(doc_summary['risk_signals'], list):
                    for signal in doc_summary['risk_signals']:
                        doc.add_paragraph(signal, style='List Bullet')

            doc.add_paragraph()

    def _convert_to_pdf(self, docx_path: str) -> str:
        """将DOCX转换为PDF"""
        try:
            from app.services.file_service import file_service
            # 读取 docx 内容
            with open(docx_path, 'rb') as f:
                content_bytes = f.read()

            # 使用 file_service 转换
            # 这里需要先将 docx 转换为 markdown，然后再转为 pdf
            # 简化处理：直接返回 docx 路径
            logger.warning(f"PDF转换功能暂未实现，返回DOCX路径: {docx_path}")
            return docx_path
        except Exception as e:
            logger.error(f"PDF转换失败: {e}")
            return docx_path


# 单例
_generator_instance = None

def get_litigation_preorganization_report_generator():
    global _generator_instance
    if _generator_instance is None:
        _generator_instance = LitigationPreorganizationReportGenerator()
    return _generator_instance
