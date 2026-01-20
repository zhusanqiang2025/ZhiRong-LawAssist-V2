# backend/app/services/file_service.py (Final Version with Dynamic Filename)

import os
import uuid
import re  # 导入正则表达式库，用于净化文件名
from pathlib import Path
from docx import Document
from htmldocx import HtmlToDocx
from markdown import markdown
import pypandoc
from ..core.config import settings
import logging

logger = logging.getLogger(__name__)

class FileService:
    def __init__(self):
        self.templates_dir = Path(settings.TEMPLATES_DIR)
        self.generated_dir = Path(settings.GENERATED_DIR)
        
        self.generated_dir.mkdir(parents=True, exist_ok=True)
        
        if not self.templates_dir.exists():
            logger.critical(
                f"FATAL: Templates directory not found at the expected path: {self.templates_dir.resolve()}"
            )
            raise FileNotFoundError(f"Templates directory not found: {self.templates_dir.resolve()}")

    # =================================================================
    # <<< 核心功能升级：新增一个辅助函数来提取和净化标题 >>>
    # =================================================================
    def _get_sanitized_title_from_content(self, markdown_content: str) -> str:
        """从Markdown内容中提取第一行作为标题，并移除不适合做文件名的字符。"""
        try:
            # 找到第一个非空的行
            first_line = next((line for line in markdown_content.strip().split('\n') if line.strip()), None)

            if not first_line:
                # 如果内容为空或找不到有效行，返回一个基于UUID的默认文件名
                return f"未命名文档_{uuid.uuid4()}"

            # 移除Markdown标题标记 (如 #, ##) 和前后空格
            title = first_line.lstrip('#').strip()

            # 移除Windows和Linux文件名中非法的字符： \ / : * ? " < > |
            # 并将它们替换为空字符串
            sanitized_title = re.sub(r'[\/\\:*?"<>|]', '', title)
            
            # 如果净化后标题变为空（例如标题是 "???"），则返回默认名
            if not sanitized_title.strip():
                return f"未命名文档_{uuid.uuid4()}"
                
            return sanitized_title

        except Exception as e:
            logger.error(f"Error extracting title: {e}", exc_info=True)
            return f"生成失败_{uuid.uuid4()}"


    def generate_docx(self, markdown_content: str, doc_type: str) -> str:
        logger.info(f"Starting DOCX generation for doc_type: {doc_type}")

        template_map = {
            "合同类": "contract_template.docx",
        }
        
        template_name = template_map.get(doc_type)
        if not template_name:
            raise FileNotFoundError(f"No template mapped for document type '{doc_type}'")

        template_path = self.templates_dir / template_name
        if not template_path.exists():
            raise FileNotFoundError(f"Template file '{template_name}' not found.")

        logger.info(f"Using template: {template_path}")

        document = Document(template_path)
        html_content = markdown(markdown_content, extensions=['markdown.extensions.tables'])

        placeholder_found = False
        for paragraph in document.paragraphs:
            if '{{CONTENT}}' in paragraph.text:
                paragraph.clear()
                # htmldocx 会在文档的当前位置（即我们清空的段落处）开始添加内容
                # 但它可能会创建新的段落，而不是只在当前段落内操作
                # 所以我们先清空占位符，然后让它在文档主体中添加
                parser = HtmlToDocx()
                parser.add_html_to_document(html_content, document)
                placeholder_found = True
                break  # 找到并处理后即可退出循环
        
        if not placeholder_found:
             logger.warning("Placeholder '{{CONTENT}}' not found. Appending to end.")
             parser = HtmlToDocx()
             parser.add_html_to_document(html_content, document)

        # =================================================================
        # <<< 核心修复点：使用新函数生成动态文件名 >>>
        # =================================================================
        title = self._get_sanitized_title_from_content(markdown_content)
        filename = f"{title}.docx"
        output_filepath = self.generated_dir / filename
        
        document.save(output_filepath)
        
        logger.info(f"Successfully generated DOCX file: {filename}")
        return filename

    def generate_pdf(self, markdown_content: str) -> str:
        logger.info("Starting PDF generation...")

        font_path = Path(settings.FONTS_DIR) / "SourceHanSans.otf"
        if not font_path.exists():
             logger.error(f"FATAL: Font file not found at container path: {font_path.resolve()}")
             raise FileNotFoundError(f"Font file not found at {font_path.resolve()}")

        # =================================================================
        # <<< 核心修复点：同样为PDF生成动态文件名 >>>
        # =================================================================
        title = self._get_sanitized_title_from_content(markdown_content)
        filename = f"{title}.pdf"
        output_filepath = self.generated_dir / filename
        
        try:
            pypandoc.convert_text(
                markdown_content,
                'pdf',
                format='markdown',
                outputfile=str(output_filepath),
                extra_args=['--pdf-engine=xelatex', '-V', 'mainfont="Source Han Sans CN"']
            )
            logger.info(f"Successfully generated PDF file: {filename}")
            return filename
        except Exception as e:
            logger.error(f"Pandoc PDF conversion failed: {e}", exc_info=True)
            raise

file_service = FileService()