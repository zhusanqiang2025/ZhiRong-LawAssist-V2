# backend/app/services/document_preprocessor.py
"""
合同文件预处理中心
负责将各种格式的合同文件转换为统一的 .docx 格式
可被其他模块复用

支持外部服务：
- MinerU PDF 解析服务
- OCR 文字识别服务（支持公网/本地图片、PDF）
"""
import os
import uuid
import shutil
import logging
import io
import requests
import re
from typing import Tuple, Optional, Dict, List
from enum import Enum
from pathlib import Path
from PIL import Image
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)


class DocumentFormat(str, Enum):
    """支持的文档格式"""
    DOC = "doc"          # legacy Word
    DOCX = "docx"        # modern Word
    PDF = "pdf"          # PDF
    TXT = "txt"          # 纯文本
    RTF = "rtf"          # Rich Text Format
    ODT = "odt"          # OpenDocument Text
    # 图片格式
    JPG = "jpg"          # JPEG 图片
    JPEG = "jpeg"        # JPEG 图片
    PNG = "png"          # PNG 图片
    BMP = "bmp"          # BMP 图片
    TIFF = "tiff"        # TIFF 图片
    GIF = "gif"          # GIF 图片
    UNKNOWN = "unknown"


class ConversionResult(str, Enum):
    """转换结果状态"""
    SUCCESS = "success"
    FAILED = "failed"
    ALREADY_CONVERTED = "already_converted"
    UNSUPPORTED_FORMAT = "unsupported_format"


class DocumentPreprocessor:
    """
    合同文件预处理中心

    功能：
    1. 格式检测：自动识别文件格式
    2. 格式转换：统一转换为 .docx 格式
    3. 文档后处理：清理页码、多余空格、不正常换行等
    4. 质量检查：验证转换后的文件完整性
    5. 元数据提取：提取文档基本信息（页数、字数等）
    """

    # 后处理选项
    ENABLE_POSTPROCESSING = True  # 是否启用后处理
    ENABLE_AI_POSTPROCESSING = True  # 是否启用 AI 辅助后处理
    REMOVE_PAGE_NUMBERS = True    # 删除页码
    CLEAN_SPACES = True           # 清理多余空格
    FIX_LINE_BREAKS = True        # 修复不正常换行
    REMOVE_EMPTY_PARAGRAPHS = True # 删除空段落

    # AI 后处理策略
    AI_CONFIDENCE_THRESHOLD = 0.7  # 规则置信度阈值
    AI_ONLY_AMBIGUOUS = True  # 仅对不确定的内容使用 AI
    AI_BATCH_SIZE = 5  # 批量处理时每次发送的段落数

    # 页码识别模式（常见的页码格式）
    PAGE_NUMBER_PATTERNS = [
        r'^\s*[\-—]*\s*\d+\s*[\-—]*\s*$',                    # - 1 - 或 — 2 —
        r'^\s*第\s*\d+\s*[页页张张]\s*$',                     # 第1页
        r'^\s*Page\s*\d+\s*$',                              # Page 1
        r'^\s*\d+\s*/\s*\d+\s*$',                          # 1/10
        r'^\s*[第]*\d+\s*页\s*$',                          # 第1页、1页
        r'^\s*\d+\s*$',                                     # 单独数字（可能是页码）
        r'^\s*[-–—]{2,}\s*\d+\s*[-–—]{2,}\s*$',           # -- 1 --
    ]

    # 支持的输入格式
    SUPPORTED_INPUT_FORMATS = {
        DocumentFormat.DOC,
        DocumentFormat.DOCX,
        DocumentFormat.PDF,
        DocumentFormat.TXT,
        DocumentFormat.RTF,
        DocumentFormat.ODT,
        # 图片格式
        DocumentFormat.JPG,
        DocumentFormat.JPEG,
        DocumentFormat.PNG,
        DocumentFormat.BMP,
        DocumentFormat.TIFF,
        DocumentFormat.GIF,
    }

    # 图片格式集合
    IMAGE_FORMATS = {
        DocumentFormat.JPG,
        DocumentFormat.JPEG,
        DocumentFormat.PNG,
        DocumentFormat.BMP,
        DocumentFormat.TIFF,
        DocumentFormat.GIF,
    }

    # 无需转换的格式
    NO_CONVERSION_NEEDED = {DocumentFormat.DOCX}

    def __init__(self, upload_dir: str = None, mineru_url: str = None, ocr_url: str = None):
        """
        初始化预处理中心

        Args:
            upload_dir: 文件上传目录，默认为 storage/uploads
            mineru_url: MinerU PDF 解析服务地址
            ocr_url: OCR 文字识别服务地址
        """
        self.upload_dir = upload_dir or "storage/uploads"
        os.makedirs(self.upload_dir, exist_ok=True)

        # 外部服务配置
        self.mineru_url = mineru_url
        self.ocr_url = ocr_url

        # 从环境变量读取配置（如果未传入）
        if not self.mineru_url:
            from app.core.config import settings
            self.mineru_url = settings.MINERU_API_URL if settings.MINERU_ENABLED else None
        if not self.ocr_url:
            from app.core.config import settings
            self.ocr_url = settings.OCR_API_URL if settings.OCR_ENABLED else None

        # 记录服务状态
        if self.mineru_url:
            logger.info(f"MinerU 服务已启用: {self.mineru_url}")
        else:
            logger.info("MinerU 服务未启用，将使用本地 PDF 解析")

        if self.ocr_url:
            logger.info(f"OCR 服务已启用: {self.ocr_url}")
        else:
            logger.info("OCR 服务未启用，将使用本地 OCR（如果可用）")

    def detect_format(self, file_path: str) -> DocumentFormat:
        """
        检测文件格式

        Args:
            file_path: 文件路径

        Returns:
            DocumentFormat 枚举值
        """
        if not os.path.exists(file_path):
            return DocumentFormat.UNKNOWN

        # 通过文件扩展名判断
        ext = Path(file_path).suffix.lower().lstrip('.')

        try:
            return DocumentFormat(ext)
        except ValueError:
            return DocumentFormat.UNKNOWN

    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """
        验证文件有效性

        Args:
            file_path: 文件路径

        Returns:
            (是否有效, 错误消息)
        """
        if not os.path.exists(file_path):
            return False, "文件不存在"

        if os.path.getsize(file_path) == 0:
            return False, "文件为空"

        # 检查文件大小限制（50MB）
        max_size = 50 * 1024 * 1024
        if os.path.getsize(file_path) > max_size:
            return False, f"文件过大（超过 {max_size // 1024 // 1024}MB）"

        file_format = self.detect_format(file_path)
        if file_format == DocumentFormat.UNKNOWN:
            return False, "不支持的文件格式"

        if file_format not in self.SUPPORTED_INPUT_FORMATS:
            return False, f"不支持的文件格式: {file_format.value}"

        return True, ""

    def needs_conversion(self, file_path: str) -> bool:
        """
        判断文件是否需要转换

        Args:
            file_path: 文件路径

        Returns:
            是否需要转换
        """
        file_format = self.detect_format(file_path)
        return file_format not in self.NO_CONVERSION_NEEDED

    def convert_to_docx(
        self,
        file_path: str,
        output_filename: str = None,
        force: bool = False
    ) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """
        将文件转换为 .docx 格式

        Args:
            file_path: 输入文件路径
            output_filename: 输出文件名（可选）
            force: 是否强制转换（即使已经是 .docx）

        Returns:
            (转换结果, 输出文件路径, 元数据)
        """
        # 验证文件
        is_valid, error_msg = self.validate_file(file_path)
        if not is_valid:
            return ConversionResult.FAILED, error_msg, None

        file_format = self.detect_format(file_path)

        # 已经是 .docx 且不强制转换
        if file_format == DocumentFormat.DOCX and not force:
            return ConversionResult.ALREADY_CONVERTED, file_path, self._extract_metadata(file_path)

        # 生成输出文件名
        if output_filename is None:
            base_name = Path(file_path).stem
            output_filename = f"{base_name}.converted.docx"

        output_path = os.path.join(self.upload_dir, output_filename)

        # 根据输入格式选择转换方法
        try:
            if file_format == DocumentFormat.DOC:
                return self._convert_doc_to_docx(file_path, output_path)
            elif file_format == DocumentFormat.PDF:
                return self._convert_pdf_to_docx(file_path, output_path)
            elif file_format == DocumentFormat.TXT:
                return self._convert_txt_to_docx(file_path, output_path)
            elif file_format == DocumentFormat.RTF:
                return self._convert_rtf_to_docx(file_path, output_path)
            elif file_format == DocumentFormat.ODT:
                return self._convert_odt_to_docx(file_path, output_path)
            elif file_format in self.IMAGE_FORMATS:
                return self._convert_image_to_docx(file_path, output_path)
            else:
                return ConversionResult.UNSUPPORTED_FORMAT, "", None

        except Exception as e:
            logger.error(f"文件转换失败: {file_path} -> {output_path}, 错误: {str(e)}")
            return ConversionResult.FAILED, str(e), None

    def _convert_doc_to_docx(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """转换 .doc 到 .docx"""
        from app.services.converter import convert_doc_to_docx

        filename = os.path.basename(input_path)
        success, docx_filename, msg = convert_doc_to_docx(filename)

        if success and docx_filename:
            converted_path = os.path.join(self.upload_dir, docx_filename)
            # 后处理转换后的文档
            if self.ENABLE_POSTPROCESSING:
                converted_path = self._postprocess_docx(converted_path)
            metadata = self._extract_metadata(converted_path)
            return ConversionResult.SUCCESS, converted_path, metadata
        else:
            return ConversionResult.FAILED, msg, None

    def _postprocess_docx(self, docx_path: str) -> str:
        """
        对转换后的 docx 文档进行后处理

        功能：
        1. 删除页码（单独一行的数字或页码格式）
        2. 清理多余空格
        3. 修复不正常换行
        4. 删除空段落
        5. [AI增强] 使用 AI 识别不确定的页码和段落边界

        Args:
            docx_path: docx 文件路径

        Returns:
            处理后的文件路径（同输入路径）
        """
        try:
            from docx import Document
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            logger.info(f"开始后处理文档: {docx_path}")
            doc = Document(docx_path)

            # 统计处理前的段落数
            original_para_count = len(doc.paragraphs)
            removed_count = 0
            modified_count = 0

            # 需要删除的段落索引（倒序收集以便安全删除）
            paragraphs_to_remove = []

            # 需要使用 AI 判断的段落（索引 + 原因）
            ai_pending_paragraphs = []  # [(index, reason), ...]

            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()

                # 1. 删除空段落
                if self.REMOVE_EMPTY_PARAGRAPHS and not text:
                    # 检查段落是否真的为空（没有图片、表格等内容）
                    if not self._has_content(para):
                        paragraphs_to_remove.append(i)
                        removed_count += 1
                        continue

                # 2. 删除页码（规则判断 + AI 增强）
                if self.REMOVE_PAGE_NUMBERS:
                    rule_result = self._is_page_number_with_confidence(text, i, original_para_count)

                    if rule_result["is_page_number"]:
                        # 规则明确判断为页码，直接删除
                        if rule_result["confidence"] >= self.AI_CONFIDENCE_THRESHOLD:
                            paragraphs_to_remove.append(i)
                            removed_count += 1
                            continue
                        # 置信度低，加入 AI 待判断列表
                        elif self.ENABLE_AI_POSTPROCESSING:
                            ai_pending_paragraphs.append((i, "potential_page_number", text))

                # 3. 清理多余空格
                if self.CLEAN_SPACES and text:
                    cleaned_text = self._clean_spaces(text)
                    if cleaned_text != text:
                        # 更新段落文本
                        self._set_paragraph_text(para, cleaned_text)
                        modified_count += 1

                # 4. 修复不正常换行（合并连续短段落）
                if self.FIX_LINE_BREAKS and text and len(text) < 50:
                    # 检查是否应该与下一段合并
                    if i + 1 < len(doc.paragraphs):
                        next_text = doc.paragraphs[i + 1].text.strip()
                        # 如果下一段也不是空的，且不是明显的段落边界
                        if next_text:
                            boundary_result = self._is_paragraph_boundary_with_confidence(text, next_text)
                            if not boundary_result["is_boundary"]:
                                # 规则明确判断应该合并
                                if boundary_result["confidence"] >= self.AI_CONFIDENCE_THRESHOLD:
                                    # 合并段落
                                    combined_text = text + " " + next_text
                                    self._set_paragraph_text(para, combined_text)
                                    # 标记下一段删除
                                    if i + 1 not in paragraphs_to_remove:
                                        paragraphs_to_remove.append(i + 1)
                                    modified_count += 1
                                # 置信度低，加入 AI 待判断列表
                                elif self.ENABLE_AI_POSTPROCESSING:
                                    ai_pending_paragraphs.append((i, "should_merge", text, next_text))

            # 5. AI 辅助判断（如果启用）
            if self.ENABLE_AI_POSTPROCESSING and ai_pending_paragraphs:
                logger.info(f"使用 AI 分析 {len(ai_pending_paragraphs)} 个不确定的段落")
                ai_results = self._ai_analyze_paragraphs(doc, ai_pending_paragraphs)

                # 处理 AI 判断结果
                for idx, should_remove in ai_results.items():
                    if should_remove:
                        if idx not in paragraphs_to_remove:
                            paragraphs_to_remove.append(idx)
                            removed_count += 1

            # 倒序删除段落（避免索引问题）
            for idx in sorted(paragraphs_to_remove, reverse=True):
                try:
                    # 获取段落的父元素
                    para_element = doc.paragraphs[idx]._element
                    para_element.getparent().remove(para_element)
                except Exception as e:
                    logger.warning(f"删除段落 {idx} 失败: {str(e)}")

            # 保存处理后的文档
            doc.save(docx_path)

            logger.info(f"文档后处理完成: {docx_path}, "
                       f"删除 {removed_count} 个段落, 修改 {modified_count} 个段落")

            return docx_path

        except Exception as e:
            logger.error(f"文档后处理失败: {str(e)}")
            return docx_path

    def _is_page_number(self, text: str, index: int, total_count: int) -> bool:
        """
        判断文本是否为页码（简化版本，保持兼容性）

        Args:
            text: 段落文本
            index: 段落索引
            total_count: 总段落数

        Returns:
            是否为页码
        """
        result = self._is_page_number_with_confidence(text, index, total_count)
        return result["is_page_number"]

    def _is_page_number_with_confidence(self, text: str, index: int, total_count: int) -> Dict:
        """
        判断文本是否为页码（带置信度）

        Args:
            text: 段落文本
            index: 段落索引
            total_count: 总段落数

        Returns:
            {"is_page_number": bool, "confidence": float, "reason": str}
        """
        if not text:
            return {"is_page_number": False, "confidence": 1.0, "reason": "empty"}

        # 高置信度模式：明确的页码格式
        high_confidence_patterns = [
            (r'^\s*第\s*\d+\s*[页页张张]\s*$', "standard_page_cn"),  # 第1页
            (r'^\s*Page\s*\d+\s*$', "standard_page_en"),  # Page 1
            (r'^\s*\d+\s*/\s*\d+\s*$', "page_fraction"),  # 1/10
            (r'^\s*[-–—]{2,}\s*\d+\s*[-–—]{2,}\s*$', "dashed_number"),  # -- 1 --
        ]

        for pattern, reason in high_confidence_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return {"is_page_number": True, "confidence": 0.95, "reason": reason}

        # 中等置信度模式：单独的数字
        if re.match(r'^\s*[\-—]*\s*\d+\s*[\-—]*\s*$', text):
            # 单独的数字，需要根据位置判断
            if text.isdigit() and len(text) <= 3:
                # 如果在文档中间位置，更可能是页码
                if 5 < index < total_count - 5:
                    return {"is_page_number": True, "confidence": 0.6, "reason": "lone_number_mid"}
                # 在开头或结尾，可能是标题或序号
                else:
                    return {"is_page_number": False, "confidence": 0.5, "reason": "lone_number_edge"}

        return {"is_page_number": False, "confidence": 1.0, "reason": "no_match"}

    def _has_content(self, paragraph) -> bool:
        """
        检查段落是否包含实际内容（图片、形状等）

        Args:
            paragraph: python-docx 段落对象

        Returns:
            是否有内容
        """
        # 检查段落中的所有运行
        for run in paragraph.runs:
            # 检查是否有图片
            if 'graphic' in run._element.xml:
                return True
            # 检查是否有形状
            if 'shape' in run._element.xml:
                return True

        return False

    def _clean_spaces(self, text: str) -> str:
        """
        清理文本中的多余空格

        Args:
            text: 输入文本

        Returns:
            清理后的文本
        """
        # 替换多个连续空格为单个空格
        text = re.sub(r' +', ' ', text)
        # 替换多个连续制表符为单个空格
        text = re.sub(r'\t+', ' ', text)
        # 清理行首行尾空格
        text = text.strip()
        # 处理中文和英文/数字之间的空格（保留一个）
        # 中文字符后跟空格再跟英文/数字
        text = re.sub(r'([\u4e00-\u9fff]) +([a-zA-Z0-9])', r'\1\2', text)
        # 英文/数字后跟空格再跟中文字符
        text = re.sub(r'([a-zA-Z0-9]) +([\u4e00-\u9fff])', r'\1\2', text)

        return text

    def _is_paragraph_boundary(self, current_text: str, next_text: str) -> bool:
        """
        判断两个段落之间是否应该保持分段（简化版本，保持兼容性）

        Args:
            current_text: 当前段落文本
            next_text: 下一段落文本

        Returns:
            是否应该保持分段
        """
        result = self._is_paragraph_boundary_with_confidence(current_text, next_text)
        return result["is_boundary"]

    def _is_paragraph_boundary_with_confidence(self, current_text: str, next_text: str) -> Dict:
        """
        判断两个段落之间是否应该保持分段（带置信度）

        Args:
            current_text: 当前段落文本
            next_text: 下一段落文本

        Returns:
            {"is_boundary": bool, "confidence": float, "reason": str}
        """
        # 高置信度：当前段以句号等结尾
        if current_text.rstrip().endswith(('.', '。', '!', '！', '?', '？')):
            return {"is_boundary": True, "confidence": 0.9, "reason": "sentence_ending"}

        # 高置信度：下一段以常见段落起始词开头
        starters = ['首先', '其次', '最后', '另外', '此外', '总之', '因此',
                   '第', '一、', '二、', '三、', '（', '1.', '2.', '3.']

        for starter in starters:
            if next_text.startswith(starter):
                return {"is_boundary": True, "confidence": 0.95, "reason": f"starter_{starter}"}

        # 中等置信度：英文首字母大写
        if next_text and next_text[0].isupper():
            if current_text and current_text[-1].islower():
                return {"is_boundary": True, "confidence": 0.7, "reason": "english_case_change"}

        # 低置信度：没有明确边界标记
        return {"is_boundary": False, "confidence": 0.4, "reason": "no_clear_boundary"}

    def _set_paragraph_text(self, paragraph, new_text: str):
        """
        设置段落文本（保留原有格式）

        Args:
            paragraph: python-docx 段落对象
            new_text: 新文本
        """
        # 清除现有文本
        for run in paragraph.runs:
            run.text = ""

        # 如果没有运行了，添加一个新的
        if not paragraph.runs:
            paragraph.add_run(new_text)
        else:
            # 在第一个运行中设置新文本
            paragraph.runs[0].text = new_text

    def _ai_analyze_paragraphs(self, doc, pending_paragraphs: List[Tuple]) -> Dict[int, bool]:
        """
        使用 AI 分析不确定的段落

        Args:
            doc: python-docx 文档对象
            pending_paragraphs: 待分析的段落列表 [(index, reason, ...), ...]

        Returns:
            {paragraph_index: should_remove_or_merge} 字典
        """
        from app.services.ai_document_helper import get_ai_helper

        ai_helper = get_ai_helper()

        if not ai_helper.is_available():
            logger.warning("AI 服务不可用，跳过 AI 分析")
            return {}

        results = {}
        total_paras = len(doc.paragraphs)

        # 按批次处理
        batch_size = self.AI_BATCH_SIZE

        for i in range(0, len(pending_paragraphs), batch_size):
            batch = pending_paragraphs[i:i + batch_size]

            # 准备批量分类的输入
            paragraphs_text = []
            contexts = []

            for item in batch:
                idx = item[0]
                reason = item[1]

                if reason == "potential_page_number":
                    # 页码判断
                    text = item[2] if len(item) > 2 else ""
                    paragraphs_text.append(text)
                    contexts.append({
                        "index": idx,
                        "total": total_paras,
                        "task": "page_number_detection"
                    })

                elif reason == "should_merge":
                    # 段落合并判断
                    current_text = item[2] if len(item) > 2 else ""
                    next_text = item[3] if len(item) > 3 else ""
                    # 组合文本用于判断
                    combined = f"{current_text}\n{next_text}"
                    paragraphs_text.append(combined)
                    contexts.append({
                        "index": idx,
                        "total": total_paras,
                        "task": "paragraph_merge",
                        "current_text": current_text,
                        "next_text": next_text
                    })

            # 调用 AI 批量分类
            try:
                ai_results = ai_helper.batch_classify_paragraphs(paragraphs_text, contexts)

                # 处理 AI 返回结果
                for j, (item, ai_result) in enumerate(zip(batch, ai_results)):
                    idx = item[0]
                    reason = item[1]

                    if "error" in ai_result:
                        logger.warning(f"段落 {idx} 的 AI 分析失败: {ai_result['error']}")
                        continue

                    if reason == "potential_page_number":
                        # AI 判断是否为页码
                        is_page_number = ai_result.get("is_page_number", False)
                        confidence = ai_result.get("confidence", 0.5)

                        # 高置信度时，根据判断结果决定是否删除
                        if confidence >= self.AI_CONFIDENCE_THRESHOLD:
                            results[idx] = is_page_number  # 如果是页码则删除

                    elif reason == "should_merge":
                        # AI 判断是否应该合并（如果应该合并，则下一段需要删除）
                        should_merge = ai_result.get("should_merge_next", False)
                        confidence = ai_result.get("confidence", 0.5)

                        if confidence >= self.AI_CONFIDENCE_THRESHOLD and should_merge:
                            # 应该合并，标记下一段删除
                            idx_next = item[0] + 1
                            results[idx_next] = True

            except Exception as e:
                logger.error(f"AI 批量分析失败: {str(e)}")
                # 继续处理下一批

        return results

    def _convert_txt_to_docx(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """转换纯文本到 .docx"""
        try:
            from docx import Document

            # 检测编码
            encoding = self._detect_encoding(input_path)

            # 读取文本内容
            with open(input_path, 'r', encoding=encoding) as f:
                content = f.read()

            # 创建 Word 文档
            doc = Document()
            doc.add_paragraph(content)

            # 保存
            doc.save(output_path)

            # 后处理
            if self.ENABLE_POSTPROCESSING:
                self._postprocess_docx(output_path)

            metadata = self._extract_metadata(output_path)
            return ConversionResult.SUCCESS, output_path, metadata

        except Exception as e:
            return ConversionResult.FAILED, f"文本转换失败: {str(e)}", None

    def _convert_rtf_to_docx(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """转换 RTF 到 .docx"""
        # RTF 可以通过 LibreOffice 转换
        return self._convert_with_libreoffice(input_path, output_path)

    def _convert_odt_to_docx(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """转换 ODT 到 .docx"""
        # ODT 可以通过 LibreOffice 转换
        return self._convert_with_libreoffice(input_path, output_path)

    def _convert_with_libreoffice(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """使用 LibreOffice 进行通用转换"""
        import subprocess
        import time

        try:
            # 获取输入文件的绝对路径
            input_abs_path = os.path.abspath(input_path)
            output_abs_dir = os.path.abspath(self.upload_dir)

            cmd = [
                "soffice",
                "--headless",
                "--convert-to", "docx:MS Word 2007 XML",
                "--outdir", output_abs_dir,
                input_abs_path
            ]

            logger.info(f"LibreOffice 转换命令: {' '.join(cmd)}")

            result = subprocess.run(
                cmd,
                check=False,  # 改为 False，以便我们检查返回码
                timeout=60,
                capture_output=True,
                text=True
            )

            logger.info(f"LibreOffice 返回码: {result.returncode}")
            logger.info(f"LibreOffice stdout: {result.stdout}")
            if result.stderr:
                logger.warning(f"LibreOffice stderr: {result.stderr}")

            # 等待文件系统同步
            time.sleep(1.0)

            # 获取期望的输出文件名
            base_name = Path(input_abs_path).stem
            expected_filename = f"{base_name}.docx"
            expected_path = os.path.join(output_abs_dir, expected_filename)

            # 1. 首先检查期望的输出路径
            if os.path.exists(expected_path):
                logger.info(f"找到期望的输出文件: {expected_path}")
                # 后处理
                if self.ENABLE_POSTPROCESSING:
                    self._postprocess_docx(expected_path)
                metadata = self._extract_metadata(expected_path)
                return ConversionResult.SUCCESS, expected_path, metadata

            # 2. 列出目录中的所有文件用于调试
            all_files = os.listdir(output_abs_dir)
            logger.info(f"目录中的所有文件: {all_files}")

            # 3. 查找可能的输出文件
            all_docx_files = [f for f in all_files if f.endswith('.docx')]
            logger.info(f"目录中的所有 .docx 文件: {all_docx_files}")

            # 尝试多种匹配模式
            possible_files = [
                f for f in all_docx_files
                if os.path.splitext(f)[0] == base_name
            ]

            # 如果没有找到，尝试使用时间戳找到最新的文件
            if not possible_files and all_docx_files:
                # 查找最近创建的 .docx 文件（在最近 10 秒内）
                import time as time_module
                current_time = time_module.time()
                for f in all_docx_files:
                    file_path = os.path.join(output_abs_dir, f)
                    file_mtime = os.path.getmtime(file_path)
                    if current_time - file_mtime < 10:  # 10秒内创建的文件
                        possible_files.append(f)
                        logger.info(f"找到最近创建的文件: {f}")

            if possible_files:
                # 选择最新的文件（按修改时间）
                possible_files_with_time = [
                    (f, os.path.getmtime(os.path.join(output_abs_dir, f)))
                    for f in possible_files
                ]
                possible_files_with_time.sort(key=lambda x: x[1], reverse=True)
                newest_file = possible_files_with_time[0][0]

                converted_path = os.path.join(output_abs_dir, newest_file)
                logger.info(f"找到转换后的文件: {newest_file}")

                # 后处理
                if self.ENABLE_POSTPROCESSING:
                    self._postprocess_docx(converted_path)

                metadata = self._extract_metadata(converted_path)
                return ConversionResult.SUCCESS, converted_path, metadata

            error_msg = f"转换失败：未找到输出文件 (期望: {expected_filename})"
            if "export filter" in result.stderr:
                error_msg = "LibreOffice 导出过滤器配置错误，请联系管理员"

            logger.error(f"{error_msg}。目录中的文件: {all_files}")
            return ConversionResult.FAILED, error_msg, None

        except subprocess.TimeoutExpired:
            return ConversionResult.FAILED, "转换超时（超过60秒）", None
        except Exception as e:
            logger.error(f"转换异常: {str(e)}")
            return ConversionResult.FAILED, f"转换失败: {str(e)}", None

    def _convert_image_to_docx(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """
        转换图片到 .docx（使用OCR识别文字）

        支持的格式：jpg, jpeg, png, bmp, tiff, gif

        注意：OCR 功能需要安装系统依赖（RapidOCR 或 pytesseract）
        如果未安装，将返回提示信息
        """
        try:
            # 使用 OCR 提取图片中的文字
            ocr_result = self._ocr_image(input_path)

            if ocr_result["success"] and ocr_result["text"] and len(ocr_result["text"].strip()) >= 10:
                text = ocr_result["text"]
            else:
                # OCR 不可用或未识别到文字
                return ConversionResult.FAILED, ocr_result.get("message", "无法识别图片中的文字"), None

            # 创建 Word 文档
            from docx import Document
            from docx.shared import Inches, Pt

            doc = Document()

            # 添加标题
            title = doc.add_paragraph(os.path.basename(input_path))
            title.runs[0].font.size = Pt(14)
            title.runs[0].font.bold = True

            # 添加识别的文字
            doc.add_paragraph("【OCR 识别结果】")
            doc.add_paragraph(text)

            # 尝试添加原始图片
            try:
                # 检查图片尺寸
                with Image.open(input_path) as img:
                    width, height = img.size
                    # 计算合适的显示尺寸（最大6英寸宽）
                    max_width = 6.0
                    if width > 0:
                        scale = min(max_width / width, 1.0)
                        display_width = width * scale
                    else:
                        display_width = max_width

                doc.add_paragraph("\n【原始图片】")
                doc.add_picture(input_path, width=Inches(display_width))
            except Exception as e:
                logger.warning(f"添加原始图片到文档失败: {str(e)}")

            # 保存文档
            doc.save(output_path)

            # 后处理
            if self.ENABLE_POSTPROCESSING:
                self._postprocess_docx(output_path)

            metadata = self._extract_metadata(output_path)
            metadata["ocr_used"] = True
            metadata["original_image_size"] = os.path.getsize(input_path)

            return ConversionResult.SUCCESS, output_path, metadata

        except Exception as e:
            logger.error(f"图片 OCR 转换失败: {str(e)}")
            return ConversionResult.FAILED, f"图片处理失败: {str(e)}", None

    def _ocr_image(self, image_path: str) -> Dict:
        """
        使用 OCR 从图片中提取文字

        优先使用外部 OCR 服务，回退到本地 OCR

        Returns:
            Dict with keys:
                - success (bool): 是否成功识别
                - text (str): 识别的文字内容
                - message (str): 错误或提示信息
        """
        # 方法1：使用外部 OCR 服务（优先）
        if self.ocr_url:
            try:
                from app.core.config import settings

                with open(image_path, 'rb') as f:
                    files = {'file': f}
                    response = requests.post(
                        self.ocr_url,
                        files=files,
                        timeout=settings.OCR_API_TIMEOUT
                    )

                if response.status_code == 200:
                    result = response.json()
                    # 假设 OCR 服务返回格式: {"text": "...", "status": "success"}
                    if result.get("status") == "success" and result.get("text"):
                        return {
                            "success": True,
                            "text": result["text"],
                            "message": "OCR 识别成功（外部服务）"
                        }
                    else:
                        logger.warning(f"外部 OCR 服务返回异常: {result}")
                else:
                    logger.warning(f"外部 OCR 服务返回错误: {response.status_code}")

            except requests.exceptions.Timeout:
                logger.warning("外部 OCR 服务请求超时")
            except requests.exceptions.RequestException as e:
                logger.warning(f"外部 OCR 服务请求失败: {str(e)}")
            except Exception as e:
                logger.warning(f"外部 OCR 服务调用异常: {str(e)}")

        # 方法2：使用 rapidocr（推荐，更快）- 需要系统 OpenGL 库
        try:
            from rapidocr_onnxruntime import RapidOCR
            ocr = RapidOCR()
            result, _ = ocr(image_path)
            if result:
                # 按行合并文字
                text_lines = [line[1] for line in result]
                text = "\n".join(text_lines)
                return {
                    "success": True,
                    "text": text,
                    "message": "OCR 识别成功（RapidOCR）"
                }
        except ImportError:
            logger.info("RapidOCR 未安装，尝试使用 pytesseract")
        except Exception as e:
            logger.warning(f"RapidOCR 失败: {str(e)}")

        # 方法3：使用 pytesseract - 需要系统安装 tesseract-ocr
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(image_path)
            # 设置为中英文混合识别
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            if text and len(text.strip()) >= 10:
                return {
                    "success": True,
                    "text": text,
                    "message": "OCR 识别成功（pytesseract）"
                }
            else:
                return {
                    "success": False,
                    "text": "",
                    "message": "无法识别图片中的文字，请确保图片清晰且包含文字内容"
                }
        except ImportError:
            logger.warning("pytesseract 未安装")
        except Exception as e:
            logger.warning(f"pytesseract OCR 失败: {str(e)}")

        # 所有 OCR 方法都不可用
        return {
            "success": False,
            "text": "",
            "message": "OCR 功能不可用：请配置外部 OCR 服务或安装 RapidOCR/pytesseract"
        }

    def _is_pdf_scanned(self, pdf_path: str) -> bool:
        """
        检测 PDF 是否为扫描版（图片型）

        Returns:
            True 如果是扫描版，False 如果是文字版
        """
        try:
            doc = fitz.open(pdf_path)
            page_count = doc.page_count

            # 检查前3页或全部页数（如果少于3页）
            pages_to_check = min(3, page_count)
            text_pages = 0

            for page_num in range(pages_to_check):
                page = doc[page_num]
                text = page.get_text()

                # 如果页面文字很少，可能是扫描版
                if len(text.strip()) > 50:
                    text_pages += 1

            doc.close()

            # 如果检查的页面中文字页少于一半，判定为扫描版
            return text_pages < (pages_to_check / 2)

        except Exception as e:
            logger.warning(f"检测 PDF 类型失败: {str(e)}")
            return False

    def _convert_pdf_to_docx(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """转换 PDF 到 .docx（优先使用 MinerU 服务）"""
        try:
            # 方法1：优先使用 MinerU 服务
            if self.mineru_url:
                logger.info(f"使用 MinerU 服务处理 PDF: {input_path}")
                result = self._convert_pdf_with_mineru(input_path, output_path)
                if result[0] == ConversionResult.SUCCESS:
                    return result
                else:
                    logger.warning(f"MinerU 转换失败: {result[1]}，尝试使用本地方法")

            # 方法2：检测是否为扫描版 PDF
            is_scanned = self._is_pdf_scanned(input_path)

            if is_scanned:
                logger.info(f"检测到扫描版 PDF，使用 OCR 处理: {input_path}")
                return self._convert_scanned_pdf_to_docx(input_path, output_path)
            else:
                # 文字版 PDF，尝试使用 PyMuPDF 直接提取文字创建文档
                logger.info(f"检测到文字版 PDF，使用 PyMuPDF 提取文字: {input_path}")
                return self._convert_text_pdf_with_pymupdf(input_path, output_path)

        except Exception as e:
            logger.error(f"PDF 转换失败: {str(e)}")
            return ConversionResult.FAILED, f"PDF 转换失败: {str(e)}", None

    def _convert_pdf_with_mineru(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """使用 MinerU 服务转换 PDF"""
        try:
            from app.core.config import settings

            # 读取 PDF 文件
            with open(input_path, 'rb') as f:
                files = {'file': (os.path.basename(input_path), f, 'application/pdf')}
                response = requests.post(
                    self.mineru_url,
                    files=files,
                    timeout=settings.MINERU_API_TIMEOUT
                )

            if response.status_code == 200:
                # MinerU 返回 markdown 或 docx 内容
                content_type = response.headers.get('Content-Type', '')

                if 'application/json' in content_type:
                    # JSON 响应格式
                    result = response.json()
                    if result.get("status") == "success":
                        # 检查返回的内容类型
                        if "docx_content" in result:
                            # Base64 编码的 docx 内容
                            import base64
                            docx_content = base64.b64decode(result["docx_content"])
                            with open(output_path, 'wb') as f:
                                f.write(docx_content)
                            # 后处理
                            if self.ENABLE_POSTPROCESSING:
                                self._postprocess_docx(output_path)
                            metadata = self._extract_metadata(output_path)
                            metadata["pdf_type"] = "text"
                            metadata["mineru_used"] = True
                            return ConversionResult.SUCCESS, output_path, metadata
                        elif "text" in result:
                            # 纯文本内容，创建 docx
                            return self._create_docx_from_text(result["text"], output_path, input_path)
                    else:
                        return ConversionResult.FAILED, result.get("message", "MinerU 处理失败"), None

                elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
                    # 直接返回 docx 文件
                    with open(output_path, 'wb') as f:
                        f.write(response.content)
                    # 后处理
                    if self.ENABLE_POSTPROCESSING:
                        self._postprocess_docx(output_path)
                    metadata = self._extract_metadata(output_path)
                    metadata["pdf_type"] = "text"
                    metadata["mineru_used"] = True
                    return ConversionResult.SUCCESS, output_path, metadata

                else:
                    # 其他响应类型，尝试解析为文本
                    text_content = response.text
                    if len(text_content) > 100:
                        return self._create_docx_from_text(text_content, output_path, input_path)

            return ConversionResult.FAILED, f"MinerU 服务返回错误: {response.status_code}", None

        except requests.exceptions.Timeout:
            return ConversionResult.FAILED, "MinerU 服务请求超时", None
        except requests.exceptions.RequestException as e:
            return ConversionResult.FAILED, f"MinerU 服务请求失败: {str(e)}", None
        except Exception as e:
            logger.warning(f"MinerU 转换异常: {str(e)}")
            return ConversionResult.FAILED, f"MinerU 转换失败: {str(e)}", None

    def _create_docx_from_text(self, text: str, output_path: str, original_filename: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """从文本内容创建 docx 文档"""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(os.path.basename(original_filename), 0)

            # 按段落添加文本
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

            doc.save(output_path)

            # 后处理
            if self.ENABLE_POSTPROCESSING:
                self._postprocess_docx(output_path)

            metadata = self._extract_metadata(output_path)
            metadata["pdf_type"] = "text"
            metadata["mineru_used"] = True
            return ConversionResult.SUCCESS, output_path, metadata

        except Exception as e:
            return ConversionResult.FAILED, f"创建 docx 文档失败: {str(e)}", None

    def _convert_text_pdf_with_pymupdf(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """使用 PyMuPDF 提取 PDF 文字并创建 docx"""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(os.path.basename(input_path), 0)

            # 打开 PDF
            pdf_doc = fitz.open(input_path)
            total_text = []
            page_count = pdf_doc.page_count

            for page_num in range(page_count):
                page = pdf_doc[page_num]
                text = page.get_text()
                if text.strip():
                    total_text.append(f"--- 第 {page_num + 1} 页 ---\n{text}")

            pdf_doc.close()

            if not total_text:
                return ConversionResult.FAILED, "PDF 中未提取到文字内容", None

            # 添加所有文字
            doc.add_paragraph("\n".join(total_text))
            doc.save(output_path)

            # 后处理
            if self.ENABLE_POSTPROCESSING:
                self._postprocess_docx(output_path)

            metadata = self._extract_metadata(output_path)
            metadata["pdf_type"] = "text"
            metadata["extraction_method"] = "PyMuPDF"
            metadata["page_count"] = page_count

            return ConversionResult.SUCCESS, output_path, metadata

        except Exception as e:
            logger.warning(f"PyMuPDF 提取失败: {str(e)}，尝试使用 LibreOffice")
            # 回退到 LibreOffice
            return self._convert_text_pdf_to_docx(input_path, output_path)

    def _convert_text_pdf_to_docx(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """转换文字版 PDF 到 .docx（使用 LibreOffice）"""
        # 直接复用 _convert_with_libreoffice 方法
        result, converted_path, metadata = self._convert_with_libreoffice(input_path, output_path)

        if result == ConversionResult.SUCCESS and metadata:
            metadata["pdf_type"] = "text"

        return result, converted_path, metadata

    def _convert_scanned_pdf_to_docx(self, input_path: str, output_path: str) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """转换扫描版（图片型）PDF 到 .docx（使用OCR）"""
        try:
            from docx import Document
            from docx.shared import Inches
            import io

            doc = Document()

            # 添加标题
            doc.add_heading(os.path.basename(input_path), 0)
            doc.add_paragraph("【OCR 识别结果 - 扫描版 PDF】")

            # 打开 PDF
            pdf_doc = fitz.open(input_path)
            total_text = []
            page_count = pdf_doc.page_count
            ocr_unavailable = False

            for page_num in range(page_count):
                page = pdf_doc[page_num]

                # 方法1：尝试直接提取文字
                text = page.get_text()
                if len(text.strip()) > 20:
                    total_text.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
                else:
                    # 方法2：渲染页面为图片，然后 OCR
                    try:
                        # 渲染页面为图片（300 DPI）
                        mat = fitz.Matrix(2.0, 2.0)  # 2x 缩放
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")

                        # 保存临时图片
                        temp_img_path = os.path.join(self.upload_dir, f"temp_page_{page_num}.png")
                        with open(temp_img_path, 'wb') as f:
                            f.write(img_data)

                        # OCR 识别
                        ocr_result = self._ocr_image(temp_img_path)
                        if ocr_result["success"] and ocr_result["text"]:
                            total_text.append(f"--- 第 {page_num + 1} 页 ---\n{ocr_result['text']}")
                        elif not ocr_result["success"] and "不可用" in ocr_result.get("message", ""):
                            # OCR 完全不可用
                            ocr_unavailable = True
                            total_text.append(f"--- 第 {page_num + 1} 页 ---\n[OCR 功能不可用]")

                        # 清理临时图片
                        try:
                            os.remove(temp_img_path)
                        except:
                            pass

                    except Exception as e:
                        logger.warning(f"页面 {page_num + 1} OCR 失败: {str(e)}")
                        total_text.append(f"--- 第 {page_num + 1} 页 ---\n[无法识别]")

            pdf_doc.close()

            # 添加所有识别的文字
            if total_text:
                doc.add_paragraph("\n".join(total_text))
            else:
                doc.add_paragraph("[未能识别到文字内容，请检查PDF图片质量]")

            # 保存文档
            doc.save(output_path)

            # 后处理
            if self.ENABLE_POSTPROCESSING:
                self._postprocess_docx(output_path)

            metadata = self._extract_metadata(output_path)
            metadata["pdf_type"] = "scanned"
            metadata["ocr_used"] = True
            metadata["page_count"] = page_count

            # 如果 OCR 不可用，返回警告而不是成功
            if ocr_unavailable:
                return ConversionResult.FAILED, "OCR 功能不可用：扫描版 PDF 需要安装 RapidOCR 或 pytesseract 及其系统依赖", None

            return ConversionResult.SUCCESS, output_path, metadata

        except Exception as e:
            logger.error(f"扫描版 PDF 转换失败: {str(e)}")
            return ConversionResult.FAILED, f"扫描版 PDF 处理失败: {str(e)}", None

    def _detect_encoding(self, file_path: str) -> str:
        """检测文本文件编码"""
        import chardet

        with open(file_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            return result.get('encoding', 'utf-8')

    def _extract_metadata(self, file_path: str) -> Dict:
        """
        提取文档元数据
        Args:
            file_path: 文件路径
        Returns:
            元数据字典
        """
        # 修复说明：添加了 "file_path": file_path 字段
        metadata = {
            "file_path": file_path,  # ✅ [关键修复] 添加完整文件路径，防止下游字典键覆盖
            "original_filename": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "format": self.detect_format(file_path).value,
        }
        # 如果是 .docx，提取更多信息
        if file_path.endswith('.docx'):
            try:
                from docx import Document
                doc = Document(file_path)
                # 统计段落数
                metadata["paragraph_count"] = len(doc.paragraphs)
                # 统计表数
                metadata["table_count"] = len(doc.tables)
                # 统计页数（近似）
                metadata["page_count"] = self._estimate_page_count(doc)
                # 统计字数
                metadata["word_count"] = sum(len(p.text.split()) for p in doc.paragraphs)
            except Exception as e:
                logger.warning(f"提取 .docx 元数据失败: {str(e)}")
        return metadata

    def _estimate_page_count(self, doc) -> int:
        """估算文档页数"""
        # 简单估算：每页约 500 字
        total_words = sum(len(p.text.split()) for p in doc.paragraphs)
        return max(1, total_words // 500)

    def preprocess(
        self,
        file_path: str,
        output_filename: str = None,
        delete_original: bool = False
    ) -> Tuple[ConversionResult, str, Optional[Dict]]:
        """
        预处理文件（完整流程）

        流程：
        1. 验证文件
        2. 检测格式
        3. 转换为 .docx
        4. 提取元数据
        5. 可选：删除原始文件

        Args:
            file_path: 输入文件路径
            output_filename: 输出文件名（可选）
            delete_original: 是否删除原始文件

        Returns:
            (转换结果, 输出文件路径, 元数据)
        """
        # 验证文件
        is_valid, error_msg = self.validate_file(file_path)
        if not is_valid:
            return ConversionResult.FAILED, error_msg, None

        # 转换文件
        result, output_path, metadata = self.convert_to_docx(file_path, output_filename)

        # 转换成功，可选删除原文件
        if result == ConversionResult.SUCCESS and delete_original:
            try:
                os.remove(file_path)
                logger.info(f"已删除原始文件: {file_path}")
            except Exception as e:
                logger.warning(f"删除原始文件失败: {str(e)}")

        return result, output_path, metadata

    def batch_preprocess(
        self,
        file_paths: List[str],
        delete_originals: bool = False
    ) -> List[Dict]:
        """
        批量预处理文件

        Args:
            file_paths: 文件路径列表
            delete_originals: 是否删除原始文件

        Returns:
            处理结果列表
        """
        results = []

        for file_path in file_paths:
            result, output_path, metadata = self.preprocess(
                file_path,
                delete_original=delete_originals
            )

            results.append({
                "original_path": file_path,
                "result": result.value,
                "output_path": output_path,
                "metadata": metadata
            })

        return results

    def extract_text(self, file_path: str) -> str:
        """
        提取文件纯文本内容 (Markdown 格式)
        使用更健壮的方法处理格式不规范的文档
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_format = self.detect_format(file_path)
        content = []

        try:
            # 1. Word (.docx) - 使用健壮的方法
            if file_format == DocumentFormat.DOCX:
                # 使用 docx2python 作为更健壮的替代方案
                try:
                    from docx2python import docx2python

                    # 提取文本内容
                    result = docx2python(file_path)

                    # 处理正文内容
                    if result.body:
                        for paragraph in result.body:
                            # paragraph 是嵌套列表，需要展平
                            paragraph_text = []
                            for line in paragraph:
                                if line:  # 跳过空行
                                    # line 是单元格列表，用制表符分隔
                                    paragraph_text.append('\t'.join(line))
                            if paragraph_text:
                                content.append(' '.join(paragraph_text))

                    # 处理表格内容（如果有的话）
                    if result.tables:
                        for table in result.tables:
                            for row in table:
                                row_text = [cell if cell else '' for cell in row]
                                content.append(' | '.join(row_text))

                except ImportError:
                    # docx2python 不可用，使用 python-docx 并处理异常
                    logger.info("docx2python 不可用，使用 python-docx（可能无法处理格式不规范的文档）")
                    content = self._extract_text_with_python_docx(file_path)
                except Exception as e:
                    logger.warning(f"docx2python 处理失败，尝试使用 python-docx: {str(e)}")
                    content = self._extract_text_with_python_docx(file_path)

            # 2. Word (.doc) -> 先转 docx 再提
            elif file_format == DocumentFormat.DOC:
                # 复用 convert_to_docx 转换到临时目录
                temp_docx = file_path + ".temp.docx"
                res, path, _ = self._convert_doc_to_docx(file_path, temp_docx)
                if res == ConversionResult.SUCCESS:
                    return self.extract_text(path) # 递归调用

            # 3. PDF
            elif file_format == DocumentFormat.PDF:
                # 复用 _is_pdf_scanned 判断
                if self._is_pdf_scanned(file_path):
                    # 扫描版：复用 OCR 逻辑
                    # 这里为了简化，直接调用 _convert_scanned_pdf_to_docx 再提文本有点重
                    # 建议直接调用 fitz 提取图片再 OCR，或者直接返回提示
                    return self._extract_text_from_scanned_pdf(file_path)
                else:
                    # 文字版：直接用 fitz 提取
                    doc = fitz.open(file_path)
                    for page in doc:
                        text = page.get_text()
                        if text:
                            content.append(text)
                    doc.close()

            # 4. 图片
            elif file_format in self.IMAGE_FORMATS:
                ocr_res = self._ocr_image(file_path)
                if ocr_res["success"]:
                    content.append(ocr_res["text"])

            # 5. 纯文本
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()

        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            return ""

        return "\n\n".join(content)

    def _extract_text_with_python_docx(self, file_path: str) -> list:
        """
        使用 python-docx 提取文本（带异常处理）
        这是对格式不规范文档的降级方案
        """
        content = []
        try:
            from docx import Document
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.document import Document as DocxDocument
            from docx.table import _Cell, Table
            from lxml import etree

            # 方法1：尝试使用 Document 类（可能因表格格式问题失败）
            try:
                doc = Document(file_path)

                # 提取段落
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        # 简单恢复标题层级
                        if 'Heading 1' in para.style.name:
                            content.append(f"# {text}")
                        elif 'Heading 2' in para.style.name:
                            content.append(f"## {text}")
                        elif 'Heading 3' in para.style.name:
                            content.append(f"### {text}")
                        else:
                            content.append(text)

                # 提取表格（这可能失败）
                try:
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells]
                            content.append(" | ".join(row_text))
                except Exception as e:
                    logger.warning(f"表格提取失败，跳过表格: {str(e)}")

            except Exception as doc_error:
                # 方法2：直接解析 XML，绕过 python-docx 的表格验证
                logger.warning(f"python-docx 解析失败，尝试直接解析 XML: {str(doc_error)}")
                content = self._extract_text_from_xml(file_path)

        except ImportError:
            logger.error("python-docx 未安装")
        except Exception as e:
            logger.error(f"python-docx 提取文本失败: {str(e)}")

        return content

    def _extract_text_from_xml(self, file_path: str) -> list:
        """
        直接从 docx XML 中提取文本（最健壮的方法）
        绕过 python-docx 的验证，直接解析 XML
        """
        content = []
        try:
            import zipfile
            from xml.etree import ElementTree as ET

            # docx 文件是一个 zip 压缩包
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                # 读取主文档内容
                xml_content = zip_ref.read('word/document.xml')

            # 解析 XML
            root = ET.fromstring(xml_content)

            # 定义命名空间
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }

            # 提取所有段落
            paragraphs = root.findall('.//w:p', namespaces)

            for para in paragraphs:
                # 提取段落中的所有文本
                texts = []
                for text_elem in para.findall('.//w:t', namespaces):
                    if text_elem.text:
                        texts.append(text_elem.text)

                if texts:
                    paragraph_text = ''.join(texts).strip()
                    if paragraph_text:
                        content.append(paragraph_text)

            # 提取表格（绕过 tblGrid 验证）
            tables = root.findall('.//w:tbl', namespaces)
            for table in tables:
                rows = table.findall('.//w:tr', namespaces)
                for row in rows:
                    cells = row.findall('.//w:tc', namespaces)
                    row_texts = []
                    for cell in cells:
                        cell_texts = []
                        for text_elem in cell.findall('.//w:t', namespaces):
                            if text_elem.text:
                                cell_texts.append(text_elem.text)
                        row_texts.append(''.join(cell_texts).strip())
                    content.append(' | '.join(row_texts))

        except Exception as e:
            logger.error(f"XML 解析失败: {str(e)}")

        return content

    def _extract_text_from_scanned_pdf(self, pdf_path: str) -> str:
        """辅助方法：从扫描 PDF 提取文本"""
        # 简单实现：只提取前5页 OCR，避免太慢
        doc = fitz.open(pdf_path)
        texts = []
        for i, page in enumerate(doc):
            if i >= 5: break
            pix = page.get_pixmap()
            temp_img = f"{pdf_path}_{i}.png"
            pix.save(temp_img)
            ocr_res = self._ocr_image(temp_img)
            if ocr_res["success"]:
                texts.append(ocr_res["text"])
            os.remove(temp_img)
        doc.close()
        return "\n\n".join(texts)


# 单例实例，便于其他模块复用
_preprocessor_instance = None


def get_preprocessor() -> DocumentPreprocessor:
    """获取预处理中心单例"""
    global _preprocessor_instance
    if _preprocessor_instance is None:
        _preprocessor_instance = DocumentPreprocessor()
    return _preprocessor_instance