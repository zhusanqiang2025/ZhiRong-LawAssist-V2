# backend/app/services/contract_generation/rag/template_indexer.py
"""
模板索引服务

负责将合同模板从数据库索引到向量存储。
处理 .docx 文件解析、文本预处理和元数据提取。
"""
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from datetime import datetime

from sqlalchemy.orm import Session

from app.models.contract_template import ContractTemplate

# 可选导入 ChromaDB (在 Python 3.14 + pydantic v2 环境下可能不可用)
try:
    from .vector_store import TemplateVectorStore, get_vector_store
    _chromadb_available = True
except (ImportError, Exception):
    TemplateVectorStore = None
    get_vector_store = None
    _chromadb_available = False

from .pgvector_store import PgVectorStore, get_pgvector_store
from app.core.config import settings

logger = logging.getLogger(__name__)


class TemplateIndexer:
    """
    模板索引服务

    负责将数据库中的合同模板索引到向量存储中。
    """

    def __init__(self, vector_store: Optional[Union[TemplateVectorStore, PgVectorStore]] = None):
        """
        初始化模板索引器

        Args:
            vector_store: 向量存储实例,默认使用全局单例 (优先pgvector)
        """
        # 优先使用pgvector,如果不可用则回退到ChromaDB
        if vector_store is None:
            try:
                from app.database import SessionLocal
                db = SessionLocal()
                # 测试pgvector是否可用
                test_store = get_pgvector_store()
                test_store.get_stats(db)
                db.close()
                vector_store = test_store
                logger.info("TemplateIndexer initialized (using pgvector)")
            except Exception as e:
                logger.warning(f"pgvector不可用,回退到ChromaDB: {e}")
                if get_vector_store is None:
                    raise RuntimeError("既无法使用pgvector,也无法使用ChromaDB (ChromaDB未安装)")
                vector_store = get_vector_store()
                logger.info("TemplateIndexer initialized (using ChromaDB fallback)")

        self.vector_store = vector_store
        logger.info("TemplateIndexer initialized")

    # ==================== 文档解析 ====================

    def _parse_docx_file(self, file_path: str) -> str:
        """
        解析 .docx 文件，提取文本内容

        Args:
            file_path: 文档文件路径

        Returns:
            str: 提取的文本内容

        Raises:
            FileNotFoundError: 文件不存在
            Exception: 解析失败
        """
        try:
            from docx import Document

            doc = Document(file_path)

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # 跳过空段落
                    paragraphs.append(text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n".join(paragraphs)

            logger.debug(f"Parsed DOCX file: {file_path} ({len(full_text)} chars)")
            return full_text

        except Exception as e:
            logger.error(f"Failed to parse DOCX file {file_path}: {e}")
            raise

    def _parse_markdown_file(self, file_path: str) -> tuple:
        """
        解析 .md 文件，提取文本内容和元数据

        Args:
            file_path: Markdown 文件路径

        Returns:
            tuple: (文本内容, 元数据字典)
        """
        try:
            import re
            import yaml

            # 读取文件内容
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 解析 YAML Frontmatter
            metadata = {}
            text_content = content

            # 检查是否有 YAML Frontmatter（以 --- 开头和结尾）
            if content.startswith('---'):
                # 找到第二个 ---
                end_match = re.search(r'\n---\s*\n', content)
                if end_match:
                    yaml_text = content[4:end_match.start()]  # 去掉开头的 ---
                    text_content = content[end_match.end():]  # 去掉 YAML 部分

                    # 解析 YAML
                    try:
                        metadata = yaml.safe_load(yaml_text) or {}
                    except yaml.YAMLError as e:
                        logger.warning(f"Failed to parse YAML frontmatter in {file_path}: {e}")

            logger.debug(f"Parsed Markdown file: {file_path} ({len(text_content)} chars)")
            logger.debug(f"Extracted metadata: {metadata}")

            return text_content, metadata

        except Exception as e:
            logger.error(f"Failed to parse Markdown file {file_path}: {e}")
            # 返回空内容和空元数据
            return "", {}

    def _parse_template_file(self, file_path: str) -> tuple:
        """
        解析模板文件，自动识别格式并提取文本内容

        支持 .docx 和 .md 格式，对于 .md 文件会解析 YAML Frontmatter

        Args:
            file_path: 模板文件路径

        Returns:
            tuple: (文本内容, 元数据字典)
        """
        # 根据文件扩展名选择解析方法
        file_ext = Path(file_path).suffix.lower()

        if file_ext == '.docx':
            # DOCX 文件没有 YAML 元数据
            content = self._parse_docx_file(file_path)
            return content, {}
        elif file_ext in ['.md', '.markdown']:
            # Markdown 文件，解析 YAML Frontmatter
            return self._parse_markdown_file(file_path)
        else:
            # 尝试作为纯文本文件读取
            logger.warning(f"Unknown file format: {file_ext}, trying to read as plain text")
            return self._parse_markdown_file(file_path)

    def _prepare_indexing_text(
        self,
        template: ContractTemplate,
        parsed_content: str,
        file_metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        准备用于索引的文本内容

        将模板名称、描述、类别、文件元数据等与文件内容组合，
        提升检索时的语义匹配效果。

        Args:
            template: 合同模板模型实例
            parsed_content: 从文件解析的文本内容
            file_metadata: 从 Markdown 文件解析的 YAML 元数据

        Returns:
            str: 组合后的索引文本
        """
        # 构建增强的索引文本
        parts = []

        # 1. 模板名称（重要，权重最高）
        if template.name:
            parts.append(f"模板名称：{template.name}")

        # 2. 类别和子类别
        if template.category:
            parts.append(f"类别：{template.category}")
        if template.subcategory:
            parts.append(f"子类别：{template.subcategory}")

        # 3. 文件元数据（如果有）
        if file_metadata:
            # category: 兜底模板/民法典典型合同等
            if file_metadata.get('category'):
                parts.append(f"模板分类：{file_metadata['category']}")
            # type: 免责声明/技术服务合同等
            if file_metadata.get('type'):
                parts.append(f"合同类型：{file_metadata['type']}")
            # scenario: 风险控制/技术开发等
            if file_metadata.get('scenario'):
                parts.append(f"使用场景：{file_metadata['scenario']}")
            # tags: 标签列表
            tags = file_metadata.get('tags', [])
            if tags and isinstance(tags, list):
                tags_str = "、".join(tags)
                parts.append(f"标签：{tags_str}")

        # 4. 描述
        if template.description:
            parts.append(f"描述：{template.description}")

        # 5. 关键词
        if template.keywords:
            keywords_str = "、".join(template.keywords)
            parts.append(f"关键词：{keywords_str}")

        # 6. 使用场景（数据库字段）
        if template.usage_scenario:
            parts.append(f"使用场景：{template.usage_scenario}")

        # 7. 司法管辖区
        if template.jurisdiction:
            parts.append(f"适用地区：{template.jurisdiction}")

        # 8. 文件内容（限制长度，避免嵌入超过模型限制）
        # BGE-M3 支持 8192 tokens，约 3 万汉字
        max_content_length = 25000  # 保留余量
        if len(parsed_content) > max_content_length:
            parsed_content = parsed_content[:max_content_length] + "..."
        parts.append(f"内容：\n{parsed_content}")

        return "\n\n".join(parts)

    def _extract_metadata(self, template: ContractTemplate) -> Dict[str, Any]:
        """
        提取模板元数据（用于存储在向量数据库中）

        优先使用知识图谱法律特征，如果模板已关联知识图谱。

        Args:
            template: 合同模板模型实例

        Returns:
            Dict[str, Any]: 元数据字典（包含知识图谱法律特征）
        """
        # 尝试从 metadata_info 中获取知识图谱关联信息
        knowledge_graph_features = None
        if template.metadata_info and isinstance(template.metadata_info, dict):
            kg_link = template.metadata_info.get("knowledge_graph_link", {})
            if kg_link:
                knowledge_graph_features = kg_link.get("legal_features")

        # 构建基础元数据
        metadata = {
            # 基础信息
            "name": template.name,
            "category": template.category,
            "subcategory": template.subcategory or "",
            "description": template.description or "",
            "tags": ",".join(template.tags) if template.tags else "",
            "keywords": ",".join(template.keywords) if template.keywords else "",
            "file_name": template.file_name,
            "file_url": template.file_url,
            "file_type": template.file_type,

            # 权限信息
            "is_public": template.is_public,
            "owner_id": template.owner_id or 0,

            # 统计信息
            "download_count": template.download_count,
            "rating": template.rating,
            "status": template.status,

            # 语言和地区
            "language": template.language,
            "jurisdiction": template.jurisdiction or "",

            # 时间戳
            "created_at": template.created_at.isoformat() if template.created_at else "",
            "indexed_at": datetime.now().isoformat(),

            # ==================== 知识图谱法律特征（推荐使用） ====================
            "knowledge_graph_linked": knowledge_graph_features is not None,
        }

        # 如果有知识图谱特征，使用它；否则回退到已废弃的V2字段
        if knowledge_graph_features:
            # 使用知识图谱法律特征
            metadata.update({
                # 知识图谱核心法律特征
                "kg_transaction_nature": knowledge_graph_features.get("transaction_nature", ""),
                "kg_contract_object": knowledge_graph_features.get("contract_object", ""),
                "kg_stance": knowledge_graph_features.get("stance", ""),
                "kg_consideration_type": knowledge_graph_features.get("consideration_type", ""),
                "kg_consideration_detail": knowledge_graph_features.get("consideration_detail", ""),
                "kg_transaction_characteristics": knowledge_graph_features.get("transaction_characteristics", ""),
                "kg_usage_scenario": knowledge_graph_features.get("usage_scenario", ""),
                "kg_legal_basis": ",".join(knowledge_graph_features.get("legal_basis", [])),
                # 匹配的合同类型名称
                "kg_matched_contract_type": template.metadata_info.get("knowledge_graph_link", {}).get("matched_contract_type", ""),
            })
        else:
            # 回退到已废弃的V2字段（用于向后兼容）
            metadata.update({
                # V2 四维法律特征（已废弃）
                "transaction_nature": template.transaction_nature or "",
                "contract_object": template.contract_object or "",
                "stance": template.stance or "",
                "complexity": template.complexity or "",
                # V2 扩展字段（已废弃）
                "transaction_consideration": template.transaction_consideration or "",
                "transaction_characteristics": template.transaction_characteristics or "",
                # V2 结构锚点字段（已废弃）
                "primary_contract_type": template.primary_contract_type or "",
                "delivery_model": template.delivery_model or "",
                "payment_model": template.payment_model or "",
                "risk_level": template.risk_level or "",
                "is_recommended": template.is_recommended,
            })

        return metadata

    # ==================== 单个模板索引 ====================

    def index_template(
        self,
        template: ContractTemplate,
        file_path: Optional[str] = None,
        reindex: bool = False
    ) -> bool:
        """
        索引单个模板到向量存储

        Args:
            template: 合同模板模型实例
            file_path: 模板文件路径，默认使用 template.file_url
            reindex: 是否重新索引（如果已存在则更新）

        Returns:
            bool: 是否成功
        """
        try:
            # 确定文件路径
            if file_path is None:
                file_path = template.file_url

            # 检查文件是否存在
            if not Path(file_path).exists():
                logger.error(f"Template file not found: {file_path}")
                return False

            # 检查文件格式，只索引 Markdown 文件
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in ['.md', '.markdown']:
                logger.debug(f"Skipping non-Markdown file: {file_path} (format: {file_ext})")
                return False  # 跳过非 Markdown 文件

            # 解析文档（支持 .docx 和 .md 格式，返回内容和元数据）
            parsed_content, file_metadata = self._parse_template_file(file_path)

            # 准备索引文本（传入文件元数据以增强索引）
            indexing_text = self._prepare_indexing_text(template, parsed_content, file_metadata)

            # 提取元数据
            metadata = self._extract_metadata(template)

            # 添加到向量存储
            if reindex:
                self.vector_store.update_template(
                    template_id=template.id,
                    text=indexing_text,
                    metadata=metadata,
                    user_id=template.owner_id,
                    is_public=template.is_public
                )
                logger.info(f"Re-indexed template: {template.name} ({template.id})")
            else:
                self.vector_store.add_template(
                    template_id=template.id,
                    text=indexing_text,
                    metadata=metadata,
                    user_id=template.owner_id,
                    is_public=template.is_public
                )
                logger.info(f"Indexed template: {template.name} ({template.id})")

            return True

        except Exception as e:
            logger.error(f"Failed to index template {template.id}: {e}")
            return False

    # ==================== 批量索引 ====================

    def index_templates_batch(
        self,
        templates: List[ContractTemplate],
        reindex: bool = False
    ) -> Dict[str, int]:
        """
        批量索引模板

        Args:
            templates: 合同模板列表
            reindex: 是否重新索引

        Returns:
            Dict[str, int]: 索引统计 {"success": 成功数, "failed": 失败数}
        """
        success_count = 0
        failed_count = 0

        logger.info(f"Starting batch indexing of {len(templates)} templates")

        for template in templates:
            if self.index_template(template, reindex=reindex):
                success_count += 1
            else:
                failed_count += 1

        # 持久化向量存储
        self.vector_store.persist()

        logger.info(f"Batch indexing completed: {success_count} succeeded, {failed_count} failed")

        return {
            "success": success_count,
            "failed": failed_count,
            "total": len(templates)
        }

    def index_all_templates(
        self,
        db: Session,
        category: Optional[str] = None,
        is_public: Optional[bool] = None,
        reindex: bool = False
    ) -> Dict[str, int]:
        """
        索引数据库中的所有模板（或符合条件的模板）

        Args:
            db: 数据库会话
            category: 按类别过滤，None 表示全部
            is_public: 按权限过滤，None 表示全部
            reindex: 是否重新索引

        Returns:
            Dict[str, int]: 索引统计
        """
        # 构建查询
        query = db.query(ContractTemplate)

        if category:
            query = query.filter(ContractTemplate.category == category)

        if is_public is not None:
            query = query.filter(ContractTemplate.is_public == is_public)

        # 只索引激活状态的模板
        query = query.filter(ContractTemplate.status == "active")

        templates = query.all()

        logger.info(f"Found {len(templates)} templates to index from database")

        return self.index_templates_batch(templates, reindex=reindex)

    def index_user_templates(
        self,
        db: Session,
        user_id: int,
        reindex: bool = False
    ) -> Dict[str, int]:
        """
        索引指定用户的所有私有模板

        Args:
            db: 数据库会话
            user_id: 用户 ID
            reindex: 是否重新索引

        Returns:
            Dict[str, int]: 索引统计
        """
        templates = db.query(ContractTemplate).filter(
            ContractTemplate.owner_id == user_id,
            ContractTemplate.is_public == False,
            ContractTemplate.status == "active"
        ).all()

        logger.info(f"Found {len(templates)} private templates for user {user_id}")

        return self.index_templates_batch(templates, reindex=reindex)

    # ==================== 删除和清理 ====================

    def remove_template(self, template_id: str, user_id: Optional[int] = None, is_public: bool = True) -> bool:
        """
        从向量存储中移除模板

        Args:
            template_id: 模板 ID
            user_id: 用户 ID（私有模板必需）
            is_public: 是否为公共模板

        Returns:
            bool: 是否成功
        """
        try:
            self.vector_store.delete_template(
                template_id=template_id,
                user_id=user_id,
                is_public=is_public
            )
            logger.info(f"Removed template {template_id} from vector store")
            return True
        except Exception as e:
            logger.error(f"Failed to remove template {template_id}: {e}")
            return False

    def clear_all_indexed(self) -> bool:
        """
        清空所有已索引的模板（危险操作！）

        Returns:
            bool: 是否成功
        """
        try:
            self.vector_store.clear_collection(is_public=True)
            logger.warning("Cleared all indexed templates from vector store")
            return True
        except Exception as e:
            logger.error(f"Failed to clear indexed templates: {e}")
            return False

    def clear_user_indexed(self, user_id: int) -> bool:
        """
        清空指定用户的所有已索引模板

        Args:
            user_id: 用户 ID

        Returns:
            bool: 是否成功
        """
        try:
            self.vector_store.delete_user_collection(user_id)
            logger.info(f"Cleared all indexed templates for user {user_id}")
            return True
        except Exception as e:
            logger.error(f"Failed to clear user {user_id} indexed templates: {e}")
            return False

    # ==================== 维护和监控 ====================

    def get_indexing_stats(self, db: Session) -> Dict[str, Any]:
        """
        获取索引统计信息

        Args:
            db: 数据库会话

        Returns:
            Dict[str, Any]: 索引统计
        """
        # 数据库中的模板数量
        db_total = db.query(ContractTemplate).filter(
            ContractTemplate.status == "active"
        ).count()

        db_public = db.query(ContractTemplate).filter(
            ContractTemplate.is_public == True,
            ContractTemplate.status == "active"
        ).count()

        db_private = db_total - db_public

        # 向量存储中的模板数量
        vector_stats = self.vector_store.get_collection_stats(is_public=True)

        return {
            "database": {
                "total": db_total,
                "public": db_public,
                "private": db_private
            },
            "vector_store": {
                "public_indexed": vector_stats["count"],
                "collection_name": vector_stats["name"]
            },
            "coverage": f"{(vector_stats['count'] / db_public * 100):.1f}%" if db_public > 0 else "N/A"
        }

    def rebuild_index(self, db: Session) -> Dict[str, int]:
        """
        重建整个索引（清空后重新索引所有模板）

        Args:
            db: 数据库会话

        Returns:
            Dict[str, int]: 索引统计
        """
        logger.warning("Rebuilding entire template index...")

        # 清空现有索引
        self.clear_all_indexed()

        # 重新索引所有公共模板
        return self.index_all_templates(db, is_public=True, reindex=False)


# ==================== 单例模式 ====================

_template_indexer_instance: Optional[TemplateIndexer] = None


def get_template_indexer() -> TemplateIndexer:
    """
    获取模板索引器单例

    Returns:
        TemplateIndexer: 模板索引器实例
    """
    global _template_indexer_instance
    if _template_indexer_instance is None:
        _template_indexer_instance = TemplateIndexer()
    return _template_indexer_instance
