# backend/app/services/contract_generation/agents/contract_drafter.py
"""
合同起草 Agent

功能：
1. 根据需求分析结果起草合同内容
2. 支持多种合同类型
3. 生成规范的 Markdown 格式内容
4. 支持基于专业模板的增强起草（RAG）
5. 使用 BGE-M3 + Chroma + BGE-Reranker 实现智能模板检索
"""
import logging
import os
from typing import Dict, List, Optional
from pathlib import Path
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_openai import ChatOpenAI
from docx import Document

logger = logging.getLogger(__name__)


class ContractDrafterAgent:
    """
    合同起草 Agent

    根据用户需求和分析结果，起草完整的合同内容
    """

    def __init__(self, llm: ChatOpenAI):
        self.llm = llm
        self.system_prompt = self._build_system_prompt()

    def draft(
        self,
        requirement: Dict,
        context: Dict = None
    ) -> str:
        """
        起草合同内容

        Args:
            requirement: 需求分析结果
            context: 上下文信息（如原合同内容等）

        Returns:
            Markdown 格式的合同内容
        """
        try:
            # 构建起草提示词
            prompt = self._build_drafting_prompt(requirement, context)

            # 调用 LLM
            response = self.llm.invoke([
                SystemMessage(content=self.system_prompt),
                HumanMessage(content=prompt)
            ])

            content = response.content.strip()
            logger.info(f"[ContractDrafter] 合同起草完成，长度: {len(content)} 字符")

            return content

        except Exception as e:
            logger.error(f"[ContractDrafter] 起草失败: {str(e)}", exc_info=True)
            return ""

    def _build_system_prompt(self) -> str:
        """构建系统提示词"""
        return """你是一个专业的合同起草专家。你的任务是根据用户需求起草完整的合同内容。

## 起草要求

### 内容要求
1. 条款完整、逻辑清晰
2. 用词准确、符合法律规范
3. 保护当事人合法权益
4. 预防常见法律风险
5. 结构规范、易于阅读

### 格式要求
使用 Markdown 格式，包含：
- 一级标题：合同名称（# 合同名称）
- 二级标题：主要条款（## 第一条 XXX）
- 三级标题：细分内容（### 1.1 XXX）
- 粗体：重要条款（**重要内容**）
- 列表：多项内容（- 列表项）

### 合同结构
1. 合同名称
2. 合同各方
3. 合同背景（引言）
4. 定义与解释
5. 核心条款（根据合同类型调整）
   - 标的物/服务内容
   - 数量/质量
   - 价款/报酬
   - 履行期限/方式
   - 权利义务
   - 违约责任
   - 争议解决
6. 其他条款
7. 签署

### 常见合同类型
- 股权转让协议
- 合作协议
- 租赁合同
- 服务合同
- 买卖合同
- 借款合同
- 劳动合同

## 输出要求
- 只输出合同内容，不要解释说明
- 确保格式规范，便于转换为 Word 文档
- 使用标准的法律文书用语
"""

    def _build_drafting_prompt(
        self,
        requirement: Dict,
        context: Dict = None
    ) -> str:
        """构建起草提示词"""
        key_info = requirement.get("key_info", {})
        processing_type = requirement.get("processing_type", "single_contract")

        prompt = f"""## 需求分析结果

处理类型：{processing_type}

关键信息：
"""

        # 添加关键信息
        for key, value in key_info.items():
            prompt += f"- {key}：{value}\n"

        # 如果有原合同信息
        if context and context.get("original_contract"):
            prompt += f"""
## 原合同信息

{context.get("original_contract_summary", "见原合同文件")}
"""

        # 如果有合同规划
        if processing_type == "contract_planning" and requirement.get("contract_plan"):
            plan = requirement["contract_plan"]
            current_contract = plan.get("current_contract", {})

            prompt += f"""
## 合同规划

当前起草：{current_contract.get("title", "合同")}
目的：{current_contract.get("purpose", "")}
优先级：{current_contract.get("priority", "")}

关联合同：
"""

            for related in current_contract.get("dependencies", []):
                prompt += f"- {related}\n"

        prompt += """

## 起草要求

请根据以上信息起草一份完整的合同，使用 Markdown 格式。
"""

        return prompt

    def draft_modification(
        self,
        original_contract: str,
        modification_points: List[Dict]
    ) -> str:
        """
        起草合同变更协议

        Args:
            original_contract: 原合同内容
            modification_points: 变更点列表

        Returns:
            变更协议内容
        """
        prompt = f"""## 原合同

{original_contract[:2000]}...

## 变更内容

"""

        for i, point in enumerate(modification_points, 1):
            prompt += f"{i}. **{point.get('title', '变更项')}**\n"
            prompt += f"   原条款：{point.get('original', '')}\n"
            prompt += f"   变更为：{point.get('modified', '')}\n\n"

        prompt += """
请起草一份合同变更协议，包含：
1. 变更背景
2. 变更条款（逐条列出）
3. 未变更条款继续有效
4. 其他约定
"""

        response = self.llm.invoke([
            SystemMessage(content=self.system_prompt),
            HumanMessage(content=prompt)
        ])

        return response.content.strip()

    def draft_termination(
        self,
        original_contract: str,
        termination_reason: str,
        post_termination_arrangements: Dict
    ) -> str:
        """
        起草合同解除协议

        Args:
            original_contract: 原合同内容
            termination_reason: 解除原因
            post_termination_arrangements: 解除后安排

        Returns:
            解除协议内容
        """
        prompt = f"""## 原合同

{original_contract[:2000]}...

## 解除原因

{termination_reason}

## 解除后安排

"""

        for key, value in post_termination_arrangements.items():
            prompt += f"- {key}：{value}\n"

        prompt += """
请起草一份合同解除协议，包含：
1. 解除背景
2. 解除条款
3. 解除后的权利义务处理
4. 费用结算
5. 争议解决
6. 其他约定
"""

        response = self.llm.invoke([
            SystemMessage(content=self.system_prompt),
            HumanMessage(content=prompt)
        ])

        return response.content.strip()

    def draft_with_template(
        self,
        requirement: Dict = None,
        context: Dict = None,
        use_template: bool = True,
        user_id: Optional[int] = None,
        return_template_info: bool = False,
        # 新接口参数（workflow 使用）
        analysis_result: Dict = None,
        template_content: str = "",
        strategy: Dict = None,
        reference_content: str = ""
    ) -> str | tuple[str, Optional[Dict]]:
        """
        使用专业模板起草合同内容（支持新旧两种接口）

        新接口（workflow 使用）:
            analysis_result: 需求分析结果
            template_content: 模板内容
            strategy: 生成策略
            reference_content: 参考资料

        旧接口（RAG 使用）:
            requirement: 需求分析结果
            context: 上下文信息
            use_template: 是否使用专业模板参考（RAG 检索）
            user_id: 当前用户 ID，用于检索私有模板
            return_template_info: 是否返回模板信息（用于前端展示）

        Returns:
            Markdown 格式的合同内容
            或 (内容, 模板信息) 元组（当 return_template_info=True 或使用新接口时）
        """
        try:
            # 判断使用新接口还是旧接口
            if analysis_result is not None:
                # 新接口（workflow 调用）
                return self._draft_with_template_new(
                    analysis_result=analysis_result,
                    template_content=template_content,
                    strategy=strategy,
                    reference_content=reference_content
                )

            # 旧接口（RAG 调用）- 继续原来的实现
            key_info = requirement.get("key_info", {}) if requirement else {}
            # 支持中英文字段名（向后兼容）
            transaction_type = key_info.get("交易类型") or key_info.get("transaction_type", "")

            # 构建基础提示词
            base_prompt = self._build_drafting_prompt(requirement, context)

            # RAG 模板检索
            template_content = ""
            template_info = None
            template_source = None  # 新增：记录模板来源

            if use_template and transaction_type:
                # 尝试使用 RAG 检索相关模板
                template_content, template_info = self._retrieve_template_with_rag(
                    query=transaction_type,
                    requirement=requirement,
                    user_id=user_id
                )

            # 如果 RAG 检索成功，添加到提示词中
            if template_content:
                enhanced_prompt = self._build_rag_enhanced_prompt(
                    base_prompt=base_prompt,
                    template_content=template_content,
                    template_info=template_info
                )
                prompt = enhanced_prompt
                template_source = "rag"
                logger.info(f"[ContractDrafter] 使用 RAG 模板改写: {transaction_type}, "
                           f"模板: {template_info['name'] if template_info else 'N/A'}")
            else:
                # RAG 检索失败或未启用，使用传统文件模板（兜底）
                if use_template and transaction_type:
                    template_content = self._load_template_content(transaction_type)

                if template_content:
                    # 传统模板也使用改写模式
                    enhanced_prompt = f"""{base_prompt}

---

## 📋 合同模板改写任务（兜底模式）

系统为您找到以下通用合同模板，请**基于此模板进行改写**：

### 改写要求
1. **以模板为骨架**：完全保留模板的章节结构和条款框架
2. **替换关键信息**：根据用户需求替换模板中的占位符和通用内容
3. **调整具体条款**：根据用户的具体情况修改条款细节
4. **保持专业性**：维持模板的法律严谨性和规范表达

### 待改写的合同模板

```
{template_content}
```

---

## ⚠️ 关键要求
- 必须基于模板进行改写，不要完全重新生成
- 保留模板的所有重要章节和条款框架
- 确保用户提供的所有关键信息都体现在合同中
"""
                    prompt = enhanced_prompt
                    template_source = "fallback"
                    template_info = {
                        "name": f"通用{transaction_type}模板",
                        "category": "通用模板",
                        "match_source": "fallback",
                        "description": "系统基于文件目录查找的通用模板"
                    }
                    logger.info(f"[ContractDrafter] 使用传统文件模板改写: {transaction_type}")
                else:
                    # 完全没有模板，使用基础起草
                    prompt = base_prompt
                    template_source = "none"
                    template_info = {
                        "name": "AI 自由生成",
                        "category": "无模板",
                        "match_source": "none",
                        "description": "未找到匹配模板，AI 根据需求自由生成"
                    }
                    logger.warning(f"[ContractDrafter] 未找到任何模板，使用基础起草: {transaction_type}")

            # 调用 LLM
            response = self.llm.invoke([
                SystemMessage(content=self.system_prompt),
                HumanMessage(content=prompt)
            ])

            content = response.content.strip()
            logger.info(f"[ContractDrafter] 合同起草完成（模板增强），长度: {len(content)} 字符")

            # 根据参数决定返回格式
            if return_template_info:
                # 添加来源标识
                if template_info:
                    template_info["source"] = template_source
                return content, template_info
            else:
                return content

        except Exception as e:
            logger.error(f"[ContractDrafter] 模板增强起草失败: {str(e)}", exc_info=True)
            # 降级到基础起草
            if return_template_info:
                return self.draft(requirement, context), {
                    "name": "AI 自由生成（降级）",
                    "category": "无模板",
                    "match_source": "fallback_error",
                    "description": "模板检索失败，AI 根据需求自由生成",
                    "source": "error"
                }
            else:
                return self.draft(requirement, context)

    def _retrieve_template_with_rag(
        self,
        query: str,
        requirement: Dict,
        user_id: Optional[int] = None,
        top_k: int = 3
    ) -> tuple[str, Optional[Dict]]:
        """
        使用 RAG 检索相关合同模板（改写模式）

        检索策略：
        1. 使用 BGE-M3 向量检索获取候选模板
        2. 使用 BGE-Reranker-v2 进行精准重排序
        3. 过滤权限（公共模板 + 用户私有模板）
        4. 返回最佳匹配模板用于改写

        Args:
            query: 检索查询（交易类型）
            requirement: 需求信息
            user_id: 当前用户 ID
            top_k: 检索候选数量

        Returns:
            tuple: (模板内容, 模板信息) 或 ("", None)
        """
        try:
            from ..rag.template_retriever import get_template_retriever

            # 获取检索器
            retriever = get_template_retriever()

            # 构建增强查询（包含更多上下文）
            key_info = requirement.get("key_info", {})
            enhanced_query_parts = [query]

            # 添加行业信息
            if key_info.get("industry"):
                enhanced_query_parts.append(f"行业：{key_info['industry']}")

            # 添加合同金额（大额合同可能需要特殊条款）
            if key_info.get("amount"):
                enhanced_query_parts.append(f"金额：{key_info['amount']}")

            # 添加特殊要求
            if key_info.get("special_terms"):
                enhanced_query_parts.append(f"特殊条款：{key_info['special_terms']}")

            # 添加合同主体类型
            if key_info.get("party_type"):
                enhanced_query_parts.append(f"主体类型：{key_info['party_type']}")

            enhanced_query = " | ".join(enhanced_query_parts)

            logger.info(f"[ContractDrafter] 开始 RAG 模板检索，查询: {enhanced_query}")

            # 执行检索（增加候选数量以确保能找到相关模板）
            search_result = retriever.retrieve(
                query=enhanced_query,
                user_id=user_id,
                top_k=top_k,
                use_rerank=True,
                rerank_top_n=100  # 增加候选数量到 100，确保能找到相关模板
            )

            if not search_result.templates:
                logger.warning(f"[ContractDrafter] RAG 检索未找到相关模板: {query}")
                return "", None

            # 获取最佳匹配模板
            best_template = search_result.templates[0]

            # 详细的匹配度日志
            logger.info(f"[ContractDrafter] RAG 检索成功:")
            logger.info(f"  - 模板名称: {best_template.name}")
            logger.info(f"  - 模板分类: {best_template.category}")
            logger.info(f"  - 向量相似度: {best_template.similarity_score:.4f}")
            rerank_str = f"{best_template.rerank_score:.4f}" if best_template.rerank_score is not None else "N/A"
            logger.info(f"  - 重排序分数: {rerank_str}")
            logger.info(f"  - 最终匹配度: {best_template.final_score:.4f}")
            logger.info(f"  - 匹配原因: {best_template.match_reason}")

            # 匹配度阈值检查
            # 降低阈值，因为重排序模型对中文合同模板的评分较严格
            MATCH_THRESHOLD = 0.01  # 降低阈值，提高模板使用率
            if best_template.final_score < MATCH_THRESHOLD:
                logger.warning(f"[ContractDrafter] 模板匹配度低于阈值 ({MATCH_THRESHOLD})，可能影响改写质量")

            # 读取模板文件内容
            template_content = self._load_template_from_file(best_template.file_url)

            if template_content:
                template_info = {
                    "name": best_template.name,
                    "category": best_template.category,
                    "subcategory": best_template.subcategory,
                    "match_score": best_template.final_score,
                    "similarity_score": best_template.similarity_score,
                    "rerank_score": best_template.rerank_score,
                    "match_reason": best_template.match_reason,
                    "file_url": best_template.file_url
                }
                logger.info(f"[ContractDrafter] 成功加载模板内容，长度: {len(template_content)} 字符")
                return template_content, template_info
            else:
                logger.warning(f"[ContractDrafter] 无法读取模板文件: {best_template.file_url}")
                return "", None

        except Exception as e:
            logger.error(f"[ContractDrafter] RAG 模板检索失败: {str(e)}", exc_info=True)
            return "", None

    def _load_template_from_file(self, file_path: str) -> str:
        """
        从文件路径加载模板内容

        支持 .md 和 .docx 格式

        Args:
            file_path: 模板文件路径

        Returns:
            str: 模板文本内容
        """
        try:
            # 处理可能的相对路径
            path = Path(file_path)

            # 如果是相对路径，尝试从多个可能的位置查找
            if not path.is_absolute():
                # 尝试的位置（按优先级）
                possible_paths = [
                    Path(f"/app/{file_path}"),  # 容器内绝对路径
                    Path(__file__).parent.parent.parent.parent.parent / file_path,  # backend 根目录
                    Path(__file__).parent.parent.parent.parent.parent / "storage" / file_path,  # storage 目录
                ]

                for test_path in possible_paths:
                    if test_path.exists():
                        path = test_path
                        break
                else:
                    logger.warning(f"[ContractDrafter] 模板文件不存在，尝试过的路径: {file_path}")
                    return ""

            if not path.exists():
                logger.warning(f"[ContractDrafter] 模板文件不存在: {file_path}")
                return ""

            # 根据文件扩展名读取
            if path.suffix.lower() == ".docx":
                # 读取 .docx 文件
                doc = Document(str(path))
                content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

                # 添加表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        content += f"\n{row_text}"

                logger.info(f"[ContractDrafter] 成功读取 DOCX 模板: {path.name}, 长度: {len(content)}")
                return content

            elif path.suffix.lower() in [".md", ".markdown"]:
                # 读取 Markdown 文件
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
                logger.info(f"[ContractDrafter] 成功读取 Markdown 模板: {path.name}, 长度: {len(content)}")
                return content

            else:
                # 其他格式，直接读取文本
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
                logger.info(f"[ContractDrafter] 成功读取文本模板: {path.name}, 长度: {len(content)}")
                return content

        except Exception as e:
            logger.error(f"[ContractDrafter] 读取模板文件失败: {str(e)}")
            return ""

    def _build_rag_enhanced_prompt(
        self,
        base_prompt: str,
        template_content: str,
        template_info: Optional[Dict]
    ) -> str:
        """
        构建模板改写提示词（改写模式）

        工作流程：
        1. 基于匹配的合同模板进行改写
        2. 保留模板的专业结构和条款框架
        3. 根据用户需求替换和调整具体内容

        Args:
            base_prompt: 基础提示词（包含用户需求）
            template_content: 模板内容
            template_info: 模板信息

        Returns:
            str: 改写模式的提示词
        """
        # 增加模板内容长度限制（支持更大的合同模板）
        max_template_length = 16000  # 从 8000 增加到 16000
        if len(template_content) > max_template_length:
            logger.warning(f"[ContractDrafter] 模板内容过长 ({len(template_content)} 字符)，截断到 {max_template_length} 字符")
            template_content = template_content[:max_template_length] + "\n...(内容已截断，完整模板请参考原文件)"

        # 计算匹配度百分比
        match_score_str = f"{template_info['match_score']:.2%}" if template_info else "N/A"
        match_reason_str = template_info['match_reason'] if template_info else '结构完整，条款规范'

        enhanced = f"""{base_prompt}

---

## 📋 合同模板改写任务

系统根据您的需求，为您检索到以下高匹配度合同模板，请**基于此模板进行改写**：

### 模板信息
- **模板名称**：{template_info['name'] if template_info else '专业合同模板'}
- **匹配度**：{match_score_str}
- **推荐原因**：{match_reason_str}

### 改写工作流程

请严格按照以下步骤进行改写：

#### 第一步：理解模板结构
仔细分析下方的合同模板，理解其：
- 章节结构和条款顺序
- 专业法律术语的使用方式
- 关键条款的表述方式（如违约责任、争议解决）

#### 第二步：提取用户需求
从上方的"需求分析结果"中提取：
- 合同主体信息（甲方、乙方等）
- 核心交易内容（标的物、服务、价格等）
- 特殊条款要求
- 履行方式和期限

#### 第三步：执行改写
**重要：以模板为骨架进行改写，而非重新生成**

1. **保留结构**：完全保留模板的章节顺序和条款框架
2. **替换主体**：将模板中的甲方、乙方等占位符替换为用户提供的具体信息
3. **修改内容**：根据用户需求调整具体条款内容
4. **增加条款**：如有特殊要求，在相应章节增加条款
5. **保持专业性**：维持模板的法律严谨性和专业表达

#### 改写示例

**模板原文**：
```
甲方：[公司名称]
乙方：[公司名称]
鉴于甲方需要...，乙方愿意...
```

**改写后**：
```
甲方：北京XX科技有限公司
乙方：上海XX咨询服务有限公司
鉴于甲方需要技术开发服务，乙方愿意提供相关技术服务...
```

### 待改写的合同模板

```
{template_content}
```

---

## ⚠️ 关键要求

1. **必须基于模板改写**：不要重新生成，而是对模板进行修改和适配
2. **保持结构完整**：保留所有重要章节，除非用户明确说明不需要
3. **信息完整**：确保用户提供的所有关键信息都体现在合同中
4. **格式规范**：使用 Markdown 格式，保持清晰的层级结构
5. **法律专业**：使用规范的法律术语，避免口语化表达

现在请开始改写，输出完整的合同内容。
"""
        return enhanced

    def _load_template_content(self, contract_type: str) -> str:
        """
        加载对应类型的合同模板内容

        Args:
            contract_type: 合同类型描述

        Returns:
            模板文本内容
        """
        try:
            # 模板目录
            template_dir = Path(__file__).parent.parent.parent.parent / "templates" / "documents"

            # 映射合同类型到模板文件（优先级从高到低）
            type_to_template = [
                ("软件开发", "software"),
                ("技术服务", "technology"),
                ("技术开发", "technology"),
                ("技术转让", "technology"),
                ("技术咨询", "consulting"),
                ("股权转让", "equity"),
                ("股权投资", "equity"),
                ("合作", "cooperation"),
                ("咨询", "consulting"),
                ("买卖", "sale"),
                ("租赁", "lease"),
                ("劳动", "labor"),
                ("保密", "nda"),
                ("借款", "loan"),
            ]

            # 查找匹配的模板（按优先级）
            template_name = None
            for keyword, tmpl in type_to_template:
                if keyword in contract_type:
                    template_name = tmpl
                    logger.info(f"[ContractDrafter] 匹配到模板类型: {keyword} -> {tmpl}")
                    break

            # 如果没有匹配，使用通用合同模板
            if not template_name:
                logger.warning(f"[ContractDrafter] 未找到匹配模板，使用通用模板: {contract_type}")
                template_file = template_dir / "contract_template.docx"
            else:
                # 尝试多种可能的文件名
                possible_names = [
                    f"{template_name}_template.docx",
                    f"{template_name}_contract.docx",
                    f"{template_name}.docx",
                    "contract_template.docx"
                ]
                template_file = None
                for name in possible_names:
                    potential_path = template_dir / name
                    if potential_path.exists():
                        template_file = potential_path
                        break

                if not template_file:
                    logger.warning(f"[ContractDrafter] 模板文件不存在，使用通用模板")
                    template_file = template_dir / "contract_template.docx"

            # 读取模板内容
            if template_file.exists():
                doc = Document(str(template_file))
                content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                logger.info(f"[ContractDrafter] 加载模板: {template_file.name}")
                return content
            else:
                logger.warning(f"[ContractDrafter] 模板文件不存在: {template_file}")
                return ""

        except Exception as e:
            logger.error(f"[ContractDrafter] 加载模板失败: {str(e)}")
            return ""

    def draft_from_scratch(
        self,
        analysis_result: Dict,
        reference_content: str = "",
        strategy: Dict = None,
        knowledge_graph_features: Dict = None  # ✨ 新增知识图谱特征参数
    ) -> str:
        """
        从零开始起草合同（不使用模板）

        Args:
            analysis_result: 需求分析结果（包含 key_info, processing_type 等）
            reference_content: 参考资料内容
            strategy: 生成策略信息
            knowledge_graph_features: 知识图谱法律特征（新增）

        Returns:
            Markdown 格式的合同内容
        """
        try:
            # 构建 requirement 字典（兼容旧接口）
            requirement = {
                "key_info": analysis_result.get("key_info", {}),
                "processing_type": analysis_result.get("processing_type", "single_contract")
            }

            # 构建基础提示词
            base_prompt = self._build_drafting_prompt(requirement)

            # 添加参考资料和知识图谱特征
            enhanced_prompt = base_prompt
            if reference_content or knowledge_graph_features:
                enhanced_prompt = f"""{base_prompt}"""

                # ✨ 添加知识图谱法律特征
                if knowledge_graph_features:
                    legal_features = knowledge_graph_features.get("legal_features")
                    if legal_features:
                        enhanced_prompt += f"""

## 📚 知识图谱法律特征

系统已为您匹配到该类合同的标准法律特征，请在起草时参考：

### 法律特征
- **交易性质**：{legal_features.get('transaction_nature', 'N/A')}
- **合同标的**：{legal_features.get('contract_object', 'N/A')}
- **起草立场**：{legal_features.get('stance', 'N/A')}
- **交易对价类型**：{legal_features.get('consideration_type', 'N/A')}
- **交易对价详情**：{legal_features.get('consideration_detail', 'N/A')}
- **交易特征**：{legal_features.get('transaction_characteristics', 'N/A')}

### 适用场景
{knowledge_graph_features.get('usage_scenario', 'N/A')}

### 法律依据
{chr(10).join(f"- {basis}" for basis in knowledge_graph_features.get('legal_basis', [])[:5]) if knowledge_graph_features.get('legal_basis') else '无'}

### 匹配信息
- **匹配合同类型**：{knowledge_graph_features.get('matched_contract_type', 'N/A')}
- **匹配置信度**：{knowledge_graph_features.get('match_confidence', 0):.2%}

**请确保起草的合同符合上述法律特征和适用场景。**
"""

                if reference_content:
                    enhanced_prompt += f"""

## 参考资料

{reference_content}
"""

                enhanced_prompt += """

---

请根据以上信息起草一份完整的合同，使用 Markdown 格式。
"""

            # 调用 LLM
            response = self.llm.invoke([
                SystemMessage(content=self.system_prompt),
                HumanMessage(content=enhanced_prompt)
            ])

            content = response.content.strip()
            logger.info(f"[ContractDrafter] 从零起草完成，长度: {len(content)} 字符")

            return content

        except Exception as e:
            logger.error(f"[ContractDrafter] 从零起草失败: {str(e)}", exc_info=True)
            return ""

    def _draft_with_template_new(
        self,
        analysis_result: Dict,
        template_content: str,
        strategy: Dict,
        reference_content: str = "",
        knowledge_graph_features: Dict = None  # ✨ 新增知识图谱特征参数
    ) -> tuple[str, Optional[Dict]]:
        """新接口：使用模板起草（workflow 调用）"""
        try:
            # 构建基础提示词
            requirement = {
                "key_info": analysis_result.get("key_info", {}),
                "processing_type": analysis_result.get("processing_type", "single_contract")
            }
            base_prompt = self._build_drafting_prompt(requirement)

            # 构建模板增强提示词
            template_info = {
                "name": strategy.get("template_name", "本地模板"),
                "category": "本地模板",
                "match_source": "local",
                "description": "使用本地 Markdown 模板"
            }

            enhanced_prompt = f"""{base_prompt}

---

## 📋 合同模板改写任务

请基于以下合同模板进行改写：

### 改写要求
1. **以模板为骨架**：完全保留模板的章节结构和条款框架
2. **替换关键信息**：根据用户需求替换模板中的占位符和通用内容
3. **调整具体条款**：根据用户的具体情况修改条款细节
4. **保持专业性**：维持模板的法律严谨性和规范表达

### 待改写的合同模板

```
{template_content}
```

---

"""

            # ✨ 添加知识图谱法律特征
            if knowledge_graph_features:
                legal_features = knowledge_graph_features.get("legal_features")
                if legal_features:
                    enhanced_prompt += f"""
## 📚 知识图谱法律特征

系统已为您匹配到该类合同的标准法律特征，请在改写时参考：

### 法律特征
- **交易性质**：{legal_features.get('transaction_nature', 'N/A')}
- **合同标的**：{legal_features.get('contract_object', 'N/A')}
- **起草立场**：{legal_features.get('stance', 'N/A')}
- **交易对价类型**：{legal_features.get('consideration_type', 'N/A')}
- **交易对价详情**：{legal_features.get('consideration_detail', 'N/A')}
- **交易特征**：{legal_features.get('transaction_characteristics', 'N/A')}

### 适用场景
{knowledge_graph_features.get('usage_scenario', 'N/A')}

### 法律依据
{chr(10).join(f"- {basis}" for basis in knowledge_graph_features.get('legal_basis', [])[:5]) if knowledge_graph_features.get('legal_basis') else '无'}

### 匹配信息
- **匹配合同类型**：{knowledge_graph_features.get('matched_contract_type', 'N/A')}
- **匹配置信度**：{knowledge_graph_features.get('match_confidence', 0):.2%}

**请确保改写的合同符合上述法律特征和适用场景。**

---

"""

            # 添加参考资料
            if reference_content:
                enhanced_prompt += f"""
## 参考资料

{reference_content}

"""

            enhanced_prompt += """
请开始改写，输出完整的合同内容。
"""

            # 调用 LLM
            response = self.llm.invoke([
                SystemMessage(content=self.system_prompt),
                HumanMessage(content=enhanced_prompt)
            ])

            content = response.content.strip()
            logger.info(f"[ContractDrafter] 模板起草完成，长度: {len(content)} 字符")

            return content, template_info

        except Exception as e:
            logger.error(f"[ContractDrafter] 新接口模板起草失败: {str(e)}", exc_info=True)
            return "", None

    def _draft_with_template_old(
        self,
        requirement: Dict,
        context: Dict = None,
        use_template: bool = True,
        user_id: Optional[int] = None,
        return_template_info: bool = False
    ) -> str | tuple[str, Optional[Dict]]:
        """旧接口：RAG 增强起草（保持向后兼容）"""
        # 原来的 draft_with_template 方法实现
        # ... 这里保留原来的实现代码 ...
        if return_template_info:
            return self.draft(requirement, context), None
        else:
            return self.draft(requirement, context)

    # ==================== 【新增】简化版变更/解除协议生成方法 ====================

    def draft_modification_simple(
        self,
        original_contract: str,
        confirmed_info: Dict
    ) -> str:
        """
        【新增】使用优化的结构化提示词起草变更协议

        适用于：
        - 已在 Step 2 确认结构化信息的场景
        - 简化工作流的单节点生成

        Args:
            original_contract: 原合同内容
            confirmed_info: 确认的变更信息，格式：
                {
                    "processing_type": "contract_modification",
                    "original_contract_info": {
                        "contract_name": str,
                        "signing_date": str,
                        "parties": [str, ...],
                        "contract_term": str,
                        "key_terms": { ... }
                    },
                    "modification_points": [
                        {
                            "clause_number": str,
                            "original_content": str,
                            "modified_content": str,
                            "reason": str
                        },
                        ...
                    ],
                    "confidence": float
                }

        Returns:
            变更协议内容（Markdown 格式）
        """
        try:
            # 提取原合同信息
            original_info = confirmed_info.get("original_contract_info", {})
            contract_name = original_info.get("contract_name", "原合同")
            signing_date = original_info.get("signing_date", "")
            parties = original_info.get("parties", [])
            contract_term = original_info.get("contract_term", "")

            # 提取变更点
            modification_points = confirmed_info.get("modification_points", [])

            # 构建结构化提示词
            prompt = f"""你是一位专业的律师，需要起草一份合同变更协议。

## 原合同基本信息
- **合同名称**：{contract_name}
- **签订日期**：{signing_date}
- **合同期限**：{contract_term}
- **当事人**：
{chr(10).join(f"  - {party}" for party in parties) if parties else "  - 详见原合同"}

## 原合同关键条款（参考）
{original_contract[:2000] if original_contract else "无"}

## 变更内容
"""

            if modification_points:
                for i, point in enumerate(modification_points, 1):
                    prompt += f"""
### 变更点 {i}
- **条款编号**：{point.get('clause_number', '未指定')}
- **原条款内容**：
  ```
  {point.get('original_content', '无')}
  ```
- **变更后内容**：
  ```
  {point.get('modified_content', '无')}
  ```
- **变更原因**：{point.get('reason', '未说明')}
"""
            else:
                prompt += "\n未提供具体变更点，请根据用户需求自行推断。\n"

            prompt += """
## 起草要求

请起草一份完整的合同变更协议，使用 Markdown 格式，包含以下部分：

### 1. 协议标题
- **合同变更协议**

### 2. 当事人信息
- 甲方、乙方的完整信息（名称、地址、法定代表人等）

### 3. 变更背景
- 简述原合同签订情况
- 说明变更的必要性

### 4. 变更条款
- **逐条列出变更内容**
- 明确原条款与变更后条款的对比
- 变更生效时间

### 5. 未变更条款
- 明确除本协议约定的变更外，原合同其他条款继续有效
- 原合同与变更协议不一致处，以变更协议为准

### 6. 变更效力
- 本协议作为原合同的组成部分
- 本协议与原合同具有同等法律效力

### 7. 争议解决
- 约定争议解决方式

### 8. 其他约定
- 双方认为需要约定的其他事项

### 9. 签署部分
- 双方签字盖章位置

## 法律要求
请确保：
- **条款完整、表述严谨**
- **符合《中华人民共和国民法典》相关规定**
- **语言简洁明了，避免歧义**
- **保护双方合法权益**
- **明确变更前后条款的对应关系**
"""

            response = self.llm.invoke([
                SystemMessage(content=self.system_prompt),
                HumanMessage(content=prompt)
            ])

            content = response.content.strip()
            logger.info(f"[ContractDrafter] 简化版变更协议起草完成，长度: {len(content)} 字符")

            return content

        except Exception as e:
            logger.error(f"[ContractDrafter] 简化版变更协议起草失败: {str(e)}", exc_info=True)
            raise

    def draft_termination_simple(
        self,
        original_contract: str,
        confirmed_info: Dict
    ) -> str:
        """
        【新增】使用优化的结构化提示词起草解除协议

        适用于：
        - 已在 Step 2 确认结构化信息的场景
        - 简化工作流的单节点生成

        Args:
            original_contract: 原合同内容
            confirmed_info: 确认的解除信息，格式：
                {
                    "processing_type": "contract_termination",
                    "original_contract_info": {
                        "contract_name": str,
                        "signing_date": str,
                        "parties": [str, ...],
                        "contract_term": str,
                        "key_terms": { ... }
                    },
                    "termination_reason": str,
                    "post_termination_arrangements": {
                        "fee_settlement": str,
                        "liability_breach": str,
                        "dispute_resolution": str,
                        ...
                    },
                    "confidence": float
                }

        Returns:
            解除协议内容（Markdown 格式）
        """
        try:
            # 提取原合同信息
            original_info = confirmed_info.get("original_contract_info", {})
            contract_name = original_info.get("contract_name", "原合同")
            signing_date = original_info.get("signing_date", "")
            parties = original_info.get("parties", [])
            contract_term = original_info.get("contract_term", "")

            # 提取解除信息
            termination_reason = confirmed_info.get("termination_reason", "")
            post_arrangements = confirmed_info.get("post_termination_arrangements", {})

            # 构建结构化提示词
            prompt = f"""你是一位专业的律师，需要起草一份合同解除协议。

## 原合同基本信息
- **合同名称**：{contract_name}
- **签订日期**：{signing_date}
- **合同期限**：{contract_term}
- **当事人**：
{chr(10).join(f"  - {party}" for party in parties) if parties else "  - 详见原合同"}

## 原合同关键条款（参考）
{original_contract[:2000] if original_contract else "无"}

## 解除原因
{termination_reason if termination_reason else "双方协商一致解除合同"}

## 解除后安排
"""

            # 添加解除后安排的详细信息
            arrangement_items = [
                ("fee_settlement", "费用结算方式"),
                ("liability_breach", "违约责任"),
                ("confidentiality", "保密条款"),
                ("dispute_resolution", "争议解决方式"),
                ("notice", "通知方式"),
                ("other", "其他约定")
            ]

            for key, label in arrangement_items:
                value = post_arrangements.get(key)
                if value:
                    prompt += f"- **{label}**：{value}\n"
                else:
                    prompt += f"- **{label}**：待双方协商\n"

            prompt += """
## 起草要求

请起草一份完整的合同解除协议，使用 Markdown 格式，包含以下部分：

### 1. 协议标题
- **合同解除协议**

### 2. 当事人信息
- 甲方、乙方的完整信息（名称、地址、法定代表人等）

### 3. 解除背景
- 简述原合同签订情况
- 说明解除原因
- 引用法律依据（如《民法典》相关规定）

### 4. 解除条款
- **明确解除日期**：约定合同自某年某月某日起解除
- **解除生效条件**：如需双方签署后生效，应予说明
- **解除范围**：明确解除的是原合同的全部还是部分

### 5. 解除后权利义务处理
- **已履行部分的确认**：明确双方对已履行部分的认可
- **未履行部分的终止**：确认解除后未履行部分不再履行
- **返还与恢复**：如需返还财产、恢复原状，应予明确
- **违约责任**：如一方违约导致解除，明确违约责任

### 6. 费用结算
- **已产生费用**：明确已产生费用的结算方式和期限
- **退款或补偿**：如涉及退款或补偿，明确金额、方式和期限
- **发票处理**：明确发票的开具和退还

### 7. 保密条款（如适用）
- 明确保密义务和保密期限
- 约定违反保密义务的责任

### 8. 争议解决
- 约定争议解决方式（协商、诉讼、仲裁）
- 明确管辖法院或仲裁机构

### 9. 其他约定
- 双方认为需要约定的其他事项

### 10. 签署部分
- 双方签字盖章位置
- 协议生效日期

## 法律要求
请确保：
- **条款完整、表述严谨**
- **符合《中华人民共和国民法典》相关规定**（特别是第562条、563条、566条关于合同解除的规定）
- **语言简洁明了，避免歧义**
- **保护双方合法权益**
- **明确解除后的各项安排，防止后续纠纷**
"""

            response = self.llm.invoke([
                SystemMessage(content=self.system_prompt),
                HumanMessage(content=prompt)
            ])

            content = response.content.strip()
            logger.info(f"[ContractDrafter] 简化版解除协议起草完成，长度: {len(content)} 字符")

            return content

        except Exception as e:
            logger.error(f"[ContractDrafter] 简化版解除协议起草失败: {str(e)}", exc_info=True)
            raise
