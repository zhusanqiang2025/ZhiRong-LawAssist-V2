# backend/app/services/document_templates.py
"""
Word 模板管理服务
支持基于模板生成文档
"""
import os
import logging
from pathlib import Path
from typing import Optional
from docx import Document

logger = logging.getLogger(__name__)


class TemplateManager:
    """
    模板管理器

    管理各类文档模板，提供基于模板的文档生成功能
    """

    def __init__(self):
        """初始化模板管理器"""
        # 模板根目录
        self.template_root = Path(__file__).parent.parent.parent / "templates" / "documents"
        self.template_root.mkdir(parents=True, exist_ok=True)

        # 可用的模板类型
        # 注意：python-docx 只支持 .docx 格式，不支持 .dotx 模板格式
        self.available_templates = {
            "contract": "contract_template.docx",    # 合同模板
            "letter": "letter_template.docx",        # 函件模板
            "judicial": "judicial_template.docx",    # 司法文书模板
            "simple": "simple_template.docx",        # 简单模板
        }

        logger.info(f"模板目录: {self.template_root}")
        self._scan_templates()

    def _scan_templates(self):
        """扫描可用的模板文件"""
        actual_templates = {}

        for doc_type, filename in self.available_templates.items():
            template_path = self.template_root / filename
            if template_path.exists():
                actual_templates[doc_type] = str(template_path)
                logger.info(f"找到模板: {doc_type} -> {filename}")
            else:
                logger.warning(f"模板不存在: {doc_type} -> {filename}")

        self.available_templates = actual_templates
        logger.info(f"可用模板数量: {len(actual_templates)}")

    def get_template_path(self, doc_type: str) -> Optional[str]:
        """
        获取指定类型文档的模板路径

        Args:
            doc_type: 文档类型 (contract/letter/judicial/simple)

        Returns:
            模板文件路径，不存在返回 None
        """
        return self.available_templates.get(doc_type)

    def load_template(self, doc_type: str) -> Optional[Document]:
        """
        加载指定类型的模板

        Args:
            doc_type: 文档类型

        Returns:
            Document 对象，失败返回 None
        """
        template_path = self.get_template_path(doc_type)

        if not template_path:
            logger.warning(f"模板不存在: {doc_type}")
            return None

        try:
            doc = Document(template_path)
            logger.info(f"成功加载模板: {doc_type} from {template_path}")
            return doc
        except Exception as e:
            logger.error(f"加载模板失败 {doc_type}: {str(e)}")
            return None

    def generate_from_template(
        self,
        doc_type: str,
        replacements: dict,
        output_path: str
    ) -> bool:
        """
        基于模板生成文档

        Args:
            doc_type: 文档类型
            replacements: 替换字典，格式：{"占位符": "替换值"}
            output_path: 输出文件路径

        Returns:
            是否成功
        """
        doc = self.load_template(doc_type)

        if not doc:
            return False

        try:
            # 替换文档中的占位符
            self._replace_placeholders(doc, replacements)

            # 保存文档
            doc.save(output_path)
            logger.info(f"文档生成成功: {output_path}")
            return True

        except Exception as e:
            logger.error(f"文档生成失败: {str(e)}")
            return False

    def _replace_placeholders(self, doc: Document, replacements: dict):
        """
        替换文档中的占位符

        支持的占位符格式：
        - {{variable_name}}
        - ${variable_name}
        - %variable_name%

        Args:
            doc: Document 对象
            replacements: 替换字典
        """
        import re

        logger.info(f"[DEBUG] 开始替换占位符，替换字典: {list(replacements.keys())}")

        # 定义占位符正则模式
        patterns = [
            (r'\{\{(\w+)\}\}', 1),  # {{variable}}
            (r'\$\{(\w+)\}', 1),    # ${variable}
            (r'%(\w+)%', 1),         # %variable%
        ]

        replaced_count = 0

        # 遍历所有段落
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                original_text = run.text
                for pattern, group in patterns:
                    matches = re.finditer(pattern, run.text)
                    for match in matches:
                        placeholder = match.group(group)
                        if placeholder in replacements:
                            run.text = re.sub(pattern, replacements[placeholder], run.text)
                            replaced_count += 1
                            logger.info(f"[DEBUG] 替换: {{{{placeholder}}}} -> '{replacements[placeholder]}'")
                if original_text != run.text:
                    logger.info(f"[DEBUG] Run 文本变化: {repr(original_text)} -> {repr(run.text)}")

        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for pattern, group in patterns:
                                matches = re.finditer(pattern, run.text)
                                for match in matches:
                                    placeholder = match.group(group)
                                    if placeholder in replacements:
                                        run.text = re.sub(pattern, replacements[placeholder], run.text)
                                        replaced_count += 1
                                        logger.info(f"[DEBUG] 表格替换: {{{{placeholder}}}} -> '{replacements[placeholder]}'")

        logger.info(f"[DEBUG] 占位符替换完成，共替换 {replaced_count} 处")

    def is_template_available(self, doc_type: str) -> bool:
        """检查指定类型的模板是否可用"""
        return doc_type in self.available_templates

    def get_available_template_types(self) -> list:
        """获取所有可用的模板类型"""
        return list(self.available_templates.keys())


# 单例
_template_manager: Optional[TemplateManager] = None


def get_template_manager() -> TemplateManager:
    """获取模板管理器单例"""
    global _template_manager
    if _template_manager is None:
        _template_manager = TemplateManager()
    return _template_manager
