import os
from io import BytesIO
from pdf2docx import Converter
from docx import Document
from docx.shared import RGBColor

class DocGenService:
    @staticmethod
    def generate_revised_docx(original_pdf_path: str, review_items: list):
        """
        生成带有模拟修订痕迹的 Word 文档
        :param original_pdf_path: 原始 PDF 路径
        :param review_items: 数据库中的 ContractReviewItem 列表
        """
        # 1. 临时转换 PDF -> Docx
        temp_docx = original_pdf_path.replace(".pdf", "_temp.docx")
        try:
            cv = Converter(original_pdf_path)
            cv.convert(temp_docx, start=0, end=None)
            cv.close()
            
            # 2. 打开生成的 Word
            doc = Document(temp_docx)
            
            # 3. 筛选出有效的修改 (已批准 或 已手动解决)
            valid_items = [
                i for i in review_items 
                if i.item_status in ["Approved", "Solved"]
            ]
            
            # 4. 应用修订样式
            for item in valid_items:
                target_text = item.quote
                # 优先使用人工定稿，其次用 AI 建议
                new_text = item.final_text if item.final_text else item.suggestion
                
                if not target_text or not new_text:
                    continue

                for para in doc.paragraphs:
                    if target_text in para.text:
                        parts = para.text.split(target_text)
                        para.clear()
                        
                        # 前文
                        para.add_run(parts[0])
                        
                        # 原文 (删除线 + 灰色)
                        run_old = para.add_run(target_text)
                        run_old.font.strike = True
                        run_old.font.color.rgb = RGBColor(169, 169, 169)
                        
                        # 新文 (下划线 + 红色)
                        run_new = para.add_run(f" {new_text} ")
                        run_new.font.underline = True
                        run_new.font.color.rgb = RGBColor(255, 0, 0)
                        run_new.bold = True
                        
                        # 后文
                        if len(parts) > 1:
                            para.add_run(parts[1])
                        break
            
            # 5. 保存到内存流
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # 清理临时文件
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
                
            return buffer

        except Exception as e:
            print(f"Doc generation failed: {e}")
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            return None