# backend/app/services/document_renderer.py
"""
专业文档渲染服务
将结构化文档渲染为规范格式的 Word 文档

支持两种渲染模式：
1. 基于模板渲染（推荐，使用 template.docx）
2. 代码生成渲染（兜底，纯代码创建）
"""
import logging
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.services.common.document_templates import get_template_manager

logger = logging.getLogger(__name__)


class DocumentRenderer:
    """
    文档渲染器

    将结构化的文档数据渲染为规范的 Word 文档
    优先使用模板，模板不存在时使用代码生成
    """

    def __init__(self):
        """初始化渲染器"""
        self.template_manager = get_template_manager()

    def render_contract(self, structure, docx_path: str) -> bool:
        """
        渲染合同文档

        Args:
            structure: DocumentStructure 对象
            docx_path: 输出文件路径

        Returns:
            是否成功
        """
        try:
            from app.services.document_structurer import DocumentStructure

            if not isinstance(structure, DocumentStructure):
                logger.error("结构参数类型错误")
                return False

            # 方案1：优先使用模板
            if self.template_manager.is_template_available("contract"):
                return self._render_from_template(
                    "contract",
                    structure,
                    docx_path
                )

            # 方案2：兜底使用代码生成
            doc = Document()

            # 设置文档样式
            self._setup_styles(doc)

            # 添加标题
            self._add_title(doc, structure.title)

            # 添加元数据（合同编号、日期等）
            self._add_metadata(doc, structure.metadata)

            # 添加合同方信息
            self._add_parties(doc, structure.parties)

            # 添加章节
            self._add_sections(doc, structure.sections)

            # 添加签署区
            self._add_signature_section(doc, structure.parties)

            # 保存文档
            doc.save(docx_path)

            logger.info(f"合同文档渲染成功: {docx_path}")
            return True

        except Exception as e:
            logger.error(f"合同文档渲染失败: {str(e)}")
            return False

    def render_letter(self, structure, docx_path: str) -> bool:
        """
        渲染函件文档

        Args:
            structure: DocumentStructure 对象
            docx_path: 输出文件路径

        Returns:
            是否成功
        """
        try:
            from app.services.document_structurer import DocumentStructure

            doc = Document()

            # 设置文档样式
            self._setup_letter_styles(doc)

            # 添加标题
            self._add_letter_title(doc, structure.title)

            # 添加称谓
            self._add_letter_salutation(doc)

            # 添加正文章节
            self._add_sections(doc, structure.sections)

            # 添加结语和签名
            self._add_letter_closing(doc)

            # 保存文档
            doc.save(docx_path)

            logger.info(f"函件文档渲染成功: {docx_path}")
            return True

        except Exception as e:
            logger.error(f"函件文档渲染失败: {str(e)}")
            return False

    def _setup_styles(self, doc: Document):
        """设置合同文档样式"""
        # 正文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)

        # 标题1样式
        if 'Heading 1' in doc.styles:
            h1 = doc.styles['Heading 1']
            h1.font.name = '黑体'
            h1.font.size = Pt(16)
            h1.font.bold = True
            h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1.paragraph_format.space_before = Pt(12)
            h1.paragraph_format.space_after = Pt(12)

        # 标题2样式
        if 'Heading 2' in doc.styles:
            h2 = doc.styles['Heading 2']
            h2.font.name = '黑体'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.paragraph_format.space_before = Pt(6)
            h2.paragraph_format.space_after = Pt(6)

        # 标题3样式
        if 'Heading 3' in doc.styles:
            h3 = doc.styles['Heading 3']
            h3.font.name = '黑体'
            h3.font.size = Pt(12)
            h3.font.bold = True
            h3.paragraph_format.space_before = Pt(3)
            h3.paragraph_format.space_after = Pt(3)

    def _setup_letter_styles(self, doc: Document):
        """设置函件文档样式"""
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)

        # 函件标题
        try:
            letter_title = doc.styles.add_style('LetterTitle', 1)  # 1=paragraph style
            letter_title.font.name = '黑体'
            letter_title.font.size = Pt(18)
            letter_title.font.bold = True
            letter_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass  # 样式可能已存在

    def _add_title(self, doc: Document, title: str):
        """添加文档标题"""
        # 添加空行
        doc.add_paragraph()

        # 添加标题
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置标题格式
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(18)
            run.font.bold = True

        # 添加分隔线
        doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_letter_title(self, doc: Document, title: str):
        """添加函件标题"""
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(18)
            run.font.bold = True

    def _add_metadata(self, doc: Document, metadata: dict):
        """添加元数据"""
        if not metadata:
            return

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts = []
        if metadata.get('contract_number'):
            parts.append(f"合同编号：{metadata['contract_number']}")
        if metadata.get('date'):
            parts.append(f"日期：{metadata['date']}")

        if parts:
            p.add_run('  |  '.join(parts))
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

    def _add_parties(self, doc: Document, parties):
        """添加合同方信息"""
        if not parties:
            return

        # 合同方标题
        doc.add_heading("合同各方", level=2)

        for party in parties:
            p = doc.add_paragraph(style='List Bullet')

            # 添加角色和名称
            role_run = p.add_run(f"{party.role}：")
            role_run.bold = True
            role_run.font.name = '黑体'

            p.add_run(party.name or "_________________")

            # 添加详细字段
            if party.fields:
                for key, value in party.fields.items():
                    detail = doc.add_paragraph()
                    detail.paragraph_format.left_indent = Inches(0.5)
                    label_run = detail.add_run(f"{key}：")
                    label_run.bold = True
                    detail.add_run(value or "_________________")

        doc.add_paragraph()

    def _add_sections(self, doc: Document, sections, parent_level: int = 0):
        """递归添加章节"""
        for section in sections:
            # 确定标题级别
            heading_level = min(section.level, 3)  # Word 最多支持 9 级，但一般只用 3 级

            # 添加标题
            if section.title:
                heading = doc.add_heading(section.title, level=heading_level)

                # 根据级别设置不同样式
                if heading_level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif heading_level == 2:
                    # 编号格式
                    heading_run = heading.runs[0] if heading.runs else heading.add_run("")
                    heading_run.font.bold = True

            # 添加内容
            if section.content:
                # 分段处理
                paragraphs = section.content.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        p.paragraph_format.first_line_indent = Inches(0.3)  # 首行缩进

            # 递归添加子章节
            if section.subsections:
                self._add_sections(doc, section.subsections, section.level)

    def _add_signature_section(self, doc: Document, parties):
        """添加签署区"""
        # 添加空行
        for _ in range(3):
            doc.add_paragraph()

        # 签署标题
        doc.add_heading("签署", level=2)

        # 计算每行显示的签署方数量
        party_count = len(parties) if parties else 2
        parties_per_row = min(party_count, 2)

        # 创建签署表格
        table = doc.add_table(rows=1, cols=parties_per_row * 2)
        table.autofit = False

        # 设置表格宽度
        table.width = Inches(6)

        for i, party in enumerate(parties or [type('Party', (), {'role': '甲方', 'name': ''})(),
                                                type('Party', (), {'role': '乙方', 'name': ''})()]):
            col = i % parties_per_row
            row = i // parties_per_row

            if row >= len(table.rows):
                table.add_row()

            cell = table.rows[row].cells[col * 2]

            # 添加角色
            p = cell.paragraphs[0]
            role_run = p.add_run(f"{party.role}（盖章）：")
            role_run.bold = True

            # 添加签名线
            cell = table.rows[row].cells[col * 2 + 1]
            p = cell.paragraphs[0]
            p.add_run("___________________")

            # 添加代表签字和日期
            if row >= len(table.rows):
                table.add_row()
            cell = table.rows[row].cells[col * 2]
            p = cell.paragraphs[0]
            p.add_run("代表签字：")

            cell = table.rows[row].cells[col * 2 + 1]
            p = cell.paragraphs[0]
            p.add_run("___________________")

            if row + 1 >= len(table.rows):
                table.add_row()
            cell = table.rows[row + 1].cells[col * 2]
            p = cell.paragraphs[0]
            p.add_run("日期：")

            cell = table.rows[row + 1].cells[col * 2 + 1]
            p = cell.paragraphs[0]
            p.add_run("____年__月__日")

    def _add_letter_salutation(self, doc: Document):
        """添加函件称谓"""
        p = doc.add_paragraph()
        p.add_run("尊敬的先生/女士：")
        p.paragraph_format.space_after = Pt(12)

    def _add_letter_closing(self, doc: Document):
        """添加函件结语"""
        # 添加空行
        for _ in range(2):
            doc.add_paragraph()

        # 结语
        p = doc.add_paragraph()
        p.add_run("此致")
        p.paragraph_format.space_after = Pt(6)

        p = doc.add_paragraph()
        p.add_run("敬礼！")
        p.paragraph_format.space_after = Pt(24)

        # 签名
        p = doc.add_paragraph()
        p.add_run("___________________")
        p.paragraph_format.space_after = Pt(6)

        p = doc.add_paragraph()
        signature_run = p.add_run("签名：")
        signature_run.bold = True
        p.add_run("___________________")
        p.paragraph_format.space_after = Pt(6)

        # 日期
        p = doc.add_paragraph()
        date_run = p.add_run("日期：")
        date_run.bold = True
        p.add_run("____年__月__日")

    def _render_from_template(
        self,
        template_type: str,
        structure,
        output_path: str
    ) -> bool:
        """
        基于模板渲染文档

        方案：将文本内容填充到模板中，保留模板的样式

        Args:
            template_type: 模板类型 (contract/letter/judicial)
            structure: 文档结构对象
            output_path: 输出路径

        Returns:
            是否成功
        """
        try:
            # 加载模板
            template = self.template_manager.load_template(template_type)
            if not template:
                logger.warning(f"模板加载失败: {template_type}")
                return False

            # 获取章节内容文本
            sections_text = self._sections_to_text(structure.sections)

            logger.info(f"[DEBUG] 准备填充 {len(sections_text)} 字符的内容到模板")

            # 策略：直接在文档末尾添加内容，不管是否有"合同条款"等关键词
            # 这样可以兼容用户自定义的只有样式的模板
            logger.info("[DEBUG] 将内容添加到模板末尾")
            self._append_content_to_document(template, sections_text)

            # 处理其他占位符（如果有的话）
            replacements = self._build_replacements(structure)
            replacements.pop("sections_content", "")  # 移除 sections_content，已单独处理

            # 只有当替换字典中有非空值时才进行替换
            has_valid_replacements = any(v for v in replacements.values() if v)
            if has_valid_replacements:
                self.template_manager._replace_placeholders(template, replacements)

            # 保存文档
            template.save(output_path)
            logger.info(f"基于模板渲染成功: {output_path}")
            return True

        except Exception as e:
            logger.error(f"模板渲染失败: {str(e)}")
            import traceback
            traceback.print_exc()
            return False

    def _find_content_insert_position(self, doc: Document):
        """
        查找内容插入位置

        查找策略：
        1. 查找包含"合同条款"、"内容"、"正文"等关键词的标题
        2. 返回该标题之后的段落索引

        Args:
            doc: Document 对象

        Returns:
            插入位置的段落索引，未找到返回 None
        """
        content_keywords = ["合同条款", "合同内容", "正文", "内容", "协议条款", "条款"]

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            for keyword in content_keywords:
                if keyword in text:
                    logger.info(f"[DEBUG] 找到内容插入位置: '{text}' (段落 {i})")
                    return i

        return None

    def _insert_content_at_position(self, doc: Document, position: int, content: str):
        """
        在指定位置插入内容

        Args:
            doc: Document 对象
            position: 插入位置的段落索引
            content: 要插入的内容
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # 将内容按行分割
        paragraphs_list = content.split('\n')
        paragraphs_list = [p.strip() for p in paragraphs_list if p.strip()]

        logger.info(f"[DEBUG] 在位置 {position} 插入 {len(paragraphs_list)} 个段落")

        # 获取插入位置段落的样式作为参考
        ref_paragraph = doc.paragraphs[position] if position < len(doc.paragraphs) else doc.paragraphs[-1]

        # 找到插入位置
        insert_after_element = doc.paragraphs[position]._p if position < len(doc.paragraphs) else None

        for i, para_text in enumerate(paragraphs_list):
            # 创建新段落 - 使用正确的方法
            new_p = OxmlElement('w:p')

            # 添加段落属性
            pPr = OxmlElement('w:pPr')
            new_p.append(pPr)

            # 尝试使用参考段落的样式
            try:
                if ref_paragraph.style:
                    style_id = ref_paragraph.style.style_id
                    style = OxmlElement('w:pStyle')
                    style.set(qn('w:val'), style_id)
                    pPr.append(style)
            except:
                pass  # 如果获取样式失败，使用默认样式

            # 添加运行
            r = OxmlElement('w:r')
            new_p.append(r)

            # 添加运行属性（继承参考段落的字体设置）
            rPr = OxmlElement('w:rPr')
            r.append(rPr)

            # 尝试继承字体设置
            try:
                if ref_paragraph.runs:
                    ref_run = ref_paragraph.runs[0]
                    if ref_run.font.name:
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), ref_run.font.name)
                        rPr.append(rFonts)
                    if ref_run.font.size:
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), str(int(ref_run.font.size * 2)))  # Word 字体大小单位是半点
                        rPr.append(sz)
            except:
                pass

            # 添加文本
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = para_text
            r.append(t)

            # 插入段落
            if insert_after_element is not None:
                insert_after_element.addnext(new_p)
            else:
                doc._element.body.append(new_p)

    def _append_content_to_document(self, doc: Document, content: str):
        """
        在文档末尾追加内容

        根据内容类型智能应用样式：
        - 标题行（如"第一条"、"一、"等）应用标题样式
        - 正文段落应用正文样式
        - 合同方信息（如"甲方："）应用主体信息样式

        Args:
            doc: Document 对象
            content: 要追加的内容
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import re

        # 将内容按行分割
        paragraphs_list = content.split('\n')
        paragraphs_list = [p.strip() for p in paragraphs_list if p.strip()]

        # 标题识别模式
        heading_patterns = [
            r'^第[一二三四五六七八九十百]+[条款章]',  # 第一条、第二章
            r'^[一二三四五六七八九十百]+[、．.]',      # 一、、二．
            r'^\d+[、．.]',                             # 1、、2.
            r'^\d+\.\d+',                              # 1.1、2.3
            r'^[（(]\d+[)）]',                         # （1）、(2)
        ]

        for para_text in paragraphs_list:
            # 判断是否是标题
            is_heading = any(re.match(pattern, para_text) for pattern in heading_patterns)

            # 判断是否是合同方信息行
            is_party_info = '：' in para_text or ':' in para_text

            # 创建新段落
            new_p = OxmlElement('w:p')

            # 添加段落属性
            pPr = OxmlElement('w:pPr')
            new_p.append(pPr)

            # 根据内容类型设置样式
            if is_heading:
                # 标题样式：使用 Heading 2
                style = OxmlElement('w:pStyle')
                style.set(qn('w:val'), 'Heading2')
                pPr.append(style)

                # 添加间距
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '120')  # 10pt before
                spacing.set(qn('w:after'), '80')    # 6-7pt after
                pPr.append(spacing)
            elif is_party_info:
                # 合同方信息：保持简洁，较小间距
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '40')
                spacing.set(qn('w:after'), '40')
                pPr.append(spacing)
            else:
                # 正文段落：添加首行缩进
                ind = OxmlElement('w:ind')
                ind.set(qn('w:firstLineChars'), '200')  # 首行缩进2字符
                pPr.append(ind)

                # 添加间距
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '40')
                spacing.set(qn('w:after'), '40')
                pPr.append(spacing)

            # 添加运行
            r = OxmlElement('w:r')
            new_p.append(r)

            # 添加运行属性
            rPr = OxmlElement('w:rPr')
            r.append(rPr)

            # 设置字体
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), '宋体')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:hAnsi'), '宋体')
            rPr.append(rFonts)

            # 设置字号
            if is_heading:
                # 标题使用较大字号（四号约14pt）
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '28')  # 14pt * 2
                rPr.append(sz)

                # 标题加粗
                b = OxmlElement('w:b')
                rPr.append(b)
            else:
                # 正文使用小四号约12pt
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '24')  # 12pt * 2
                rPr.append(sz)

            # 如果是合同方信息，标签部分加粗
            if is_party_info:
                parts = re.split(r'[：:]', para_text, 1)
                if len(parts) == 2:
                    # 标签部分（加粗）
                    r1 = OxmlElement('w:r')
                    new_p.append(r1)
                    rPr1 = OxmlElement('w:rPr')
                    r1.append(rPr1)

                    # 字体
                    rFonts1 = OxmlElement('w:rFonts')
                    rFonts1.set(qn('w:ascii'), '宋体')
                    rFonts1.set(qn('w:eastAsia'), '宋体')
                    rFonts1.set(qn('w:hAnsi'), '宋体')
                    rPr1.append(rFonts1)

                    # 字号
                    sz1 = OxmlElement('w:sz')
                    sz1.set(qn('w:val'), '24')
                    rPr1.append(sz1)

                    # 加粗
                    b1 = OxmlElement('w:b')
                    rPr1.append(b1)

                    # 文本
                    t1 = OxmlElement('w:t')
                    t1.set(qn('xml:space'), 'preserve')
                    t1.text = parts[0] + '：'
                    r1.append(t1)

                    # 值部分（正常）
                    r2 = OxmlElement('w:r')
                    new_p.append(r2)
                    rPr2 = OxmlElement('w:rPr')
                    r2.append(rPr2)

                    # 字体
                    rFonts2 = OxmlElement('w:rFonts')
                    rFonts2.set(qn('w:ascii'), '宋体')
                    rFonts2.set(qn('w:eastAsia'), '宋体')
                    rFonts2.set(qn('w:hAnsi'), '宋体')
                    rPr2.append(rFonts2)

                    # 字号
                    sz2 = OxmlElement('w:sz')
                    sz2.set(qn('w:val'), '24')
                    rPr2.append(sz2)

                    # 文本
                    t2 = OxmlElement('w:t')
                    t2.set(qn('xml:space'), 'preserve')
                    t2.text = parts[1]
                    r2.append(t2)
                else:
                    # 无法分割，按普通文本处理
                    t = OxmlElement('w:t')
                    t.set(qn('xml:space'), 'preserve')
                    t.text = para_text
                    r.append(t)
            else:
                # 添加文本
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = para_text
                r.append(t)

            # 追加段落
            doc._element.body.append(new_p)

    def _build_replacements(self, structure) -> dict:
        """
        根据文档结构构建替换字典

        Args:
            structure: DocumentStructure 对象

        Returns:
            替换字典
        """
        replacements = {}

        # 基本信息
        replacements["title"] = structure.title or "文档标题"
        replacements["doc_type"] = structure.doc_type or "合同"

        # 调试日志
        logger.info(f"[DEBUG] structure.title = {structure.title}")
        logger.info(f"[DEBUG] structure.doc_type = {structure.doc_type}")

        # 元数据
        if structure.metadata:
            logger.info(f"[DEBUG] structure.metadata = {structure.metadata}")
            replacements["contract_number"] = structure.metadata.get("contract_number", "")
            replacements["date"] = structure.metadata.get("date", "")
            replacements["keywords"] = ", ".join(structure.metadata.get("keywords", []))
            logger.info(f"[DEBUG] contract_number = '{replacements['contract_number']}'")
            logger.info(f"[DEBUG] date = '{replacements['date']}'")
        else:
            logger.warning("[DEBUG] structure.metadata is None or empty")
            replacements["contract_number"] = ""
            replacements["date"] = ""
            replacements["keywords"] = ""

        # 合同方信息
        if structure.parties:
            for i, party in enumerate(structure.parties):
                replacements[f"party{i+1}_role"] = party.role
                replacements[f"party{i+1}_name"] = party.name or "_________________"

                # 详细字段
                for key, value in party.fields.items():
                    placeholder = f"party{i+1}_{key}"
                    replacements[placeholder] = value or "_________________"

        # 章节内容（拼接成文本）
        sections_text = self._sections_to_text(structure.sections)
        replacements["sections_content"] = sections_text

        return replacements

    def _sections_to_text(self, sections, level: int = 0) -> str:
        """
        将章节列表转换为文本

        Args:
            sections: 章节列表
            level: 当前层级

        Returns:
            文本内容
        """
        lines = []

        for section in sections:
            # 添加标题
            if section.title:
                prefix = "  " * level
                lines.append(f"{prefix}{section.title}")

            # 添加内容
            if section.content:
                lines.append(section.content)

            # 递归添加子章节
            if section.subsections:
                lines.append(self._sections_to_text(section.subsections, level + 1))

        return "\n".join(lines)


_renderer_instance: Optional[DocumentRenderer] = None


def get_renderer() -> DocumentRenderer:
    """获取文档渲染器单例"""
    global _renderer_instance
    if _renderer_instance is None:
        _renderer_instance = DocumentRenderer()
    return _renderer_instance
