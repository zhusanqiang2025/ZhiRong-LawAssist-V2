# backend/app/services/docx_editor.py
"""
Word 文档编辑服务（增强版）
使用 Word 原生修订模式（Track Changes）应用 AI 审查意见
"""
import os
import re
from typing import List, Tuple, Optional
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging

logger = logging.getLogger(__name__)


class DocxEditor:
    """Word 文档编辑器，使用原生修订模式"""

    def __init__(self, docx_path: str):
        """
        初始化编辑器

        Args:
            docx_path: Word 文档路径
        """
        self.doc = Document(docx_path)
        self.docx_path = docx_path
        self.applied_count = 0
        self.not_found_count = 0

    def _create_track_revision_element(self, author: str = "AI审查", date: str = None):
        """
        创建修订记录元素（用于启用修订模式）

        Args:
            author: 修订作者
            date: 修订日期（ISO格式）

        Returns:
            w:rPr 元素包含修订信息
        """
        from datetime import datetime

        if date is None:
            date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

        # 创建修订属性元素
        rpr = OxmlElement('w:rPr')
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), '0')  # 修订ID
        ins.set(qn('w:author'), author)
        ins.set(qn('w:date'), date)
        rpr.append(ins)

        return rpr

    def _create_deletion_element(self, author: str = "AI审查", date: str = None):
        """
        创建删除修订元素

        Args:
            author: 修订作者
            date: 修订日期

        Returns:
            w:del 元素
        """
        from datetime import datetime

        if date is None:
            date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

        del_element = OxmlElement('w:del')
        del_element.set(qn('w:id'), '1')
        del_element.set(qn('w:author'), author)
        del_element.set(qn('w:date'), date)

        return del_element

    def _find_text_position_in_paragraph(self, para, search_text: str) -> Optional[int]:
        """
        在段落中查找文本位置

        Args:
            para: python-docx 段落对象
            search_text: 要查找的文本

        Returns:
            文本在段落中的起始位置，未找到返回 None
        """
        full_text = para.text
        if search_text not in full_text:
            return None

        return full_text.index(search_text)

    def apply_revision_in_paragraph(self, para, old_text: str, new_text: str) -> bool:
        """
        在段落中应用修订（使用原生修订模式）

        Args:
            para: python-docx 段落对象
            old_text: 要删除的原文
            new_text: 要插入的新文

        Returns:
            是否成功应用
        """
        full_text = para.text

        # 清理文本：移除多余空白
        old_text_clean = old_text.strip()
        new_text_clean = new_text.strip()

        logger.info(f"[DocxEditor] 段落文本长度: {len(full_text)}, 查找: {old_text_clean[:30]}...")

        if old_text_clean not in full_text:
            logger.warning(f"[DocxEditor] 原文不在段落中")
            return False

        # 找到第一个匹配的 run
        target_run = None
        target_index = -1
        run_texts = []

        for i, run in enumerate(para.runs):
            run_texts.append(f"Run{i}: '{run.text}'")
            if old_text_clean in run.text:
                target_run = run
                target_index = i
                logger.info(f"[DocxEditor] 在 Run{i} 中找到匹配")
                break

        if target_run is None:
            logger.warning(f"[DocxEditor] 未在任何单个 run 中找到完整匹配")
            logger.debug(f"[DocxEditor] 段落 runs: {run_texts}")
            return False

        # 获取目标 run 的原始属性
        original_text = target_run.text
        start_pos = original_text.index(old_text_clean)

        # 分割文本
        before_text = original_text[:start_pos]
        after_text = original_text[start_pos + len(old_text_clean):]

        logger.info(f"[DocxEditor] 文本分割 - 前缀: '{before_text[:20]}', 后缀: '{after_text[:20]}'")

        # 保存字体属性
        font_size = target_run.font.size
        font_name = target_run.font.name
        font_bold = target_run.font.bold
        font_italic = target_run.font.italic

        # 前缀
        if before_text:
            target_run.text = before_text
            logger.debug(f"[DocxEditor] 设置前缀文本: '{before_text[:30]}'")
        else:
            target_run.text = ""
            logger.debug(f"[DocxEditor] 清空 run 文本")

        # 删除的原文（红色删除线）
        deleted_run = para.add_run(old_text_clean)
        deleted_run.font.strike = True
        deleted_run.font.color.rgb = RGBColor(217, 0, 0)  # 深红色
        if font_size:
            deleted_run.font.size = font_size
        if font_name:
            deleted_run.font.name = font_name
        logger.info(f"[DocxEditor] 添加删除文本(红色删除线): '{old_text_clean[:30]}'")

        # 插入的新文（黄色高亮+下划线）
        inserted_run = para.add_run(new_text_clean)
        inserted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        inserted_run.font.underline = True
        inserted_run.font.color.rgb = RGBColor(0, 0, 0)
        if font_size:
            inserted_run.font.size = font_size
        if font_name:
            inserted_run.font.name = font_name
        logger.info(f"[DocxEditor] 添加插入文本(黄色高亮+下划线): '{new_text_clean[:30]}'")

        # 后缀
        if after_text:
            suffix_run = para.add_run(after_text)
            if font_size:
                suffix_run.font.size = font_size
            if font_name:
                suffix_run.font.name = font_name
            logger.debug(f"[DocxEditor] 设置后缀文本: '{after_text[:30]}'")

        logger.info(f"[DocxEditor] 修订应用成功")
        return True

    def apply_revision(self, quote: str, suggestion: str) -> bool:
        """
        应用单条修订建议

        Args:
            quote: 原文
            suggestion: 建议文本

        Returns:
            是否成功应用
        """
        quote = quote.strip()
        suggestion = suggestion.strip()

        if not quote or not suggestion:
            return False

        # 1. 在段落中查找
        for para in self.doc.paragraphs:
            if self.apply_revision_in_paragraph(para, quote, suggestion):
                self.applied_count += 1
                logger.info(f"✓ 已在段落中应用修订: {quote[:30]}... → {suggestion[:30]}...")
                return True

        # 2. 在表格中查找
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if self.apply_revision_in_paragraph(para, quote, suggestion):
                            self.applied_count += 1
                            logger.info(f"✓ 已在表格中应用修订: {quote[:30]}... → {suggestion[:30]}...")
                            return True

        # 3. 在页眉页脚中查找
        for section in self.doc.sections:
            for para in section.header.paragraphs:
                if self.apply_revision_in_paragraph(para, quote, suggestion):
                    self.applied_count += 1
                    logger.info(f"✓ 已在页眉中应用修订: {quote[:30]}... → {suggestion[:30]}...")
                    return True
            for para in section.footer.paragraphs:
                if self.apply_revision_in_paragraph(para, quote, suggestion):
                    self.applied_count += 1
                    logger.info(f"✓ 已在页脚中应用修订: {quote[:30]}... → {suggestion[:30]}...")
                    return True

        self.not_found_count += 1
        logger.warning(f"✗ 未找到原文: {quote[:50]}...")
        return False

    def apply_revisions(self, revisions: List[dict]) -> dict:
        """
        批量应用修订建议

        Args:
            revisions: 修订列表，每个元素包含 quote 和 suggestion

        Returns:
            应用结果统计
        """
        seen = set()  # 避免重复应用
        results = {
            "total": len(revisions),
            "applied": 0,
            "not_found": 0,
            "details": []
        }

        for rev in revisions:
            quote = rev.get("quote", "").strip()
            suggestion = rev.get("suggestion", "").strip()

            if not quote or not suggestion:
                continue

            # 去重
            key = f"{quote}|{suggestion}"
            if key in seen:
                continue
            seen.add(key)

            success = self.apply_revision(quote, suggestion)
            results["details"].append({
                "quote": quote,
                "suggestion": suggestion,
                "success": success
            })

        results["applied"] = self.applied_count
        results["not_found"] = self.not_found_count
        return results

    def save(self, output_path: str):
        """
        保存文档

        Args:
            output_path: 输出文件路径
        """
        self.doc.save(output_path)
        logger.info(f"修订版文档已保存至: {output_path}")

    @staticmethod
    def convert_pdf_to_docx(pdf_path: str) -> Tuple[bool, str, str]:
        """
        将 PDF 转换为 DOCX

        Args:
            pdf_path: PDF 文件路径

        Returns:
            (success, docx_path, message)
        """
        import subprocess

        base_name = pdf_path.rsplit('.', 1)[0]
        docx_path = f"{base_name}.converted.docx"

        # 使用 LibreOffice 转换
        try:
            cmd = [
                "soffice", "--headless", "--convert-to", "docx",
                "--outdir", os.path.dirname(pdf_path),
                pdf_path
            ]
            result = subprocess.run(cmd, check=True, timeout=60, capture_output=True)

            # 检查输出文件
            if os.path.exists(docx_path):
                return True, docx_path, "PDF 已转换为 DOCX"
            else:
                # LibreOffice 可能生成不同名称
                possible = [f for f in os.listdir(os.path.dirname(pdf_path))
                           if f.startswith(os.path.basename(base_name)) and f.endswith('.docx')]
                if possible:
                    return True, os.path.join(os.path.dirname(pdf_path), possible[0]), "PDF 已转换为 DOCX"
                return False, "", "转换失败：未找到输出文件"

        except subprocess.TimeoutExpired:
            return False, "", "PDF 转换超时"
        except subprocess.CalledProcessError as e:
            return False, "", f"PDF 转换失败: {e.stderr.decode() if e.stderr else str(e)}"
        except Exception as e:
            return False, "", f"PDF 转换出错: {str(e)}"
