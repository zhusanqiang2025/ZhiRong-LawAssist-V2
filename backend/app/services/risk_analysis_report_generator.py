# backend/app/services/risk_analysis_report_generator.py
"""
风险分析报告生成服务

将风险分析结果生成规范的 Word 报告
"""
import logging
import os
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class RiskAnalysisReportGenerator:
    """风险分析报告生成器"""

    def __init__(self):
        """初始化报告生成器"""
        self.output_dir = "storage/reports/risk_analysis"
        os.makedirs(self.output_dir, exist_ok=True)

    def generate(self, session_id: str, analysis_result: Dict[str, Any]) -> str:
        """
        生成风险分析报告

        Args:
            session_id: 会话ID
            analysis_result: 分析结果数据

        Returns:
            报告文件路径
        """
        try:
            # 创建Word文档
            doc = Document()

            # 设置文档样式
            self._setup_styles(doc)

            # 添加标题
            self._add_title(doc)

            # 添加元数据（会话信息、时间等）
            self._add_metadata(doc, session_id, analysis_result)

            # 添加总体摘要
            if analysis_result.get('summary'):
                self._add_summary_section(doc, analysis_result['summary'])

            # 添加风险分布
            if analysis_result.get('risk_distribution'):
                self._add_risk_distribution_section(doc, analysis_result['risk_distribution'])

            # 添加风险项列表
            if analysis_result.get('risk_items'):
                self._add_risk_items_section(doc, analysis_result['risk_items'])

            # 添加总体置信度
            if analysis_result.get('total_confidence') is not None:
                self._add_confidence_section(doc, analysis_result['total_confidence'])

            # 保存文档
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"risk_analysis_report_{session_id}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)

            logger.info(f"风险分析报告生成成功: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"风险分析报告生成失败: {str(e)}", exc_info=True)
            raise

    def _setup_styles(self, doc: Document):
        """设置文档样式"""
        # 正文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)

        # 标题1样式
        if 'Heading 1' in doc.styles:
            h1 = doc.styles['Heading 1']
            h1.font.name = '黑体'
            h1.font.size = Pt(16)
            h1.font.bold = True
            h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1.paragraph_format.space_before = Pt(12)
            h1.paragraph_format.space_after = Pt(12)

        # 标题2样式
        if 'Heading 2' in doc.styles:
            h2 = doc.styles['Heading 2']
            h2.font.name = '黑体'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.paragraph_format.space_before = Pt(6)
            h2.paragraph_format.space_after = Pt(6)

        # 标题3样式
        if 'Heading 3' in doc.styles:
            h3 = doc.styles['Heading 3']
            h3.font.name = '黑体'
            h3.font.size = Pt(12)
            h3.font.bold = True
            h3.paragraph_format.space_before = Pt(3)
            h3.paragraph_format.space_after = Pt(3)

    def _add_title(self, doc: Document):
        """添加文档标题"""
        # 添加空行
        doc.add_paragraph()

        # 添加标题
        heading = doc.add_heading("风险评估分析报告", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置标题格式
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(18)
            run.font.bold = True

        # 添加分隔线
        doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_metadata(self, doc: Document, session_id: str, data: Dict[str, Any]):
        """添加元数据"""
        # 添加空行
        doc.add_paragraph()

        # 元数据段落
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts = []
        parts.append(f"会话编号：{session_id}")
        parts.append(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")

        if parts:
            p.add_run('  |  '.join(parts))
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

    def _add_summary_section(self, doc: Document, summary: str):
        """添加总体摘要"""
        doc.add_heading("一、总体摘要", level=2)

        p = doc.add_paragraph(summary)
        p.paragraph_format.first_line_indent = Inches(0.3)
        doc.add_paragraph()

    def _add_risk_distribution_section(self, doc: Document, risk_distribution: Dict):
        """添加风险分布"""
        doc.add_heading("二、风险分布", level=2)

        # 按风险等级统计
        if risk_distribution.get('by_level'):
            doc.add_heading("风险等级分布", level=3)

            for level, count in risk_distribution['by_level'].items():
                p = doc.add_paragraph()
                level_run = p.add_run(f"{level}：")
                level_run.bold = True
                p.add_run(f"{count} 项")

        # 按风险类别统计
        if risk_distribution.get('by_category'):
            doc.add_heading("风险类别分布", level=3)

            for category, count in risk_distribution['by_category'].items():
                p = doc.add_paragraph()
                category_run = p.add_run(f"{category}：")
                category_run.bold = True
                p.add_run(f"{count} 项")

        doc.add_paragraph()

    def _add_risk_items_section(self, doc: Document, risk_items: List[Dict]):
        """添加风险项列表"""
        doc.add_heading("三、风险详情", level=2)

        for i, item in enumerate(risk_items, 1):
            # 风险项标题
            doc.add_heading(f"风险项 {i}：{item.get('title', '未命名')}", level=3)

            # 风险等级
            if item.get('risk_level'):
                p = doc.add_paragraph()
                level_run = p.add_run(f"风险等级：")
                level_run.bold = True
                level_text = p.add_run(item['risk_level'])

                # 根据风险等级设置颜色
                if item['risk_level'] == '高':
                    level_text.font.color.rgb = RGBColor(217, 0, 0)  # 红色
                elif item['risk_level'] == '中':
                    level_text.font.color.rgb = RGBColor(255, 128, 0)  # 橙色
                elif item['risk_level'] == '低':
                    level_text.font.color.rgb = RGBColor(0, 128, 0)  # 绿色

            # 置信度
            if item.get('confidence') is not None:
                p = doc.add_paragraph()
                p.add_run("置信度：").bold = True
                p.add_run(f"{item['confidence'] * 100:.1f}%")

            # 描述
            if item.get('description'):
                p = doc.add_paragraph()
                p.add_run("描述：").bold = True
                desc_para = doc.add_paragraph(item['description'])
                desc_para.paragraph_format.first_line_indent = Inches(0.3)

            # 原因
            if item.get('reasons'):
                p = doc.add_paragraph()
                p.add_run("原因分析：").bold = True
                reasons_para = doc.add_paragraph()
                reasons_para.paragraph_format.first_line_indent = Inches(0.3)
                for reason in item['reasons']:
                    reasons_para.add_run(f"• {reason}\n")

            # 建议
            if item.get('suggestions'):
                p = doc.add_paragraph()
                p.add_run("建议措施：").bold = True
                suggestions_para = doc.add_paragraph()
                suggestions_para.paragraph_format.first_line_indent = Inches(0.3)
                for suggestion in item['suggestions']:
                    suggestions_para.add_run(f"• {suggestion}\n")

            doc.add_paragraph()

    def _add_confidence_section(self, doc: Document, total_confidence: float):
        """添加总体置信度"""
        doc.add_heading("四、总体评估", level=2)

        p = doc.add_paragraph()
        p.add_run("总体置信度：").bold = True
        confidence_text = p.add_run(f"{total_confidence * 100:.1f}%")

        # 根据置信度设置颜色
        if total_confidence >= 0.8:
            confidence_text.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
        elif total_confidence >= 0.6:
            confidence_text.font.color.rgb = RGBColor(255, 128, 0)  # 橙色
        else:
            confidence_text.font.color.rgb = RGBColor(217, 0, 0)  # 红色


# 全局实例
_generator_instance = None


def get_risk_analysis_report_generator() -> RiskAnalysisReportGenerator:
    """获取风险分析报告生成器单例"""
    global _generator_instance
    if _generator_instance is None:
        _generator_instance = RiskAnalysisReportGenerator()
    return _generator_instance
