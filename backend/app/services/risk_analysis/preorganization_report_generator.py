# backend/app/services/preorganization_report_generator.py
"""
预整理报告生成服务

将文档预整理结果生成规范的 Word 报告
"""
import logging
import os
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


class PreorganizationReportGenerator:
    """预整理报告生成器"""

    def __init__(self):
        """初始化报告生成器"""
        self.output_dir = "storage/reports/preorganization"
        os.makedirs(self.output_dir, exist_ok=True)

    def generate(self, session_id: str, preorganization_data: Dict[str, Any]) -> str:
        """
        生成预整理报告

        Args:
            session_id: 会话ID
            preorganization_data: 预整理数据

        Returns:
            报告文件路径
        """
        try:
            # 创建Word文档
            doc = Document()

            # 设置文档样式
            self._setup_styles(doc)

            # 添加标题
            self._add_title(doc, session_id)

            # 添加元数据（会话信息、时间等）
            self._add_metadata(doc, session_id, preorganization_data)

            # 添加用户需求摘要
            if preorganization_data.get('user_requirement_summary'):
                self._add_user_requirement_section(doc, preorganization_data['user_requirement_summary'])

            # 添加文档信息
            if preorganization_data.get('documents_info'):
                self._add_documents_section(doc, preorganization_data['documents_info'])

            # 添加事实摘要
            if preorganization_data.get('fact_summary'):
                self._add_fact_summary_section(doc, preorganization_data['fact_summary'])

            # 添加合同法律特征
            if preorganization_data.get('contract_legal_features'):
                self._add_legal_features_section(doc, preorganization_data['contract_legal_features'])

            # 添加合同关系
            if preorganization_data.get('contract_relationships'):
                self._add_relationships_section(doc, preorganization_data['contract_relationships'])

            # 添加架构图
            if preorganization_data.get('architecture_diagram'):
                self._add_diagram_section(doc, preorganization_data['architecture_diagram'])

            # 添加增强分析结果（交易全景、主体画像、时间线等）
            if preorganization_data.get('enhanced_analysis_json'):
                self._add_enhanced_analysis_section(doc, preorganization_data['enhanced_analysis_json'])

            # 保存文档
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"preorganization_report_{session_id}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)

            logger.info(f"预整理报告生成成功: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"预整理报告生成失败: {str(e)}", exc_info=True)
            raise

    def _setup_styles(self, doc: Document):
        """设置文档样式"""
        # 正文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)

        # 标题1样式
        if 'Heading 1' in doc.styles:
            h1 = doc.styles['Heading 1']
            h1.font.name = '黑体'
            h1.font.size = Pt(16)
            h1.font.bold = True
            h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1.paragraph_format.space_before = Pt(12)
            h1.paragraph_format.space_after = Pt(12)

        # 标题2样式
        if 'Heading 2' in doc.styles:
            h2 = doc.styles['Heading 2']
            h2.font.name = '黑体'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.paragraph_format.space_before = Pt(6)
            h2.paragraph_format.space_after = Pt(6)

        # 标题3样式
        if 'Heading 3' in doc.styles:
            h3 = doc.styles['Heading 3']
            h3.font.name = '黑体'
            h3.font.size = Pt(12)
            h3.font.bold = True
            h3.paragraph_format.space_before = Pt(3)
            h3.paragraph_format.space_after = Pt(3)

    def _add_title(self, doc: Document, session_id: str):
        """添加文档标题"""
        # 添加空行
        doc.add_paragraph()

        # 添加标题
        heading = doc.add_heading("文档预整理报告", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置标题格式
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(18)
            run.font.bold = True

        # 添加会话ID
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"会话编号：{session_id}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # 添加分隔线
        doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_metadata(self, doc: Document, session_id: str, data: Dict[str, Any]):
        """添加元数据"""
        # 添加空行
        doc.add_paragraph()

        # 元数据段落
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts = []
        parts.append(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")

        if parts:
            p.add_run('  |  '.join(parts))
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

    def _add_user_requirement_section(self, doc: Document, summary: str):
        """添加用户需求摘要"""
        doc.add_heading("一、用户需求摘要", level=2)

        p = doc.add_paragraph(summary)
        p.paragraph_format.first_line_indent = Inches(0.3)
        doc.add_paragraph()

    def _add_documents_section(self, doc: Document, documents_info):
        """添加文档信息"""
        import json

        doc.add_heading("二、文档信息", level=2)

        # 处理 documents_info 可能是字符串的情况
        if isinstance(documents_info, str):
            try:
                documents_info = json.loads(documents_info)
            except json.JSONDecodeError:
                logger.error(f"documents_info 是字符串但无法解析为 JSON: {documents_info[:200]}")
                return

        # 确保是列表
        if not isinstance(documents_info, list):
            logger.error(f"documents_info 不是列表类型: {type(documents_info)}")
            return

        for i, doc_info in enumerate(documents_info, 1):
            # 处理 doc_info 可能是字符串的情况
            if isinstance(doc_info, str):
                try:
                    doc_info = json.loads(doc_info)
                except json.JSONDecodeError:
                    logger.warning(f"doc_info[{i}] 是字符串但无法解析，跳过")
                    continue

            # 确保 doc_info 是字典
            if not isinstance(doc_info, dict):
                logger.warning(f"doc_info[{i}] 不是字典类型: {type(doc_info)}")
                continue

            # 文档标题
            doc.add_heading(f"文档{i}：{doc_info.get('file_name', '未命名')}", level=3)

            # 文档元数据
            if doc_info.get('document_type'):
                p = doc.add_paragraph()
                p.add_run("文档类型：").bold = True
                p.add_run(doc_info['document_type'])

            if doc_info.get('page_count'):
                p = doc.add_paragraph()
                p.add_run("页数：").bold = True
                p.add_run(str(doc_info['page_count']))

            # 文档摘要
            if doc_info.get('summary'):
                p = doc.add_paragraph()
                p.add_run("摘要：").bold = True
                summary_para = doc.add_paragraph(doc_info['summary'])
                summary_para.paragraph_format.first_line_indent = Inches(0.3)

            # 质量评估
            if doc_info.get('quality_assessment'):
                qa = doc_info['quality_assessment']
                p = doc.add_paragraph()
                p.add_run("质量评估：").bold = True

                if qa.get('overall_score'):
                    p.add_run(f"综合评分：{qa['overall_score']}/100  ")

                if qa.get('clarity'):
                    p.add_run(f"清晰度：{qa['clarity']}  ")

                if qa.get('completeness'):
                    p.add_run(f"完整度：{qa['completeness']}")

            doc.add_paragraph()

    def _add_fact_summary_section(self, doc: Document, fact_summary: Dict):
        """添加事实摘要"""
        doc.add_heading("三、事实摘要", level=2)

        # 关键事实
        if fact_summary.get('key_facts'):
            doc.add_heading("关键事实", level=3)
            for fact in fact_summary['key_facts']:
                p = doc.add_paragraph(f"• {fact}", style='List Bullet')

        # 时间线
        if fact_summary.get('timeline'):
            doc.add_heading("时间线", level=3)
            for event in fact_summary['timeline']:
                p = doc.add_paragraph()
                p.add_run(f"{event.get('date', '')}：").bold = True
                p.add_run(event.get('event', ''))

        # 涉及主体
        if fact_summary.get('entities'):
            doc.add_heading("涉及主体", level=3)
            for entity in fact_summary['entities']:
                p = doc.add_paragraph(f"• {entity}", style='List Bullet')

        doc.add_paragraph()

    def _add_legal_features_section(self, doc: Document, features: Dict):
        """添加合同法律特征"""
        doc.add_heading("四、合同法律特征", level=2)

        # 合同类型
        if features.get('contract_types'):
            p = doc.add_paragraph()
            p.add_run("合同类型：").bold = True
            p.add_run('、'.join(features['contract_types']))

        # 法律条款
        if features.get('legal_clauses'):
            doc.add_heading("法律条款", level=3)
            for clause in features['legal_clauses']:
                p = doc.add_paragraph(f"• {clause}", style='List Bullet')

        # 风险点
        if features.get('risk_points'):
            doc.add_heading("识别的风险点", level=3)
            for risk in features['risk_points']:
                p = doc.add_paragraph()
                risk_run = p.add_run(f"风险：{risk.get('description', '未描述')}")
                risk_run.font.color.rgb = RGBColor(217, 0, 0)  # 红色
                if risk.get('suggestion'):
                    p = doc.add_paragraph()
                    p.add_run("建议：").bold = True
                    p.add_run(risk['suggestion'])

        doc.add_paragraph()

    def _add_relationships_section(self, doc: Document, relationships: List[Dict]):
        """添加合同关系"""
        doc.add_heading("五、合同关系", level=2)

        for rel in relationships:
            p = doc.add_paragraph()

            # 关系描述
            if rel.get('source') and rel.get('target'):
                p.add_run(f"{rel['source']} → {rel['target']}").bold = True

            if rel.get('relationship_type'):
                p.add_run(f"（{rel['relationship_type']}）")

            # 关系说明
            if rel.get('description'):
                desc_para = doc.add_paragraph(rel['description'])
                desc_para.paragraph_format.left_indent = Inches(0.3)

        doc.add_paragraph()

    def _add_diagram_section(self, doc: Document, diagram: Dict):
        """添加架构图"""
        doc.add_heading("六、架构图", level=2)

        # 图表描述
        if diagram.get('description'):
            p = doc.add_paragraph()
            p.add_run("图表说明：").bold = True
            p.add_run(diagram['description'])

        # 图表数据（文本表示）
        if diagram.get('nodes'):
            doc.add_heading("节点", level=3)
            for node in diagram['nodes']:
                p = doc.add_paragraph(f"• {node.get('label', '未命名')}（{node.get('type', '未知类型')}）", style='List Bullet')

        if diagram.get('edges'):
            doc.add_heading("连接", level=3)
            for edge in diagram['edges']:
                p = doc.add_paragraph(f"• {edge.get('from', '')} → {edge.get('to', '')}：{edge.get('label', '')}", style='List Bullet')

    def _add_enhanced_analysis_section(self, doc: Document, enhanced_analysis_json: str):
        """
        添加增强分析章节

        包含：交易全景、争议焦点、主体画像、时间线
        """
        import json
        from typing import Dict, List

        try:
            enhanced_data = json.loads(enhanced_analysis_json)
        except json.JSONDecodeError:
            logger.error("enhanced_analysis_json 解析失败")
            return

        # 1. 交易全景
        if enhanced_data.get("transaction_summary"):
            doc.add_heading("七、交易全景", level=2)

            # 交易故事叙述
            p = doc.add_paragraph(enhanced_data["transaction_summary"])
            p.paragraph_format.first_line_indent = Inches(0.3)

            # 合同状态
            if enhanced_data.get("contract_status"):
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("合同状态：").bold = True
                p.add_run(enhanced_data["contract_status"])

                # 根据状态设置颜色
                status_color_map = {
                    '磋商': RGBColor(0, 102, 204),
                    '履约': RGBColor(0, 153, 76),
                    '违约': RGBColor(217, 0, 0),
                    '终止': RGBColor(128, 128, 128)
                }
                status_color = status_color_map.get(enhanced_data["contract_status"], RGBColor(0, 0, 0))
                for run in p.runs:
                    if enhanced_data["contract_status"] in run.text:
                        run.font.color.rgb = status_color

            doc.add_paragraph()

        # 2. 争议焦点
        if enhanced_data.get("dispute_focus"):
            doc.add_heading("八、争议焦点", level=2)
            p = doc.add_paragraph(enhanced_data["dispute_focus"])
            p.paragraph_format.first_line_indent = Inches(0.3)
            doc.add_paragraph()

        # 3. 主体画像
        if enhanced_data.get("parties") and len(enhanced_data["parties"]) > 0:
            doc.add_heading("九、主体画像", level=2)

            for idx, party in enumerate(enhanced_data["parties"], 1):
                # 主体名称和角色
                doc.add_heading(f"{idx}. {party.get('role', '未知角色')}: {party.get('name', '未命名')}", level=3)

                # 核心义务
                if party.get("obligations") and len(party["obligations"]) > 0:
                    p = doc.add_paragraph()
                    p.add_run("核心义务：").bold = True
                    for obligation in party["obligations"]:
                        p = doc.add_paragraph(f"• {obligation}", style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.2)

                # 核心权利
                if party.get("rights") and len(party["rights"]) > 0:
                    p = doc.add_paragraph()
                    p.add_run("核心权利：").bold = True
                    for right in party["rights"]:
                        p = doc.add_paragraph(f"• {right}", style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.2)

                # 风险敞口
                if party.get("risk_exposure"):
                    p = doc.add_paragraph()
                    p.add_run("风险敞口：").bold = True
                    risk_run = p.add_run(party["risk_exposure"])

                    # 根据风险等级设置颜色
                    if '高' in party["risk_exposure"]:
                        risk_run.font.color.rgb = RGBColor(217, 0, 0)  # 红色
                    elif '中' in party["risk_exposure"]:
                        risk_run.font.color.rgb = RGBColor(255, 153, 0)  # 橙色
                    else:
                        risk_run.font.color.rgb = RGBColor(0, 153, 76)  # 绿色

                doc.add_paragraph()

        # 4. 时间线
        if enhanced_data.get("timeline") and len(enhanced_data["timeline"]) > 0:
            doc.add_heading("十、时间线", level=2)

            for event in enhanced_data["timeline"]:
                # 日期和事件类型
                p = doc.add_paragraph()
                date_run = p.add_run(f"{event.get('date', '日期未知')} - {event.get('type', '事件')}")
                date_run.bold = True
                date_run.font.color.rgb = RGBColor(0, 102, 204)

                # 事件描述
                if event.get("event"):
                    p = doc.add_paragraph(event["event"])
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.first_line_indent = Inches(0.2)

                # 来源文档
                if event.get("source_doc"):
                    p = doc.add_paragraph()
                    source_run = p.add_run(f"来源：{event['source_doc']}")
                    source_run.font.size = Pt(10)
                    source_run.font.color.rgb = RGBColor(128, 128, 128)
                    p.paragraph_format.left_indent = Inches(0.3)

                doc.add_paragraph()


# 全局实例
_generator_instance = None


def get_preorganization_report_generator() -> PreorganizationReportGenerator:
    """获取预整理报告生成器单例"""
    global _generator_instance
    if _generator_instance is None:
        _generator_instance = PreorganizationReportGenerator()
    return _generator_instance
