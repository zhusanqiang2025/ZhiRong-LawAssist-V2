# backend/templates/documents/create_templates.py
"""创建 Word 文档模板"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os


def create_contract_template():
    """创建合同模板"""
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    # 1. 标题
    title = doc.add_heading('{{title}}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '黑体'
        run.font.size = Pt(18)
        run.font.bold = True

    # 分隔线
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. 元数据信息
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run('合同编号：{{contract_number}}    日期：{{date}}')
    meta_run.font.size = Pt(10)

    doc.add_paragraph()

    # 3. 合同方信息
    doc.add_heading('合同各方', level=2)

    # 甲方
    p1 = doc.add_paragraph()
    r1 = p1.add_run('{{party1_role}}：')
    r1.bold = True
    p1.add_run('{{party1_name}}')

    # 乙方
    p2 = doc.add_paragraph()
    r2 = p2.add_run('{{party2_role}}：')
    r2.bold = True
    p2.add_run('{{party2_name}}')

    doc.add_paragraph()

    # 4. 章节内容占位符
    doc.add_heading('合同条款', level=2)
    content_para = doc.add_paragraph('{{sections_content}}')

    # 5. 签署区
    for _ in range(3):
        doc.add_paragraph()

    doc.add_heading('签署', level=2)

    # 创建签署表格
    table = doc.add_table(rows=3, cols=4)
    table.width = Inches(6)

    # 甲方
    table.rows[0].cells[0].text = '{{party1_role}}（盖章）：'
    table.rows[1].cells[0].text = '代表签字：'
    table.rows[2].cells[0].text = '日期：'

    # 乙方
    table.rows[0].cells[2].text = '{{party2_role}}（盖章）：'
    table.rows[1].cells[2].text = '代表签字：'
    table.rows[2].cells[2].text = '日期：'

    # 设置单元格样式
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(11)

    # 保存
    output_path = 'backend/templates/documents/contract_template.docx'
    doc.save(output_path)
    print(f'[OK] Contract template created: {output_path}')


def create_letter_template():
    """创建函件模板"""
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    # 1. 标题
    title = doc.add_heading('{{title}}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '黑体'
        run.font.size = Pt(18)
        run.font.bold = True

    doc.add_paragraph()

    # 2. 称谓
    salutation = doc.add_paragraph()
    salutation.add_run('尊敬的{{recipient_name}}：')

    doc.add_paragraph()

    # 3. 正文内容
    content_para = doc.add_paragraph('{{sections_content}}')

    # 4. 结语
    for _ in range(2):
        doc.add_paragraph()

    closing = doc.add_paragraph()
    closing.add_run('此致')
    closing.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run('敬礼！')
    p.paragraph_format.space_after = Pt(24)

    # 5. 签名
    p = doc.add_paragraph()
    p.add_run('___________________')
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    signature_run = p.add_run('签名：')
    signature_run.bold = True
    p.add_run('___________________')
    p.paragraph_format.space_after = Pt(6)

    # 6. 日期
    p = doc.add_paragraph()
    date_run = p.add_run('日期：')
    date_run.bold = True
    p.add_run('{{date}}')

    # 保存
    output_path = 'backend/templates/documents/letter_template.docx'
    doc.save(output_path)
    print(f'[OK] Letter template created: {output_path}')


def create_simple_template():
    """创建简单文档模板"""
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    # 1. 标题
    title = doc.add_heading('{{title}}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '黑体'
        run.font.size = Pt(18)
        run.font.bold = True

    # 2. 元数据
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run('日期：{{date}}')
    meta_run.font.size = Pt(10)

    doc.add_paragraph()

    # 3. 内容
    doc.add_heading('内容', level=2)
    content_para = doc.add_paragraph('{{sections_content}}')

    # 保存
    output_path = 'backend/templates/documents/simple_template.docx'
    doc.save(output_path)
    print(f'[OK] Simple template created: {output_path}')


def create_judicial_template():
    """创建司法文书模板"""
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)

    # 1. 法院名称
    court = doc.add_paragraph()
    court.alignment = WD_ALIGN_PARAGRAPH.CENTER
    court_run = court.add_run('{{court_name}}')
    court_run.font.name = '黑体'
    court_run.font.size = Pt(16)
    court_run.font.bold = True

    doc.add_paragraph()

    # 2. 文书类型和标题
    title = doc.add_heading('{{title}}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '黑体'
        run.font.size = Pt(18)
        run.font.bold = True

    # 3. 案号
    case_no = doc.add_paragraph()
    case_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    case_no_run = case_no.add_run('案号：{{case_number}}')
    case_no_run.font.size = Pt(11)

    doc.add_paragraph()

    # 4. 当事人信息
    doc.add_heading('当事人', level=2)
    parties_para = doc.add_paragraph('{{parties_info}}')

    doc.add_paragraph()

    # 5. 案件内容
    doc.add_heading('案件内容', level=2)
    content_para = doc.add_paragraph('{{sections_content}}')

    # 6. 落款
    for _ in range(3):
        doc.add_paragraph()

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.add_run('{{court_name}}')
    closing.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('{{date}}')

    # 保存
    output_path = 'backend/templates/documents/judicial_template.docx'
    doc.save(output_path)
    print(f'[OK] Judicial template created: {output_path}')


if __name__ == '__main__':
    print('Creating Word document templates...\n')

    create_contract_template()
    create_letter_template()
    create_simple_template()
    create_judicial_template()

    print('\n[OK] All templates created successfully!')
